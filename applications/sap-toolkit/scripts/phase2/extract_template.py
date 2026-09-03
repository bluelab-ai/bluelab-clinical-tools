#!/usr/bin/env python3
"""
TFL模版库表格提取工具
从模版库文档中根据代码提取对应的表格到新文档
"""

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys
import os
import copy


def get_template_path():
    """获取模版库文件路径"""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(script_dir, "模版.docx")


def load_template_library(template_path=None):
    """加载模版库并建立代码到表格的映射"""
    if template_path is None:
        template_path = get_template_path()

    doc = docx.Document(template_path)
    code_pattern = re.compile(r'^[A-Z]+_[A-Z0-9_]+$')

    elements = []
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            elements.append(('paragraph', element))
        elif tag == 'tbl':
            elements.append(('table', element))

    code_to_table_idx = {}
    table_idx = 0
    current_code = None

    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            text_elems = elem.findall('.//w:t', ns)
            text = ''.join([t.text for t in text_elems if t.text]).strip()

            if code_pattern.match(text):
                current_code = text
        elif elem_type == 'table' and current_code is not None:
            code_to_table_idx[current_code] = table_idx
            current_code = None
            table_idx += 1

    return doc, code_to_table_idx


def get_available_codes(template_path=None):
    """获取所有可用的代码列表"""
    _, code_to_table_idx = load_template_library(template_path)
    return sorted(code_to_table_idx.keys())


def extract_table_by_code(code, table_name=None, template_path=None, output_path=None):
    """
    根据代码提取表格到新文档

    参数:
        code: 表格代码，如 D_G2_ST
        table_name: 表名，显示在表格上方
        template_path: 模版库文件路径
        output_path: 输出文件路径，默认为 {code}.docx
    """
    if template_path is None:
        template_path = get_template_path()

    template_doc, code_to_table_idx = load_template_library(template_path)

    if code not in code_to_table_idx:
        print(f"错误: 代码 '{code}' 不存在于模版库中", file=sys.stderr)
        return None

    table_idx = code_to_table_idx[code]
    source_table = template_doc.tables[table_idx]

    # 深拷贝整个表格XML
    source_tbl_xml = source_table._tbl
    copied_tbl_xml = copy.deepcopy(source_tbl_xml)

    # 修改表格字体设置
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for tc in copied_tbl_xml.findall('.//' + qn('w:tc')):
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)

                # 设置字体
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:cs'), 'Times New Roman')

                # 设置字号（5号 = 10.5pt = 21）
                sz = rPr.find(qn('w:sz'))
                if sz is None:
                    sz = OxmlElement('w:sz')
                    rPr.append(sz)
                sz.set(qn('w:val'), '21')

                szCs = rPr.find(qn('w:szCs'))
                if szCs is None:
                    szCs = OxmlElement('w:szCs')
                    rPr.append(szCs)
                szCs.set(qn('w:val'), '21')

    # 创建新文档
    new_doc = docx.Document()

    # 设置页边距（左右1.91cm = 687600 EMUs）
    section = new_doc.sections[0]
    section.left_margin = 687600
    section.right_margin = 687600

    # 将表格XML直接插入到文档body中
    body = new_doc.element.body

    # 如果有表名，在表格之前添加表名段落
    if table_name:
        name_para_xml = docx.oxml.OxmlElement('w:p')
        pPr = docx.oxml.OxmlElement('w:pPr')
        jc = docx.oxml.OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)

        spacing = docx.oxml.OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'atLeast')
        pPr.append(spacing)

        name_para_xml.append(pPr)

        run_xml = docx.oxml.OxmlElement('w:r')
        rPr = docx.oxml.OxmlElement('w:rPr')
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '宋体')
        rFonts.set(qn('w:hAnsi'), '宋体')
        rFonts.set(qn('w:eastAsia'), '宋体')
        sz = docx.oxml.OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        szCs = docx.oxml.OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '24')
        rPr.append(rFonts)
        rPr.append(sz)
        rPr.append(szCs)
        run_xml.append(rPr)

        t = docx.oxml.OxmlElement('w:t')
        t.text = table_name
        run_xml.append(t)
        name_para_xml.append(run_xml)

        sectPr = body.find(qn('w:sectPr'))
        if sectPr is not None:
            sectPr.addprevious(name_para_xml)
        else:
            body.append(name_para_xml)

    # 插入表格
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(copied_tbl_xml)
    else:
        body.append(copied_tbl_xml)

    if output_path is None:
        output_path = f"{code}.docx"

    new_doc.save(output_path)
    print(f"✅ 提取表格: {code} → {output_path}", file=sys.stderr)
    return output_path


def extract_multiple_tables(codes, table_names=None, template_path=None, output_path="extracted_tables.docx"):
    """
    批量提取多个表格到同一个文档

    参数:
        codes: 代码列表
        table_names: 表名列表（可选），与codes一一对应
        template_path: 模版库文件路径
        output_path: 输出文件路径
    """
    if template_path is None:
        template_path = get_template_path()

    template_doc, code_to_table_idx = load_template_library(template_path)
    new_doc = docx.Document()

    # 设置页边距（左右1.91cm）
    section = new_doc.sections[0]
    section.left_margin = 687600
    section.right_margin = 687600

    body = new_doc.element.body
    extracted_count = 0

    for i, code in enumerate(codes):
        if code not in code_to_table_idx:
            print(f"⚠️ 代码 '{code}' 不存在，跳过", file=sys.stderr)
            continue

        table_idx = code_to_table_idx[code]
        source_table = template_doc.tables[table_idx]

        # 添加间距（除了第一个表格）
        if extracted_count > 0:
            spacing_para = OxmlElement('w:p')
            spacing_pPr = OxmlElement('w:pPr')
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '1200')
            spacing.set(qn('w:after'), '0')
            spacing_pPr.append(spacing)
            spacing_para.append(spacing_pPr)
            body.append(spacing_para)

        # 添加表名
        if table_names and i < len(table_names) and table_names[i]:
            name_para = OxmlElement('w:p')
            name_pPr = OxmlElement('w:pPr')
            name_jc = OxmlElement('w:jc')
            name_jc.set(qn('w:val'), 'left')
            name_pPr.append(name_jc)

            name_spacing = OxmlElement('w:spacing')
            name_spacing.set(qn('w:before'), '0')
            name_spacing.set(qn('w:after'), '0')
            name_spacing.set(qn('w:line'), '240')
            name_spacing.set(qn('w:lineRule'), 'atLeast')
            name_pPr.append(name_spacing)

            name_para.append(name_pPr)

            name_run = OxmlElement('w:r')
            name_rPr = OxmlElement('w:rPr')
            name_rFonts = OxmlElement('w:rFonts')
            name_rFonts.set(qn('w:ascii'), '宋体')
            name_rFonts.set(qn('w:hAnsi'), '宋体')
            name_rFonts.set(qn('w:eastAsia'), '宋体')
            name_sz = OxmlElement('w:sz')
            name_sz.set(qn('w:val'), '24')
            name_szCs = OxmlElement('w:szCs')
            name_szCs.set(qn('w:val'), '24')
            name_rPr.append(name_rFonts)
            name_rPr.append(name_sz)
            name_rPr.append(name_szCs)
            name_run.append(name_rPr)

            name_text = OxmlElement('w:t')
            name_text.text = table_names[i]
            name_run.append(name_text)
            name_para.append(name_run)
            body.append(name_para)

        # 复制表格并设置字体
        copied_tbl = copy.deepcopy(source_table._tbl)

        for tc in copied_tbl.findall('.//' + qn('w:tc')):
            for p in tc.findall(qn('w:p')):
                for r in p.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)

                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.insert(0, rFonts)
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rFonts.set(qn('w:eastAsia'), '宋体')
                    rFonts.set(qn('w:cs'), 'Times New Roman')

                    sz = rPr.find(qn('w:sz'))
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        rPr.append(sz)
                    sz.set(qn('w:val'), '21')

                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is None:
                        szCs = OxmlElement('w:szCs')
                        rPr.append(szCs)
                    szCs.set(qn('w:val'), '21')

        body.append(copied_tbl)
        extracted_count += 1

    new_doc.save(output_path)
    print(f"✅ 提取 {extracted_count}/{len(codes)} 个表格 → {output_path}", file=sys.stderr)
    return output_path


def extract_tables_from_json(json_path, output_path=None):
    """
    从模板代码结果JSON中提取所有表格

    参数:
        json_path: 模板代码结果JSON文件路径
        output_path: 输出文件路径
    """
    import json

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    tables = data.get('tables', [])

    codes = []
    names = []
    for t in tables:
        code = t.get('template_code', '')
        name = t.get('name', '')

        if isinstance(code, list):
            codes.extend(code)
            names.extend([name] * len(code))
        elif code:
            codes.append(code)
            names.append(name)

    if not codes:
        print("没有找到有效的模板代码", file=sys.stderr)
        return None

    if output_path is None:
        output_path = os.path.join(os.path.dirname(json_path), "提取的表格.docx")

    return extract_multiple_tables(codes, names, output_path=output_path)


def extract_tables_to_folder(json_path, output_dir=None):
    """
    从模板代码结果JSON中提取所有表格到单独的文件

    参数:
        json_path: 模板代码结果JSON文件路径
        output_dir: 输出目录路径
    """
    import json

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    tables = data.get('tables', [])

    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(json_path), "提取的表格")

    os.makedirs(output_dir, exist_ok=True)

    template_path = get_template_path()
    extracted = 0
    skipped = 0

    for t in tables:
        code = t.get('template_code', '')
        name = t.get('name', '')

        if not code:
            continue

        safe_name = name.replace("/", "_").replace("\\", "_")
        output_path = os.path.join(output_dir, f"{safe_name}.docx")

        if isinstance(code, list):
            # 多模板合并到一个文档
            result = extract_multiple_tables(code, [name] * len(code), template_path=template_path, output_path=output_path)
            if result:
                extracted += 1
            else:
                skipped += 1
        else:
            result = extract_table_by_code(code, name, template_path, output_path)
            if result:
                extracted += 1
            else:
                skipped += 1

    print(f"\n✅ 提取完成: {extracted} 个表格 → {output_dir}", file=sys.stderr)
    if skipped > 0:
        print(f"⚠️ 跳过 {skipped} 个（代码不存在）", file=sys.stderr)

    return output_dir


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="TFL模版库表格提取工具")
    parser.add_argument("codes", nargs="*", help="表格代码")
    parser.add_argument("--name", action="append", help="表名（可多次使用）")
    parser.add_argument("--output", "-o", help="输出文件路径")
    parser.add_argument("--list", action="store_true", help="列出所有可用代码")
    parser.add_argument("--json", help="从模板代码结果JSON提取所有表格到一个文件")
    parser.add_argument("--folder", help="从模板代码结果JSON提取每个表格到单独文件")

    args = parser.parse_args()

    if args.list:
        codes = get_available_codes()
        print(f"模版库中共有 {len(codes)} 个表格代码:")
        for i, code in enumerate(codes, 1):
            print(f"  {i:3d}. {code}")
    elif args.folder:
        extract_tables_to_folder(args.folder, args.output)
    elif args.json:
        extract_tables_from_json(args.json, args.output)
    elif len(args.codes) == 1:
        table_name = args.name[0] if args.name else None
        extract_table_by_code(args.codes[0], table_name, output_path=args.output)
    elif len(args.codes) > 1:
        extract_multiple_tables(args.codes, args.name, output_path=args.output or "extracted_tables.docx")
    else:
        parser.print_help()
