#!/usr/bin/env python3
"""
表格合并工具
根据 tables.json 合并所有表格到一个 Word 文档
- 已填充的表格使用填充版本
- 未填充的表格使用模板版本
- 自动添加表号（表1, 表2, ...）
"""

import json
import os
import sys
from datetime import datetime

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def log(msg):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}", file=sys.stderr)


def add_table_of_contents(doc):
    """Add a Word TOC field that includes Heading 1 through Heading 4."""
    for level in range(1, 5):
        style_name = f"TOC {level}"
        try:
            toc_style = doc.styles[style_name]
        except KeyError:
            toc_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        toc_style.paragraph_format.line_spacing = 1.0
        toc_style.paragraph_format.space_before = docx.shared.Pt(0)
        toc_style.paragraph_format.space_after = docx.shared.Pt(0)

    toc_heading = doc.add_paragraph("目录", style="TOC Heading")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading_run = toc_heading.runs[0]
    toc_heading_run.font.name = "Times New Roman"
    toc_heading_run.font.size = docx.shared.Pt(16)
    toc_heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    toc_paragraph = doc.add_paragraph()
    toc_run = toc_paragraph.add_run()._r
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-4" \\h \\z \\u'
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    field_result = OxmlElement("w:t")
    field_result.text = "目录将在 Microsoft Word 中打开时自动更新"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    toc_run.extend([field_begin, instruction, field_separate, field_result, field_end])

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def create_parent_header(doc, parent_name):
    """
    创建父标题段落（二级标题）

    参数:
        doc: Word 文档
        parent_name: 父标题名称
    """
    # 使用二级标题样式
    para = doc.add_paragraph(style='Heading 2')

    # 段落属性
    pPr = para._element.get_or_add_pPr()

    # 左对齐
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)

    # 段前段后间距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '360')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'atLeast')
    pPr.append(spacing)

    # 添加父标题名称
    run = para.add_run(parent_name)
    run.bold = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = docx.shared.Pt(16)

    return para


def create_category_header(doc, category_name):
    """
    创建分类标题段落（三级标题）

    参数:
        doc: Word 文档
        category_name: 分类名称
    """
    # 使用三级标题样式
    para = doc.add_paragraph(style='Heading 3')

    # 段落属性
    pPr = para._element.get_or_add_pPr()

    # 左对齐
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)

    # 段前段后间距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'atLeast')
    pPr.append(spacing)

    # 添加分类名称
    run = para.add_run(category_name)
    run.bold = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = docx.shared.Pt(14)

    return para


def create_table_caption(doc, table_name, table_index):
    """
    创建表题段落（表1 表名）

    参数:
        doc: Word 文档
        table_name: 表格名称
        table_index: 表格序号
    """
    # 使用四级标题样式
    para = doc.add_paragraph(style='Heading 4')

    # 段落属性
    pPr = para._element.get_or_add_pPr()

    # 左对齐
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)

    # 段后间距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'atLeast')
    pPr.append(spacing)

    # 添加表号（加粗）
    run_index = para.add_run(f"表{table_index} ")
    run_index.bold = True
    run_index.font.name = 'Times New Roman'
    run_index._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_index.font.size = docx.shared.Pt(10.5)

    # 添加表名
    run_name = para.add_run(table_name)
    run_name.font.name = 'Times New Roman'
    run_name._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_name.font.size = docx.shared.Pt(10.5)

    return para


def copy_table_from_doc(source_doc, target_doc, table_idx):
    """
    从源文档复制表格到目标文档

    参数:
        source_doc: 源文档
        target_doc: 目标文档
        table_idx: 表格索引
    """
    import copy

    source_table = source_doc.tables[table_idx]
    copied_tbl = copy.deepcopy(source_table._tbl)

    # 设置字体
    for tc in copied_tbl.findall('.//' + qn('w:tc')):
        for p in tc.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)

                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:cs'), 'Times New Roman')

                sz = rPr.find(qn('w:sz'))
                if sz is None:
                    sz = OxmlElement('w:sz')
                    rPr.append(sz)
                sz.set(qn('w:val'), '21')

                szCs = rPr.find(qn('w:szCs'))
                if szCs is None:
                    szCs = OxmlElement('w:szCs')
                    rPr.append(szCs)
                szCs.set(qn('w:val'), '21')

    # 插入到文档
    body = target_doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(copied_tbl)
    else:
        body.append(copied_tbl)


def create_placeholder_table(doc, table_name):
    """
    创建占位提示（用于未提取的表格）

    参数:
        doc: Word 文档
        table_name: 表格名称
    """
    # 创建提示段落
    para = doc.add_paragraph()

    # 设置段落属性
    pPr = para._element.get_or_add_pPr()

    # 左对齐
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'left')
    pPr.append(jc)

    # 段前段后间距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:after'), '120')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'atLeast')
    pPr.append(spacing)

    # 添加提示文字
    run = para.add_run('未能在CRF中找到分析项目')
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = docx.shared.Pt(10.5)
    run.italic = True  # 斜体显示

    return para


def get_parent_header(category, fallback="安全性指标"):
    """
    根据分类名称获取父标题

    参数:
        category: 分类名称
        fallback: 未匹配时的兜底父标题

    返回:
        父标题名称
    """
    # 分类到父标题的映射
    category_to_parent = {
        "病例分布（随机化人群）": "各中心病例分布和人群划分情况",
        "人口学信息": "人口学信息和基线资料",
        "病史": "人口学信息和基线资料",
        "基线信息": "人口学信息和基线资料",
        "手术日当天信息": "手术日当天信息",
        "主要疗效终点分析": "疗效指标",
        "次要疗效终点分析": "疗效指标",
        "不良事件（SS）": "安全性指标",
        "实验室检查（SS）": "安全性指标",
        "生命体征（SS）": "安全性指标",
        "心电图检查（SS）": "安全性指标",
        "合并用药（SS）": "安全性指标",
        "器械缺陷（SS）": "安全性指标",
    }

    return category_to_parent.get(category, fallback)


def merge_tables(output_dir, output_path=None, filled_dir_name="填充的表格"):
    """
    合并所有表格

    参数:
        output_dir: 阶段2输出目录
        output_path: 输出文件路径
        filled_dir_name: 填充表格目录名（默认"填充的表格"）
    """
    # 读取 tables.json
    tables_json = os.path.join(output_dir, "tables.json")
    if not os.path.exists(tables_json):
        log(f"❌ tables.json 不存在: {tables_json}")
        return None

    with open(tables_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    tables = data.get('tables', [])
    log(f"📋 tables.json: {len(tables)} 张表")

    # 读取第二个访视名称，用于生成兜底父标题
    fallback_parent = "安全性指标"
    second_visit_file = os.path.join(output_dir, "02_内容提取", "试验流程_第二个访视项目.json")
    if os.path.exists(second_visit_file):
        try:
            with open(second_visit_file, 'r', encoding='utf-8') as f:
                visit_data = json.load(f)
            visit_name = visit_data.get("visit_name", "")
            if visit_name:
                fallback_parent = f"{visit_name}信息"
                log(f"📋 第二个访视: {visit_name}，兜底父标题: {fallback_parent}")
        except Exception as e:
            log(f"⚠️ 读取第二个访视文件失败: {e}")

    # 目录路径
    filled_dir = os.path.join(output_dir, filled_dir_name)
    template_dir = os.path.join(output_dir, "模版")

    # 创建输出文档
    doc = docx.Document()

    # 设置页边距
    section = doc.sections[0]
    section.left_margin = 687600  # 1.91cm
    section.right_margin = 687600

    # 目录独占首页；正文一级标题从下一页开始。
    add_table_of_contents(doc)
    doc.add_page_break()

    # 一级编号占位符，后续可在 Word 中统一替换为 SAP 正文实际编号。
    top_level_prefix = "$"

    # 添加顶级标题（一级标题）
    top_para = doc.add_paragraph(style='Heading 1')
    top_para.paragraph_format.space_after = docx.shared.Pt(0)
    top_run = top_para.add_run(f"{top_level_prefix} 统计分析结果")
    top_run.bold = True
    top_run.font.name = 'Times New Roman'
    top_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    top_run.font.size = docx.shared.Pt(18)

    # 统计
    used_filled = 0
    used_template = 0
    used_placeholder = 0

    # 当前分类和父标题（用于添加分类标题和父标题）
    current_category = None
    current_parent = None

    # 编号计数器
    parent_counter = 0  # 二级标题计数器 (1.1, 1.2, ...)
    category_counter = 0  # 三级标题计数器 (1.1.1, 1.1.2, ...)
    table_in_category_counter = 0  # 四级表格计数器 (1.1.1.1, 1.1.1.2, ...)

    # 合并表格
    for i, table in enumerate(tables):
        table_name = table.get('name', '')
        table_category = table.get('category', '')
        table_index = i + 1

        # 检查是否需要添加分类标题
        category_changed = table_category and table_category != current_category
        if category_changed:
            # 获取父标题
            table_parent = get_parent_header(table_category, fallback=fallback_parent)

            # 新分类紧跟上一张表时，分页位于二级或三级标题之前。
            if i > 0:
                doc.add_page_break()

            # 检查是否需要添加父标题
            if table_parent != current_parent:
                # 更新二级标题计数器
                parent_counter += 1
                category_counter = 0  # 重置三级标题计数器

                # 生成二级标题编号
                parent_number = f"{top_level_prefix}.{parent_counter}"
                create_parent_header(doc, f"{parent_number} {table_parent}")
                current_parent = table_parent

            # 添加分类标题（即使分类名称与父标题相同也要添加）
            category_counter += 1
            table_in_category_counter = 0  # 重置四级表格计数器

            # 生成三级标题编号
            category_number = f"{top_level_prefix}.{parent_counter}.{category_counter}"
            create_category_header(doc, f"{category_number} {table_category}")

            current_category = table_category

        # 更新四级表格计数器
        table_in_category_counter += 1

        # 生成四级表格编号
        table_number = (
            f"{top_level_prefix}.{parent_counter}.{category_counter}."
            f"{table_in_category_counter}"
        )

        # 同一分类内的下一张表，分页位于四级表题之前。
        if i > 0 and not category_changed:
            doc.add_page_break()

        # 处理文件名中的特殊字符
        safe_table_name = table_name.replace("/", "_")

        # 查找文件
        source_path = None
        source_type = None

        # 1. 查找填充版本
        for name in [table_name, safe_table_name]:
            filled_path = os.path.join(filled_dir, f"{name}.docx")
            if os.path.exists(filled_path):
                source_path = filled_path
                source_type = "filled"
                break

        # 2. 如果没找到填充版本，查找模板
        if not source_path:
            for name in [table_name, safe_table_name]:
                template_path = os.path.join(template_dir, f"{name}.docx")
                if os.path.exists(template_path):
                    source_path = template_path
                    source_type = "template"
                    break

        # 3. 如果都没找到，使用占位表格
        if not source_path:
            create_table_caption(doc, table_name, table_number)
            create_placeholder_table(doc, table_name)
            doc.add_paragraph()
            used_placeholder += 1
            continue

        # 统计
        if source_type == "filled":
            used_filled += 1
        else:
            used_template += 1

        # 添加表题
        create_table_caption(doc, table_name, table_number)

        # 复制表格
        try:
            source_doc = docx.Document(source_path)
            if source_doc.tables:
                copy_table_from_doc(source_doc, doc, 0)
            else:
                create_placeholder_table(doc, table_name)
        except Exception as e:
            log(f"❌ 表{table_index} 复制失败: {e}")
            create_placeholder_table(doc, table_name)

        # 添加空行分隔
        doc.add_paragraph()

    # 保存文档
    if output_path is None:
        output_path = os.path.join(output_dir, "合并的表格.docx")

    doc.save(output_path)

    log(f"\n{'='*60}")
    log(f"✅ 合并完成: {output_path}")
    log(f"  已填充: {used_filled}")
    log(f"  用模板: {used_template}")
    log(f"  占位符: {used_placeholder}")
    log(f"{'='*60}")

    return output_path


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="表格合并工具")
    parser.add_argument("--output-dir", required=True, help="阶段2输出目录")
    parser.add_argument("--output", "-o", help="输出文件路径")

    args = parser.parse_args()

    if not os.path.exists(args.output_dir):
        print(f"错误: 输出目录不存在 - {args.output_dir}", file=sys.stderr)
        sys.exit(1)

    merge_tables(args.output_dir, args.output)
