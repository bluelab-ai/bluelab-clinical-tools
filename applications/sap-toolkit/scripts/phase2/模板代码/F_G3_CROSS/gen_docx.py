#!/usr/bin/env python3
"""
F_G3_CROSS 生成脚本（交叉汇总表-治疗前×治疗后-多组）
生成三线表 Word 文档
"""

import json
import sys
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def load_json(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表核心）"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    old_borders = tcPr.find(qn("w:tcBorders"))
    if old_borders is not None:
        tcPr.remove(old_borders)

    borders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if spec is None or spec == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:color"), spec.get("color", "auto"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:space"), "0")
        borders.append(el)

    tcPr.append(borders)


def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本及格式（英文 Times New Roman，中文宋体，10.5pt）"""
    para = cell.paragraphs[0]
    para.alignment = alignment

    for run in para.runs:
        run._element.getparent().remove(run._element)

    run = para.add_run(str(text))

    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")

    for tag in ["w:sz", "w:szCs"]:
        sz = rPr.find(qn(tag))
        if sz is None:
            sz = OxmlElement(tag)
            rPr.append(sz)
        sz.set(qn("w:val"), "21")


def set_cell_vertical_align(cell, val="bottom"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def set_cell_background(cell, fill="FFFFFF"):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_row_header(row):
    """设置表头行属性（不跨页分割、重复表头）"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)

    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def apply_three_line_borders(table):
    """应用三线表边框"""
    total_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_borders(
                    cell,
                    top={"sz": 4, "color": "auto"},
                    bottom={"sz": 4, "color": "000000"},
                    left="nil",
                    right="nil",
                )
            elif i == total_rows - 1:
                set_cell_borders(
                    cell,
                    top="nil",
                    bottom={"sz": 4, "color": "auto"},
                    left="nil",
                    right="nil",
                )
            else:
                set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")

            set_cell_vertical_align(cell, "bottom")
            set_cell_background(cell, "FFFFFF")


def merge_cells(table, row_idx, start_col, end_col):
    """水平合并单元格"""
    row = table.rows[row_idx]
    start_cell = row.cells[start_col]
    end_cell = row.cells[end_col]
    start_cell.merge(end_cell)


def generate_docx(filled_data, output_path):
    """从填充好的 JSON 生成 Word 文档"""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    indicators = filled_data.get("indicators", [])
    if not indicators:
        doc.save(output_path)
        return

    for ind_idx, indicator in enumerate(indicators):
        rows = indicator.get("rows", [])
        display = indicator.get("display", "")

        header_rows = []
        data_rows = []
        for row in rows:
            if row.get("label_values") and not row.get("metric"):
                header_rows.append(row)
            else:
                data_rows.append(row)

        num_cols = 9
        total_rows = len(header_rows) + len(data_rows)
        table = doc.add_table(rows=total_rows, cols=num_cols)

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        merge_cells(table, 0, 3, 7)
        merge_cells(table, 2, 0, num_cols - 1)

        header1 = header_rows[0]
        for j, text in enumerate(header1["label_values"]):
            if j >= 3 and j <= 7:
                continue
            set_cell_text(table.rows[0].cells[j], text, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table.rows[0].cells[3], "治疗后", WD_ALIGN_PARAGRAPH.LEFT)

        header2 = header_rows[1]
        for j, text in enumerate(header2["label_values"]):
            set_cell_text(table.rows[1].cells[j], text, WD_ALIGN_PARAGRAPH.LEFT)

        set_cell_text(table.rows[2].cells[0], display, WD_ALIGN_PARAGRAPH.LEFT)

        for i, row_data in enumerate(data_rows):
            row_idx = len(header_rows) + i
            group_text = row_data.get("group", "")
            metric = row_data.get("metric", "")
            placeholder = row_data.get("placeholder", "")

            set_cell_text(table.rows[row_idx].cells[0], "")
            set_cell_text(table.rows[row_idx].cells[1], group_text)
            set_cell_text(table.rows[row_idx].cells[2], metric)

            applies_to = row_data.get("applies_to", [])
            for k, col_name in enumerate(applies_to):
                col_idx = 3 + k
                if col_idx < num_cols:
                    set_cell_text(table.rows[row_idx].cells[col_idx], placeholder)

        set_row_header(table.rows[0])
        set_row_header(table.rows[1])

        apply_three_line_borders(table)

        if ind_idx < len(indicators) - 1:
            doc.add_paragraph()

    doc.save(output_path)


def main():
    data_path = os.path.join(SCRIPT_DIR, "填充结果.json")
    output_path = os.path.join(SCRIPT_DIR, "输出结果.docx")

    data = load_json(data_path)
    generate_docx(data, output_path)
    indicator_count = len(data.get("indicators", []))
    print(f"✓ 生成完成: {output_path}")
    print(f"  包含 {indicator_count} 个指标（三线表格式）")


if __name__ == "__main__":
    main()
