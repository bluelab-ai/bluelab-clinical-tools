#!/usr/bin/env python3
"""
基于填充好的 JSON 文件，生成与模版格式一致的 Word 文档（三线表）

用法:
    python gen_docx.py <filled_json> [-o 输出文件]

示例:
    python gen_docx.py output/filled.json
    python gen_docx.py output/filled.json -o result.docx
"""

import json
import sys
import os
import argparse
from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def load_json(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表核心）"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    # 移除已有边框
    old_borders = tcPr.find(qn("w:tcBorders"))
    if old_borders is not None:
        tcPr.remove(old_borders)

    borders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if spec is None or spec == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:color"), spec.get("color", "auto"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:space"), "0")
        borders.append(el)

    tcPr.append(borders)


def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本及格式（英文 Times New Roman，中文宋体，10.5pt）"""
    para = cell.paragraphs[0]
    para.alignment = alignment

    # 清空已有内容
    for run in para.runs:
        run._element.getparent().remove(run._element)

    run = para.add_run(str(text))

    # 字体设置
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)

    # 字体
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")

    # 字号: 10.5pt = 21 half-points
    for tag in ["w:sz", "w:szCs"]:
        sz = rPr.find(qn(tag))
        if sz is None:
            sz = OxmlElement(tag)
            rPr.append(sz)
        sz.set(qn("w:val"), "21")


def set_cell_vertical_align(cell, val="bottom"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def set_cell_background(cell, fill="FFFFFF"):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_row_header(row):
    """设置表头行属性（不跨页分割、重复表头）"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)

    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def apply_three_line_borders(table):
    """
    应用三线表边框：
    - 表头行：顶线 + 底线
    - 数据行：无边框
    - 最后一行：底线
    """
    total_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                # 表头行: 顶线 + 底线
                set_cell_borders(
                    cell,
                    top={"sz": 4, "color": "auto"},
                    bottom={"sz": 4, "color": "000000"},
                    left="nil",
                    right="nil",
                )
            elif i == total_rows - 1:
                # 最后一行: 底线
                set_cell_borders(
                    cell,
                    top="nil",
                    bottom={"sz": 4, "color": "auto"},
                    left="nil",
                    right="nil",
                )
            else:
                # 数据行: 无边框
                set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")

            # 通用格式
            set_cell_vertical_align(cell, "bottom")
            set_cell_background(cell, "FFFFFF")


def merge_cells(table, row_idx, start_col, end_col):
    """水平合并单元格"""
    row = table.rows[row_idx]
    start_cell = row.cells[start_col]
    end_cell = row.cells[end_col]
    start_cell.merge(end_cell)


def merge_cells_vertical(table, start_row, end_row, col_idx):
    """垂直合并单元格"""
    start_cell = table.rows[start_row].cells[col_idx]
    end_cell = table.rows[end_row].cells[col_idx]
    start_cell.merge(end_cell)


def build_header_row(columns):
    """构建表头行: ['项目', '指标', '试验组\n(N=n1)', '对照组\n(N=n1)', '合计\n(N=n1)']"""
    headers = []
    for col in columns:
        name = col["name"]
        desc = col.get("description", "")
        if "(N=n1)" in desc:
            name = f"{name}\n(N=n1)"
        headers.append(name)
    return headers


def generate_docx(filled_data, output_path):
    """从填充好的 JSON 生成 Word 文档"""
    doc = Document()

    # 页面边距: 左右 1.91cm
    section = doc.sections[0]
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    columns = filled_data.get("columns", [])
    indicators = filled_data.get("indicators", [])
    num_cols = len(columns)

    # 构建所有数据行
    all_data_rows = []
    for indicator in indicators:
        rows_data = indicator.get("rows", [])
        for row in rows_data:
            indicator_text = row.get("indicator", "")
            group_text = row.get("group", "")
            metric = row.get("metric", "")
            placeholder = row.get("placeholder", "")
            applies_to = row.get("applies_to", [])
            label_values = row.get("label_values", [])

            # 如果有 label_values，直接使用（表头行）
            if label_values:
                row_cells = label_values
            else:
                # 数据行：[项目, 组别, 治疗前, 治疗后...]
                # 项目列：空
                # 组别列：group_text（只在第一行显示）
                # 治疗前列：metric
                row_cells = ["", group_text, metric]
                for col in columns[3:]:
                    col_name = col["name"]
                    if col_name in applies_to:
                        row_cells.append(placeholder)
                    else:
                        row_cells.append("")
            all_data_rows.append(row_cells)

    # 构建表头（使用前两行）
    header_row = all_data_rows[0] if all_data_rows else []
    header_row2 = all_data_rows[1] if len(all_data_rows) > 1 else []

    # 过滤掉空行
    filtered_data_rows = []
    for row in all_data_rows:
        # 检查是否有内容
        has_content = False
        if isinstance(row, list):
            has_content = any(cell for cell in row)
        elif isinstance(row, dict):
            has_content = bool(row.get('label_values') or row.get('metric') or row.get('group'))
        if has_content:
            filtered_data_rows.append(row)

    # 创建表格（不需要额外的表头行，因为表头已经在数据中）
    total_rows = len(filtered_data_rows)
    table = doc.add_table(rows=total_rows, cols=num_cols)

    # 移除默认表格样式（使用自定义三线表）
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    # 设置表格宽度 100%
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # 单元格边距: top=0, left=108, bottom=0, right=108 dxa
    tblCellMar = tblPr.find(qn("w:tblCellMar"))
    if tblCellMar is None:
        tblCellMar = OxmlElement("w:tblCellMar")
        tblPr.append(tblCellMar)
    for side, val in [("top", "0"), ("left", "108"), ("bottom", "0"), ("right", "108")]:
        el = tblCellMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tblCellMar.append(el)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")

    # 合并单元格
    # 表头行1（第1行）：治疗后列合并（列3-7）
    merge_cells(table, 0, 3, 7)

    # 表头行2（第2行）：不合并

    # 指标名行（第3行）水平合并所有列
    merge_cells(table, 2, 0, num_cols - 1)

    # 组别列（第2列）不垂直合并，只在每组第一行显示组名

    # 填写表头行1
    for j, text in enumerate(header_row):
        # 跳过合并的单元格（列4-7）
        if j > 3 and j < 8:
            continue
        set_cell_text(table.rows[0].cells[j], text)

    # 填写表头行2
    for j, text in enumerate(header_row2):
        set_cell_text(table.rows[1].cells[j], text)

    # 填写指标名行（第3行，合并后只写第一个单元格）
    indicator_row = filtered_data_rows[2] if len(filtered_data_rows) > 2 else []
    if indicator_row:
        indicator_text = indicator_row[0] if isinstance(indicator_row, list) else indicator_row.get('label_values', [''])[0]
        set_cell_text(table.rows[2].cells[0], indicator_text)

    # 表头行属性
    set_row_header(table.rows[0])
    set_row_header(table.rows[1])

    # 填写数据行（跳过前3行：表头行1、表头行2、指标名行）
    for i, row_data in enumerate(filtered_data_rows[3:]):
        row_idx = i + 3  # 从第4行开始（索引3）
        for j in range(num_cols):
            text = row_data[j] if j < len(row_data) else ""
            # 跳过合并的单元格
            if j == 1 and not text:
                continue
            set_cell_text(table.rows[row_idx].cells[j], text)

    # 应用三线表边框
    apply_three_line_borders(table)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="基于填充好的 JSON 生成三线表 Word 文档"
    )
    parser.add_argument("filled_json", help="填充好的 JSON 文件路径")
    parser.add_argument("-o", "--output", help="输出 Word 文件路径", default=None)

    args = parser.parse_args()

    filled_data = load_json(args.filled_json)

    if args.output:
        output_path = args.output
    else:
        base = os.path.splitext(os.path.basename(args.filled_json))[0]
        output_path = f"{base}.docx"

    generate_docx(filled_data, output_path)

    indicator_count = len(filled_data.get("indicators", []))
    print(f"✓ 已生成: {output_path}")
    print(f"  包含 {indicator_count} 个指标（三线表格式）")


if __name__ == "__main__":
    main()
