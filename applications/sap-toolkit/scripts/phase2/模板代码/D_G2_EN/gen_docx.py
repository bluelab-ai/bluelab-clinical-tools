#!/usr/bin/env python3
"""
D_G2_EN 生成脚本（受试者入组与完成情况表）
生成三线表 Word 文档
"""

import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def load_json(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表核心）"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    old_borders = tcPr.find(qn("w:tcBorders"))
    if old_borders is not None:
        tcPr.remove(old_borders)

    borders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if spec is None or spec == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:color"), spec.get("color", "auto"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:space"), "0")
        borders.append(el)

    tcPr.append(borders)


def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本及格式（英文 Times New Roman，中文宋体，10.5pt）"""
    para = cell.paragraphs[0]
    para.alignment = alignment

    for run in para.runs:
        run._element.getparent().remove(run._element)

    run = para.add_run(str(text))

    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")

    for tag in ["w:sz", "w:szCs"]:
        sz = rPr.find(qn(tag))
        if sz is None:
            sz = OxmlElement(tag)
            rPr.append(sz)
        sz.set(qn("w:val"), "21")


def set_cell_vertical_align(cell, val="bottom"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def set_cell_background(cell, fill="FFFFFF"):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_row_header(row):
    """设置表头行属性"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)

    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def apply_three_line_borders(table):
    """应用三线表边框"""
    total_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_borders(
                    cell,
                    top={"sz": 4, "color": "auto"},
                    bottom={"sz": 4, "color": "000000"},
                    left="nil",
                    right="nil",
                )
            elif i == total_rows - 1:
                set_cell_borders(
                    cell,
                    top="nil",
                    bottom={"sz": 4, "color": "auto"},
                    left="nil",
                    right="nil",
                )
            else:
                set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")

            set_cell_vertical_align(cell, "bottom")
            set_cell_background(cell, "FFFFFF")


def merge_cells(table, row_idx, start_col, end_col):
    """水平合并单元格"""
    row = table.rows[row_idx]
    start_cell = row.cells[start_col]
    end_cell = row.cells[end_col]
    start_cell.merge(end_cell)


def generate_docx(filled_data, output_path):
    """从填充好的 JSON 生成 Word 文档"""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    rows = filled_data.get("rows", [])
    columns = filled_data.get("columns", [])
    num_cols = len(columns)  # 5列: 分类, 项目, 试验组, 对照组, 合计

    # 表头行
    header_row = [col["name"] for col in columns]

    # 所有行（包括分类标题行和数据行）
    total_rows = 1 + len(rows)  # 1行表头 + 所有数据行
    table = doc.add_table(rows=total_rows, cols=num_cols)

    # 设置表格宽度 100%
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # 填写表头行
    for j, text in enumerate(header_row):
        set_cell_text(table.rows[0].cells[j], text, WD_ALIGN_PARAGRAPH.CENTER)

    # 填写数据行
    for i, row_data in enumerate(rows):
        row_idx = i + 1
        label_vals = row_data.get("label_values", [])
        data_vals = row_data.get("data_values", [])
        is_header = row_data.get("is_category_header", False)

        if is_header:
            # 分类标题行：第一列填分类名，其余列空
            set_cell_text(table.rows[row_idx].cells[0], label_vals[0], WD_ALIGN_PARAGRAPH.LEFT)
            for j in range(1, num_cols):
                set_cell_text(table.rows[row_idx].cells[j], "")
        else:
            # 数据行：第一列空，第二列填项目名，后面填数据
            set_cell_text(table.rows[row_idx].cells[0], "")
            metric = label_vals[1] if len(label_vals) > 1 else ""
            set_cell_text(table.rows[row_idx].cells[1], metric)

            for k, val in enumerate(data_vals):
                col_idx = 2 + k
                if col_idx < num_cols:
                    set_cell_text(table.rows[row_idx].cells[col_idx], val, WD_ALIGN_PARAGRAPH.CENTER)

    # 表头行属性
    set_row_header(table.rows[0])

    # 应用三线表边框
    apply_three_line_borders(table)

    doc.save(output_path)


def main():
    data_path = os.path.join(SCRIPT_DIR, "填充结果.json")
    output_path = os.path.join(SCRIPT_DIR, "输出结果.docx")

    data = load_json(data_path)
    generate_docx(data, output_path)
    print(f"✓ 生成完成: {output_path}")


if __name__ == "__main__":
    main()
