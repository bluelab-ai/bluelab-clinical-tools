#!/usr/bin/env python3
"""
F_G1_QUALCOM_N_N 生成脚本（定性指标比较表-单组-无缺失）
生成三线表 Word 文档
"""

import json
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def load_json(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表核心）"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    old_borders = tcPr.find(qn("w:tcBorders"))
    if old_borders is not None:
        tcPr.remove(old_borders)

    borders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if spec is None or spec == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:color"), spec.get("color", "auto"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:space"), "0")
        borders.append(el)

    tcPr.append(borders)


def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本及格式（英文 Times New Roman，中文宋体，10.5pt）"""
    para = cell.paragraphs[0]
    para.alignment = alignment

    for run in para.runs:
        run._element.getparent().remove(run._element)

    run = para.add_run(str(text))

    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")

    for tag in ["w:sz", "w:szCs"]:
        sz = rPr.find(qn(tag))
        if sz is None:
            sz = OxmlElement(tag)
            rPr.append(sz)
        sz.set(qn("w:val"), "21")


def set_cell_vertical_align(cell, val="bottom"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def set_cell_background(cell, fill="FFFFFF"):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_row_header(row):
    """设置表头行属性"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)

    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def apply_three_line_borders(table):
    """应用三线表边框"""
    total_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_borders(
                    cell,
                    top={"sz": 4, "color": "auto"},
                    bottom={"sz": 4, "color": "000000"},
                    left="nil",
                    right="nil",
                )
            elif i == total_rows - 1:
                set_cell_borders(
                    cell,
                    top="nil",
                    bottom={"sz": 4, "color": "auto"},
                    left="nil",
                    right="nil",
                )
            else:
                set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")

            set_cell_vertical_align(cell, "bottom")
            set_cell_background(cell, "FFFFFF")


def generate_docx(filled_data, output_path):
    """从填充好的 JSON 生成 Word 文档"""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    columns = filled_data.get("columns", [])
    indicators = filled_data.get("indicators", [])
    num_cols = len(columns)

    # 构建所有数据行
    all_data_rows = []
    for indicator in indicators:
        rows_data = indicator.get("rows", [])
        for row in rows_data:
            indicator_text = row.get("indicator", "")
            metric = row.get("metric", "")
            placeholder = row.get("placeholder", "")
            applies_to = row.get("applies_to", [])

            row_cells = [indicator_text, metric]
            for col in columns[2:]:
                col_name = col["name"]
                if col_name in applies_to:
                    row_cells.append(placeholder)
                else:
                    row_cells.append("")
            all_data_rows.append(row_cells)

    # 创建表格
    total_rows = 1 + len(all_data_rows)
    table = doc.add_table(rows=total_rows, cols=num_cols)

    # 设置表格宽度 100%
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # 填写表头
    for j, col in enumerate(columns):
        set_cell_text(table.rows[0].cells[j], col["name"], WD_ALIGN_PARAGRAPH.CENTER)

    # 表头行属性
    set_row_header(table.rows[0])

    # 填写数据行
    for i, row_data in enumerate(all_data_rows):
        row_idx = i + 1
        for j in range(num_cols):
            text = row_data[j] if j < len(row_data) else ""
            set_cell_text(table.rows[row_idx].cells[j], text)

    # 应用三线表边框
    apply_three_line_borders(table)

    doc.save(output_path)


def main():
    data_path = os.path.join(SCRIPT_DIR, "填充结果.json")
    output_path = os.path.join(SCRIPT_DIR, "输出结果.docx")

    data = load_json(data_path)
    generate_docx(data, output_path)
    indicator_count = len(data.get("indicators", []))
    print(f"✓ 生成完成: {output_path}")
    print(f"  包含 {indicator_count} 个指标（三线表格式）")


if __name__ == "__main__":
    main()
