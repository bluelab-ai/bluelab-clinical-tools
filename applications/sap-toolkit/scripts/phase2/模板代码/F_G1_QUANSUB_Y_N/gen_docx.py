#!/usr/bin/env python3
"""
基于填充好的 JSON 文件，生成与模版格式一致的 Word 文档（三线表）
支持亚组合并单元格

用法:
    python gen_docx.py <filled_json> [-o 输出文件]

示例:
    python gen_docx.py output/filled.json
    python gen_docx.py output/filled.json -o result.docx
"""

import json
import sys
import os
import argparse
from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def load_json(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表核心）"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    # 移除已有边框
    old_borders = tcPr.find(qn("w:tcBorders"))
    if old_borders is not None:
        tcPr.remove(old_borders)

    borders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if spec is None or spec == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:color"), spec.get("color", "auto"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:space"), "0")
        borders.append(el)

    tcPr.append(borders)


def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本及格式（英文 Times New Roman，中文宋体，10.5pt）"""
    para = cell.paragraphs[0]
    para.alignment = alignment

    # 清空已有内容
    for run in para.runs:
        run._element.getparent().remove(run._element)

    run = para.add_run(str(text))

    # 字体设置
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)

    # 字体
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")

    # 字号: 10.5pt = 21 half-points
    for tag in ["w:sz", "w:szCs"]:
        sz = rPr.find(qn(tag))
        if sz is None:
            sz = OxmlElement(tag)
            rPr.append(sz)
        sz.set(qn("w:val"), "21")


def set_cell_vertical_align(cell, val="bottom"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def set_cell_background(cell, fill="FFFFFF"):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_row_header(row):
    """设置表头行属性（不跨页分割、重复表头）"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)

    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def apply_three_line_borders(table):
    """
    应用三线表边框：
    - 表头行：顶线 + 底线
    - 数据行：无边框
    - 最后一行：底线
    """
    total_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                # 表头行: 顶线 + 底线
                set_cell_borders(
                    cell,
                    top={"sz": 4, "color": "auto"},
                    bottom={"sz": 4, "color": "000000"},
                    left="nil",
                    right="nil",
                )
            elif i == total_rows - 1:
                # 最后一行: 底线
                set_cell_borders(
                    cell,
                    top="nil",
                    bottom={"sz": 4, "color": "auto"},
                    left="nil",
                    right="nil",
                )
            else:
                # 数据行: 无边框
                set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")

            # 通用格式
            set_cell_vertical_align(cell, "bottom")
            set_cell_background(cell, "FFFFFF")


def build_header_row(columns):
    """构建表头行: ['亚组', '项目', '指标', '结果\n(N=n1)']"""
    headers = []
    for col in columns:
        name = col["name"]
        desc = col.get("description", "")
        if "(N=n1)" in desc:
            name = f"{name}\n(N=n1)"
        headers.append(name)
    return headers


def generate_docx(filled_data, output_path):
    """从填充好的 JSON 生成 Word 文档"""
    doc = Document()

    # 页面边距: 左右 1.91cm
    section = doc.sections[0]
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    columns = filled_data.get("columns", [])
    indicators = filled_data.get("indicators", [])
    num_cols = len(columns)

    # 构建表头
    header_row = build_header_row(columns)

    # 构建所有数据行（按亚组分组，每个亚组内包含所有指标）
    all_data_rows = []

    # 获取亚组列表（从第一个指标中获取）
    if indicators:
        subgroups = indicators[0].get("subgroups", [])
        subgroup_names = [sg.get("name", "") for sg in subgroups]
    else:
        subgroup_names = []

    # 按亚组遍历
    for sg_idx, sg_name in enumerate(subgroup_names):
        # 每个亚组内遍历所有指标
        for indicator_idx, indicator in enumerate(indicators):
            display = indicator.get("display", "")
            subgroups_data = indicator.get("subgroups", [])

            if sg_idx < len(subgroups_data):
                rows_data = subgroups_data[sg_idx].get("rows", [])

                for row_idx, row in enumerate(rows_data):
                    metric = row.get("metric", "")
                    placeholder = row.get("placeholder", "")
                    applies_to = row.get("applies_to", [])

                    # 第一行显示指标名称，后续行为空
                    indicator_text = display if row_idx == 0 else ""

                    # 亚组名称只在每个亚组的第一行显示
                    sg_text = sg_name if indicator_idx == 0 and row_idx == 0 else ""

                    row_cells = [sg_text, indicator_text, metric]
                    for col in columns[3:]:
                        col_name = col["name"]
                        if col_name in applies_to:
                            row_cells.append(placeholder)
                        else:
                            row_cells.append("")
                    all_data_rows.append(row_cells)

    # 创建表格
    total_rows = 1 + len(all_data_rows)
    table = doc.add_table(rows=total_rows, cols=num_cols)

    # 移除默认表格样式（使用自定义三线表）
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    # 设置表格宽度 100%
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # 单元格边距: top=0, left=108, bottom=0, right=108 dxa
    tblCellMar = tblPr.find(qn("w:tblCellMar"))
    if tblCellMar is None:
        tblCellMar = OxmlElement("w:tblCellMar")
        tblPr.append(tblCellMar)
    for side, val in [("top", "0"), ("left", "108"), ("bottom", "0"), ("right", "108")]:
        el = tblCellMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tblCellMar.append(el)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")

    # 填写表头
    for j, text in enumerate(header_row):
        set_cell_text(table.rows[0].cells[j], text)

    # 表头行属性
    set_row_header(table.rows[0])

    # 填写数据行
    for i, row_data in enumerate(all_data_rows):
        row_idx = i + 1
        for j in range(num_cols):
            text = row_data[j] if j < len(row_data) else ""
            set_cell_text(table.rows[row_idx].cells[j], text)

    # 应用三线表边框
    apply_three_line_borders(table)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="基于填充好的 JSON 生成三线表 Word 文档（支持亚组）"
    )
    parser.add_argument("filled_json", help="填充好的 JSON 文件路径")
    parser.add_argument("-o", "--output", help="输出 Word 文件路径", default=None)

    args = parser.parse_args()

    filled_data = load_json(args.filled_json)

    if args.output:
        output_path = args.output
    else:
        base = os.path.splitext(os.path.basename(args.filled_json))[0]
        output_path = f"{base}.docx"

    generate_docx(filled_data, output_path)

    indicator_count = len(filled_data.get("indicators", []))
    print(f"✓ 已生成: {output_path}")
    print(f"  包含 {indicator_count} 个指标（三线表格式）")


if __name__ == "__main__":
    main()