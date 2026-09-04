#!/usr/bin/env python3
"""
F_G3_ANCOVA_02 文档生成脚本
基于填充好的 JSON 文件，生成协方差分析表（不含交互项）Word 文档（三线表）

用法:
    python gen_docx.py <filled_json> [-o 输出文件]
"""

import json
import sys
import os
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def load_json(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框（三线表核心）"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    old_borders = tcPr.find(qn("w:tcBorders"))
    if old_borders is not None:
        tcPr.remove(old_borders)

    borders = OxmlElement("w:tcBorders")
    for side, spec in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if spec is None or spec == "nil":
            el.set(qn("w:val"), "nil")
        else:
            el.set(qn("w:val"), "single")
            el.set(qn("w:color"), spec.get("color", "auto"))
            el.set(qn("w:sz"), str(spec.get("sz", 4)))
            el.set(qn("w:space"), "0")
        borders.append(el)

    tcPr.append(borders)


def set_cell_text(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本及格式"""
    para = cell.paragraphs[0]
    para.alignment = alignment

    for run in para.runs:
        run._element.getparent().remove(run._element)

    run = para.add_run(str(text))

    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")

    for tag in ["w:sz", "w:szCs"]:
        sz = rPr.find(qn(tag))
        if sz is None:
            sz = OxmlElement(tag)
            rPr.append(sz)
        sz.set(qn("w:val"), "21")


def set_cell_vertical_align(cell, val="bottom"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def set_cell_background(cell, fill="FFFFFF"):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)

    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_row_header(row):
    """设置表头行属性"""
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)

    cantSplit = OxmlElement("w:cantSplit")
    trPr.append(cantSplit)

    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def apply_three_line_borders(table):
    """应用三线表边框"""
    total_rows = len(table.rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_borders(
                    cell,
                    top={"sz": 4, "color": "auto"},
                    bottom={"sz": 4, "color": "000000"},
                    left="nil",
                    right="nil",
                )
            elif i == total_rows - 1:
                set_cell_borders(
                    cell,
                    top="nil",
                    bottom={"sz": 4, "color": "auto"},
                    left="nil",
                    right="nil",
                )
            else:
                set_cell_borders(cell, top="nil", bottom="nil", left="nil", right="nil")

            set_cell_vertical_align(cell, "bottom")
            set_cell_background(cell, "FFFFFF")


def generate_docx(filled_data, output_path):
    """从填充好的 JSON 生成 Word 文档"""
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    columns = filled_data.get("columns", [])
    indicators = filled_data.get("indicators", [])
    num_cols = len(columns)

    # 构建表头
    header_row = [col["name"] for col in columns]

    # 构建所有数据行
    all_data_rows = []
    for indicator in indicators:
        rows_data = indicator.get("rows", [])
        for row in rows_data:
            label_values = row.get("label_values", [])
            data_values = row.get("data_values", [])
            applies_to = row.get("applies_to", [])

            # 第一列：项目名（仅第一行显示）
            indicator_text = label_values[0] if label_values else ""
            # 第二列：指标
            metric = label_values[1] if len(label_values) > 1 else ""

            row_cells = [indicator_text, metric]
            for col in columns[2:]:
                col_name = col["name"]
                if col_name in applies_to:
                    # 找到对应的 data_value
                    idx = applies_to.index(col_name)
                    if idx < len(data_values):
                        row_cells.append(data_values[idx])
                    else:
                        row_cells.append("")
                else:
                    row_cells.append("")
            all_data_rows.append(row_cells)

    # 创建表格
    total_rows = 1 + len(all_data_rows)
    table = doc.add_table(rows=total_rows, cols=num_cols)

    # 设置表格宽度
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # 单元格边距
    tblCellMar = tblPr.find(qn("w:tblCellMar"))
    if tblCellMar is None:
        tblCellMar = OxmlElement("w:tblCellMar")
        tblPr.append(tblCellMar)
    for side, val in [("top", "0"), ("left", "108"), ("bottom", "0"), ("right", "108")]:
        el = tblCellMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tblCellMar.append(el)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")

    # 填写表头
    for j, text in enumerate(header_row):
        set_cell_text(table.rows[0].cells[j], text)

    set_row_header(table.rows[0])

    # 填写数据行
    for i, row_data in enumerate(all_data_rows):
        row_idx = i + 1
        for j in range(num_cols):
            text = row_data[j] if j < len(row_data) else ""
            set_cell_text(table.rows[row_idx].cells[j], text)

    apply_three_line_borders(table)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="基于填充好的 JSON 生成协方差分析表（不含交互项）Word 文档"
    )
    parser.add_argument("filled_json", help="填充好的 JSON 文件路径")
    parser.add_argument("-o", "--output", help="输出 Word 文件路径", default=None)

    args = parser.parse_args()

    filled_data = load_json(args.filled_json)

    if args.output:
        output_path = args.output
    else:
        base = os.path.splitext(os.path.basename(args.filled_json))[0]
        output_path = f"{base}.docx"

    generate_docx(filled_data, output_path)

    indicator_count = len(filled_data.get("indicators", []))
    print(f"✓ 已生成: {output_path}")
    print(f"  包含 {indicator_count} 个指标（三线表格式）")


if __name__ == "__main__":
    main()
