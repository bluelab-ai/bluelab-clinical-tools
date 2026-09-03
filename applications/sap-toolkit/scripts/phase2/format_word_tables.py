#!/usr/bin/env python3
"""Batch-format tables in .docx files.

Usage
-----
    pip3 install python-docx
    python3 format_word_tables.py /path/to/document.docx
    python3 format_word_tables.py /path/to/folder

The source document is not changed.  A copy ending in ``_formatted.docx`` is
created alongside each processed file.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_run_fonts(run, reset_emphasis: bool = False) -> None:
    """Set Western and East-Asian fonts on a run."""
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    if reset_emphasis:
        run.font.bold = False
        run.font.italic = False
        run.font.color.rgb = RGBColor(0, 0, 0)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)

    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_mark_font(paragraph) -> None:
    """Set the default font used by a paragraph mark, including empty cells."""
    p_pr = paragraph._p.get_or_add_pPr()
    r_pr = p_pr.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        p_pr.append(r_pr)

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")

    for tag in ("w:sz", "w:szCs"):
        size = r_pr.find(qn(tag))
        if size is None:
            size = OxmlElement(tag)
            r_pr.append(size)
        size.set(qn("w:val"), "21")


def format_paragraph(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.0       # single spacing
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    set_paragraph_mark_font(paragraph)

    for run in paragraph.runs:
        set_run_fonts(run)

    if not paragraph.runs:
        set_run_fonts(paragraph.add_run())


def set_table_auto_fit(table) -> None:
    """Enable Word's automatic table layout.

    python-docx cannot reproduce Word VBA's wdAutoFitWindow exactly.  This sets
    the OOXML table layout to automatic, so Word may adapt columns when opened.
    """
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")


def format_table(table) -> None:
    set_table_auto_fit(table)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph)
            for nested_table in cell.tables:
                format_table(nested_table)


def format_toc_heading(paragraph) -> None:
    """Preserve the generated TOC heading's requested presentation."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_run_fonts(run, reset_emphasis=True)
        run.font.size = Pt(16)


def format_non_table_text(document) -> None:
    """Format text outside tables in the body, headers, and footers."""
    for paragraph in document.paragraphs:
        if paragraph.style.name == "TOC Heading":
            format_toc_heading(paragraph)
            continue
        is_heading = paragraph.style.name in {
            "Heading 1", "Heading 2", "Heading 3", "Heading 4"
        }
        for run in paragraph.runs:
            set_run_fonts(run, reset_emphasis=True)
            if is_heading:
                run.font.bold = True

    for section in document.sections:
        for story in (section.header, section.footer):
            for paragraph in story.paragraphs:
                for run in paragraph.runs:
                    set_run_fonts(run, reset_emphasis=True)


def set_landscape_orientation(document) -> None:
    """Change every document section from portrait to landscape."""
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )


def output_path(input_path: Path) -> Path:
    return input_path.with_name(f"{input_path.stem}_formatted.docx")


def process_document(input_path: Path) -> Path:
    document = Document(input_path)
    for table in document.tables:
        format_table(table)
    format_non_table_text(document)
    set_landscape_orientation(document)

    destination = output_path(input_path)
    document.save(destination)
    return destination


def find_documents(path: Path) -> list[Path]:
    if path.is_file():
        if path.suffix.lower() != ".docx" or path.name.startswith("~$"):
            raise ValueError("Please provide a non-temporary .docx file.")
        return [path]

    if path.is_dir():
        return [
            file
            for file in path.glob("*.docx")
            if not file.name.startswith("~$")
            and not file.stem.endswith("_formatted")
        ]

    raise FileNotFoundError(f"Path not found: {path}")


def main() -> None:
    if len(sys.argv) != 2:
        print("Usage: python3 format_word_tables.py <.docx file or folder>")
        raise SystemExit(2)

    target = Path(sys.argv[1]).expanduser().resolve()
    documents = find_documents(target)
    if not documents:
        print("No .docx files found.")
        return

    for document in documents:
        saved_to = process_document(document)
        print(f"Completed: {saved_to}")


if __name__ == "__main__":
    main()
