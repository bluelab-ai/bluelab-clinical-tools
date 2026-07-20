#!/usr/bin/env python3
"""
临床试验方案统计分析要素提取 — 一体化脚本。
从 docx 方案文件直达结构化 JSON（docx → Markdown → 质量判断 → 结构化 JSON）。

整合了以下脚本的功能：
  统计内容提取.py  — 两轮提取策略 + LLM 质量判断
  提取统计分析要素.py — LLM 结构化要素提取

用法:
  # 处理单个方案（最常用）
  python 统计分析提取_一体化.py --file "方案/xxx.docx"

  # 批量处理方案目录下所有 docx
  python 统计分析提取_一体化.py --dir "方案"

  # 指定输出目录
  python 统计分析提取_一体化.py --file "方案/xxx.docx" --out-dir "我的输出"

  # Dry-run 只列出方案文件
  python 统计分析提取_一体化.py --dry-run
"""

import argparse
import glob
import json
import os
import re
import shutil
import sys
import time
from typing import Any

import anthropic
from docx import Document as DocxDocument
from docx.table import Table as DocxTable

# ============================================================================
# 路径配置
# ============================================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)

# 默认目录
DEFAULT_SCHEME_DIR = os.path.join(PROJECT_DIR, "方案")
DEFAULT_OUT_DIR = os.path.join(PROJECT_DIR, "提取输出")

# ============================================================================
# API 配置
# ============================================================================
_sys_path_add = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..')
if _sys_path_add not in sys.path:
    sys.path.insert(0, _sys_path_add)
from config import LLM_API_KEY, LLM_API_BASE, LLM_MODEL as API_MODEL

API_BASE_URL = LLM_API_BASE
API_AUTH_TOKEN = LLM_API_KEY

# ============================================================================
# 导入 提取方案内容.py 的核心函数（通过 importlib 避免子进程调用）
# ============================================================================
_extract_content_path = os.path.join(SCRIPT_DIR, "提取方案内容.py")
if os.path.exists(_extract_content_path):
    import importlib.util as _iu

    _spec = _iu.spec_from_file_location("_extract_content", _extract_content_path)
    _ec = _iu.module_from_spec(_spec)
    _spec.loader.exec_module(_ec)
    # 核心函数引用
    parse_toc = _ec.parse_toc if hasattr(_ec, 'parse_toc') else None
    extract_section_func = _ec.extract_section
    extract_section_toc = _ec.extract_section_toc
    extract_to_markdown = _ec.extract_to_markdown
else:
    raise FileNotFoundError(f"核心模块不存在: {_extract_content_path}")

# ============================================================================
# 结构化提取 Schema（与 schemas/统计分析提取_schema.json 一致）
# ============================================================================
STRUCTURED_SCHEMA = {
    "type": "object",
    "required": ["统计分析计划"],
    "properties": {
        "统计分析计划": {
            "type": "object",
            "required": ["分析人群", "统计通用原则", "样本量与检验参数", "合格标准",
                         "主要终点分析", "次要终点分析", "安全性分析"],
            "properties": {
                "分析人群": {
                    "type": "object",
                    "description": "方案定义的各分析集。3个字段固定：FAS、PPS、SS。每个字段值为从原文抄录的定义或用途说明。如原文未明确定义，末尾标注（原文未明确定义）",
                    "required": ["FAS", "PPS", "SS"],
                    "properties": {
                        "FAS": {"type": "string", "description": "全分析集的定义或用途"},
                        "PPS": {"type": "string", "description": "符合方案集的定义或用途"},
                        "SS": {"type": "string", "description": "安全性分析集的定义或用途"},
                    },
                },
                "统计通用原则": {
                    "type": "array",
                    "description": "方案中声明的统计软件、检验方向、默认显著水平、检验假设框架等",
                    "items": {"type": "string"},
                },
                "期中分析": {
                    "type": "object",
                    "description": "【仅当原文确实计划进行期中分析时填写】",
                    "properties": {
                        "分析人群": {"type": "string"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}},
                    },
                },
                "样本量与检验参数": {
                    "type": "array",
                    "description": "样本量计算依据的参数和结果。每项独立列出。格式：「参数名 — 原文数值」",
                    "items": {"type": "string"},
                },
                "合格标准": {
                    "type": "string",
                    "description": "试验结果判定为合格/不合格的统计学标准。直接抄录原文完整语句。",
                },
                "受试者分布分析": {
                    "type": "object",
                    "description": "受试者筛选、入组、完成、脱落、各分析集分布及剔除情况的统计分析",
                    "required": ["分析人群", "条目"],
                    "properties": {
                        "分析人群": {"type": "string", "description": "如 FAS"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}, "description": "逐字抄录的原文分析语句。⚠️ 原子化拆分：每个数组元素只描述一个独立的分析步骤/检验。原文一句包含多个步骤（如既说协方差又说亚组又说缺失值处理）必须拆为多条。"},
                    },
                },
                "人口学与基线特征分析": {
                    "type": "object",
                    "description": "人口学特征、病史、生命体征等基线资料的统计描述与组间比较",
                    "required": ["分析人群", "条目"],
                    "properties": {
                        "分析人群": {"type": "string", "description": "如 FAS"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}, "description": "逐字抄录的原文分析语句。⚠️ 原子化拆分：每个数组元素只描述一个独立的分析步骤/检验。原文一句包含多个步骤（如既说协方差又说亚组又说缺失值处理）必须拆为多条。"},
                    },
                },
                "主要终点分析": {
                    "type": "object",
                    "description": "针对主要疗效终点的统计分析。如原文同一分析同时在FAS和PPS进行且方法相同，分析人群写「FAS、PPS」合并为一份条目；如方法不同则分别输出为两个字段",
                    "required": ["分析人群", "条目"],
                    "properties": {
                        "分析人群": {"type": "string", "description": "如「FAS、PPS」"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}, "description": "逐字抄录的原文分析语句。每条是一个独立的分析步骤"},
                    },
                },
                "次要终点分析": {
                    "type": "object",
                    "description": "针对次要疗效终点、器械性能终点的统计分析",
                    "required": ["分析人群", "条目"],
                    "properties": {
                        "分析人群": {"type": "string", "description": "如「FAS、PPS」"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}, "description": "逐字抄录的原文分析语句。⚠️ 原子化拆分：每个数组元素只描述一个独立的分析步骤/检验。原文一句包含多个步骤（如既说协方差又说亚组又说缺失值处理）必须拆为多条。"},
                    },
                },
                "安全性分析": {
                    "type": "object",
                    "description": "安全性指标分析，包括AE/SAE、实验室检查、生命体征等",
                    "required": ["分析人群", "条目"],
                    "properties": {
                        "分析人群": {"type": "string", "description": "如 SS"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}, "description": "逐字抄录的原文分析语句。⚠️ 原子化拆分：每个数组元素只描述一个独立的分析步骤/检验。原文一句包含多个步骤（如既说协方差又说亚组又说缺失值处理）必须拆为多条。"},
                    },
                },
                "亚组分析": {
                    "type": "object",
                    "description": "【仅当原文确实计划进行亚组分析时填写】",
                    "properties": {
                        "分析人群": {"type": "string"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}},
                    },
                },
                "敏感性分析": {
                    "type": "object",
                    "description": "【仅当原文确实计划进行敏感性分析时填写】",
                    "properties": {
                        "分析人群": {"type": "string"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}},
                    },
                },
                "缺失数据处理": {
                    "type": "array",
                    "description": "【仅当原文描述了缺失数据处理方法时填写】缺失数据、不合理数据、退出/撤出数据的处理方法。这是一个纯技术性条目，不绑定分析人群。每条是一个对象，说明该处理方法适用于哪些分析和具体方法。每项格式：{\"适用范围\": \"从原文抄录该处理方法针对的分析/场景\", \"处理方法\": \"从原文抄录具体处理方式\"}。如原文写「主要疗效指标缺失数据采取LOCF」，则适用范围写「主要疗效指标」，处理方法写「末次观测值结转(LOCF)」",
                    "items": {
                        "type": "object",
                        "required": ["适用范围", "处理方法"],
                        "properties": {
                            "适用范围": {"type": "string", "description": "该处理方法适用的分析或场景，从原文抄录"},
                            "处理方法": {"type": "string", "description": "具体的处理方式，从原文抄录"},
                        },
                    },
                },
                "其他分析": {
                    "type": "object",
                    "description": "【仅当存在上述分类未涵盖的统计分析时填写】",
                    "properties": {
                        "分析人群": {"type": "string"},
                        "表格": {"type": "string", "description": "对应的输出表格，暂不填写"},
                        "条目": {"type": "array", "items": {"type": "string"}},
                    },
                },
            },
        }
    },
}

# ============================================================================
# System Prompts
# ============================================================================

# 质量判断 prompt（与统计内容提取.py 一致）
JUDGE_PROMPT = """你是临床试验方案统计分析章节的质量审核员。判断提取出的 Markdown 文本是否同时包含以下三项：

① 样本量计算/样本量估计 —— 有具体的样本量计算依据（如主要终点指标、检验水准α、把握度1-β、预期效应量、脱落率等参数）和计算结果，不能仅一句话带过
② 分析人群/分析集定义 —— 有明确的分析集定义（如FAS/全分析集、PPS/符合方案集、SS/安全性分析集等），说明了各分析集的纳入标准
③ 统计分析方法 —— 有具体的统计检验方法（如t检验、卡方检验、方差分析、协方差分析、log-rank检验、Cox回归等），说明了对主要/次要终点分别采用什么统计方法

规则：
- 三项全有 → pass=true, missing="无"
- 缺任何一项 → pass=false, missing 写明缺哪项
- 文件以 "[错误]" 开头 → pass=false, missing="上游提取失败"
- 如果样本量计算仅写了"采用XX软件计算"而无具体参数，判定为缺失
- 如果分析集仅列出名称而无定义，判定为缺失

请只根据文件实际内容判断，不要编造。"""

JUDGE_SCHEMA = {
    "type": "object",
    "required": ["pass", "missing"],
    "properties": {
        "pass": {
            "type": "boolean",
            "description": "是否同时包含①样本量计算②分析人群定义③统计分析方法。全部包含为true",
        },
        "missing": {
            "type": "string",
            "description": "三项齐全写'无'，否则说明缺少哪项",
        },
    },
}

# 结构化提取 prompt
STRUCTURED_SYSTEM_PROMPT = """你是一个临床试验方案数据录入员。从给定的「统计分析」章节 Markdown 中逐字抄录信息。

## 最高原则：逐字抄录，禁止改写

你是复制粘贴工具，禁止改写、概括、换词、转换格式。
如果原文内容以「[错误]」开头，返回空对象 {}。

## 输出结构

所有内容包裹在 "统计分析计划" 之下。

### ⚠️ 分析人群 — 绝不可遗漏，必须作为第一个字段返回

这是强制性字段。你必须在 JSON 第一个字段输出 "分析人群" 对象，包含 FAS、PPS、SS 三个字段。
如果原文没有明确定义，抄录原文中说明该分析集用途的语句代替，并标注「（原文未明确定义）」。
示例：{"分析人群": {"FAS": "原文抄录...", "PPS": "原文抄录...", "SS": "原文抄录..."}}

找出原文中关于各分析集的说明：
- FAS/全分析集：通常在"完成情况及人口学分析"或"一般原则"附近说明用途
- PPS/符合方案集：与 FAS 出现在同一段落
- SS/安全性分析集：通常在"安全性评价"附近说明用途

### 统计通用原则
数组。逐条抄录原文中关于统计软件、检验方向、默认显著性水平等声明语句。

### 样本量与检验参数
数组。逐条抄录原文中样本量计算相关的每个参数和数值。格式：「参数名 — 原文数值」。
如原文在多处提到同一参数（如脱落率在正文和独立章节各提一次），只保留一条，取表述更完整的版本。

### 合格标准
字符串。直接抄录原文中判定试验合格/不合格的完整语句。一个字都不要改。

### 主要终点分析
对象 {"分析人群": "X", "表格": "", "条目": [...]}。这是最重要的类别，必须完整抄录。
- **第一段必须从原文的假设检验陈述开始**：如原文有 H0/H1 声明、检验类型（优效性/非劣效/等效）、α 水平、β 值，必须在条目中完整包含。不可从半截句子开始抄。
- **表格**：固定输出空字符串 ""。
- **⚠️ 拆分规则同「条目原子化拆分规则」**：假设检验声明后，效应量估计、置信区间方法、统计模型、协变量、缺失值处理、敏感性分析等各自独立为一条。
- 如 FAS 和 PPS 方法相同，分析人群写 "FAS、PPS"，条目只写一份。

### 受试者分布分析、人口学与基线特征分析、次要终点分析、安全性分析
每个都是对象 {"分析人群": "X", "表格": "", "条目": [...]}。
- **分析人群**：写明 "FAS"、"PPS"、"SS"、"FAS、PPS" 等。
- **表格**：固定输出空字符串 ""，不要填任何内容。
- **条目**：逐字抄录原文语句，不包含人群标注前缀。
- 如 FAS 和 PPS 方法相同，合并为一份。方法不同则分开。
- 原文写「根据变量的特征进行分组统计描述和比较」而未指定具体方法时，原样抄录。

### ⚠️ 条目原子化拆分规则（极其重要）

**每个数组元素只能描述一个独立的分析步骤。** 原文一段话包含多个分析步骤时，必须拆成多条。

拆分判断标准：
- 一个独立的统计方法/检验 = 一条（如「采用协方差分析比较组间差异」= 一条）
- 方法的具体参数（协变量、显著性水平）= 紧跟该方法的子句，归入同一条
- 切换了分析对象或方法类型 = 新开一条

拆分示例：
  原文：「次要终点采用协方差分析比较两组差异，将年龄和基线值作为协变量，缺失数据采用LOCF方法填补」
  → 拆为两条：
    "采用协方差分析比较两组次要终点的差异，以年龄和基线值作为协变量"
    "缺失数据采用末次观测值结转(LOCF)方法填补"

  原文：「对分类变量采用卡方检验或Fisher精确检验，连续变量采用t检验或Wilcoxon秩和检验」
  → 拆为两条：
    "分类变量采用卡方检验或Fisher精确检验进行比较"
    "连续变量采用t检验或Wilcoxon秩和检验进行比较"

**不要合并**：即使原文在同一句话里，只要语义上有独立的分析步骤就必须拆。拆分后每条都要是可以独立核查的原子语句。

### 缺失数据处理（纯技术条目，不绑定分析人群）
数组，每项是一个对象 {"适用范围": "...", "处理方法": "..."}。
- **适用范围**：从原文抄录该处理针对的分析或场景。如原文写「主要疗效指标缺失数据将采取LOCF」，则适用范围为「主要疗效指标」。
- **处理方法**：从原文抄录具体的处理方式。如「末次观测值结转(LOCF)」。
- 如原文对同一数据有多种处理方法（如既提到LOCF又提到最差填补），全部列出，不可遗漏。

### 期中分析、亚组分析、敏感性分析、其他分析
仅在原文明确计划进行时才返回。格式同上：{"分析人群": "...", "条目": [...]}

## 按需返回
只返回原文中确实描述了该分析的类别字段。原文未提及或写「不适用」的类别不返回。
但以下必须返回：分析人群、统计通用原则、样本量与检验参数、合格标准、主要终点分析、次要终点分析、安全性分析、缺失数据处理。

## 重要：区分统计分析条目与附录清单
方案统计分析章节中可能会提到一些不属于TFL表格的内容：
- **受试者级别明细列表**（如「脱落病例清单」「受试者ID级别明细」「不良事件清单」「合并用药清单」等）——这些是附录中的数据列表（listing），**不是统计分析表格**，**不要抄录到任何分析类别的条目中**
- **只在「条目」中抄录统计分析方法描述**，不要抄录「XXX清单见附录X」这类指向性语句
- 如果方案的某段既描述了统计分析又提到了清单，只抄录统计分析部分，忽略清单引用

## 完整性
- 仔细阅读全文，不要遗漏任何一段描述
- 如果同一事项在多处有不同描述，全部抄录并标注「[注意：与XXX存在潜在矛盾]」
- 试验设计信息本身不构成一条分析
- 一般统计原则如仅为描述性框架未绑定具体指标，归入「统计通用原则」"""

# JSON 输出格式追加
STRUCTURED_SYSTEM_PROMPT_JSON = STRUCTURED_SYSTEM_PROMPT + """

## 输出格式要求
你必须以纯 JSON 格式返回结果，不要用 ```json``` 代码块包裹。

输出示例（注意：分析人群必须放在第一位）：
{
  "统计分析计划": {
    "分析人群": {
      "FAS": "全分析集 — 原文抄录...",
      "PPS": "符合方案集 — 原文抄录...",
      "SS": "安全性分析集 — 原文抄录..."
    },
    "统计通用原则": ["原文语句1", "原文语句2"],
    "样本量与检验参数": ["参数1 — 数值", "参数2 — 数值"],
    "合格标准": "原文完整语句",
    "主要终点分析": {
      "分析人群": "FAS、PPS",
      "表格": "",
      "条目": [
        "完整的假设检验陈述（含H0/H1/α/β）...",
        "后续效应量估计和统计模型..."
      ]
    },
    "次要终点分析": {
      "分析人群": "FAS、PPS",
      "表格": "",
      "条目": [
        "采用协方差分析比较两组次要终点的差异，以年龄和基线值作为协变量",
        "缺失数据采用末次观测值结转(LOCF)方法填补",
        "对亚组按年龄分层（<65岁/≥65岁）进行探索性分析"
      ]
    },
    "安全性分析": {
      "分析人群": "SS",
      "表格": "",
      "条目": ["原文分析语句1", "原文分析语句2"]
    },
    "缺失数据处理": [
      {"适用范围": "主要疗效指标", "处理方法": "末次观测值结转(LOCF)"}
    ]
  }
}"""

# ============================================================================
# 辅助函数
# ============================================================================

ALL_CATEGORIES = [
    "分析人群", "统计通用原则", "期中分析",
    "样本量与检验参数", "合格标准",
    "受试者分布分析", "人口学与基线特征分析",
    "主要终点分析", "次要终点分析", "安全性分析",
    "亚组分析", "敏感性分析", "缺失数据处理", "其他分析",
]


def _unwrap(result: dict) -> dict:
    """解包顶层 key，返回统计分析计划内的 dict。"""
    if "统计分析计划" in result:
        return result["统计分析计划"]
    return result


def safe_name(name: str) -> str:
    """安全的文件名前缀。"""
    return name.replace("/", "_").replace("\\", "_")[:60]


def find_toc_md(sname: str, toc_dir: str) -> str | None:
    """在 TOC 目录中查找匹配的 .md TOC 文件。"""
    pat = os.path.join(toc_dir, f"*_{sname}.md")
    matches = sorted(glob.glob(pat))
    return matches[0] if matches else None


# ============================================================================
# 第一阶段：从 docx 提取统计分析章节（两轮策略）
# ============================================================================

# 摘要章节关键词（第二轮表格提取用）
_SUMMARY_KEYWORDS = ["方案摘要", "摘要", "方案概要"]

# 表格行关键词
_TABLE_ROW_KEYWORDS = ["统计"]

# LLM 分析 TOC 的 prompt
_FIND_STATS_HEADING_PROMPT = """You are a clinical trial document analyst. Given a protocol's table of contents (TOC), identify the single highest-level heading that covers the "statistical analysis" section.

This is the chapter that contains: statistical methods, analysis populations (FAS/PPS/SS), sample size calculation, endpoint analysis methods, statistical software, hypothesis testing framework, etc.

Return a JSON object with two fields:
- "heading": the EXACT heading text as it appears in the TOC
- "core_subject": the heading text with the numbering prefix removed (only the words describing the topic)

For example, from TOC entry "11 统计分析结果" or "六、统计学考虑" or "6. Statistical Considerations":
{"heading": "11 统计分析结果", "core_subject": "统计分析结果"}

If no statistical analysis chapter is found: {"heading": "", "core_subject": ""}"""


def _find_stats_heading(client, toc_headings: list[dict]) -> tuple[str, str]:
    """用 LLM 从 TOC 中识别统计分析章节的最高级标题。

    返回 (heading, core_subject):
      - heading: TOC 中的精确标题文本（用于 TOC 匹配）
      - core_subject: 去掉编号前缀的纯标题文本（用于正文搜索）
    """
    from 提取方案内容 import toc_to_text

    toc_text = toc_to_text(toc_headings)
    if not toc_text.strip():
        return "", ""

    resp = client.messages.create(
        model=API_MODEL,
        max_tokens=256,
        temperature=0,
        system=_FIND_STATS_HEADING_PROMPT,
        messages=[{"role": "user", "content": toc_text}],
        thinking={"type": "disabled"},
    )

    text = "".join(b.text for b in (resp.content or []) if hasattr(b, 'text'))
    try:
        result = json.loads(text)
    except json.JSONDecodeError:
        result = {}
    return result.get("heading", ""), result.get("core_subject", "")


def _find_summary_para_index(doc: DocxDocument) -> int:
    """找到摘要章节的起始段落索引，找不到返回 -1。"""
    for kw in _SUMMARY_KEYWORDS:
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name if para.style else "") or ""
            is_heading = (
                "heading" in style_name.lower()
                or "标题" in style_name
                or "Heading" in style_name
            )
            if is_heading and kw in text and len(text) < 80:
                return i
    # 回退：不限 heading 样式
    for kw in _SUMMARY_KEYWORDS:
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if kw in text and len(text) < 80:
                return i
    return -1


def _extract_from_tables(docx_path: str, keywords: list[str]) -> str | None:
    """从方案摘要表格中按关键词行提取内容，拼成 Markdown。"""
    doc = DocxDocument(docx_path)
    para_idx = _find_summary_para_index(doc)

    if para_idx >= 0:
        body = doc.element.body
        start_element = doc.paragraphs[para_idx]._element
        end_element = None
        for i in range(para_idx + 1, len(doc.paragraphs)):
            style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else ""
            if style.startswith("Heading") or style.startswith("heading"):
                end_element = doc.paragraphs[i]._element
                break

        tables = []
        started = False
        for child in body:
            if child is start_element:
                started = True
                continue
            if end_element is not None and child is end_element:
                break
            if not started:
                continue
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "tbl":
                tables.append(DocxTable(child, doc))
    else:
        tables = list(doc.tables)

    if not tables:
        return None

    found_parts = []
    for table in tables:
        for row in table.rows:
            first_cell = row.cells[0].text.strip().replace("\n", " ")
            for kw in keywords:
                if kw in first_cell:
                    if len(row.cells) == 1:
                        content = f"**{first_cell}**"
                    elif len(row.cells) == 2:
                        content = f"**{first_cell}**\n\n{row.cells[1].text.strip()}"
                    else:
                        parts = [f"**{first_cell}**"]
                        for j in range(1, len(row.cells)):
                            ct = row.cells[j].text.strip()
                            if ct:
                                parts.append(f"  {ct}")
                        content = "\n".join(parts)
                    found_parts.append(f"## {kw}\n\n{content}")
                    break

    if not found_parts:
        return None

    return "# 方案摘要（表格提取）\n\n" + "\n\n".join(found_parts)


def extract_stats_chapter(
    client: anthropic.Anthropic,
    docx_path: str,
    toc_dir: str,
    out_dir: str,
    verbose: bool = True,
) -> dict:
    """
    从 docx 中提取统计分析章节。

    LLM 分析 TOC → 找到统计章节标题 → extract_section 提取 → 质量判断。

    返回: {
        "md_text": str,           # 最终提取的 Markdown 文本
        "md_path": str | None,     # Markdown 文件保存路径
        "status": str,             # "成功" / "待定"
        "heading": str | None,     # 识别的统计章节标题
        "pass": bool,              # 质量判断是否通过
        "missing": str,            # 质量判断缺失项
        "rounds": list[dict],      # 每轮提取详情
    }
    """
    base = os.path.splitext(os.path.basename(docx_path))[0]
    sname = safe_name(base)
    toc_file = find_toc_md(sname, toc_dir)

    result = {
        "md_text": "",
        "md_path": None,
        "status": "待定",
        "heading": None,
        "pass": False,
        "missing": "",
        "rounds": [],
    }

    if not toc_file:
        result["status"] = "无TOC"
        result["missing"] = "未找到章节目录文件"
        return result

    os.makedirs(out_dir, exist_ok=True)

    # ---- 第0步：LLM 从 TOC 识别统计章节标题 ----
    toc_headings = parse_toc(toc_file)
    stats_heading, core_subject = _find_stats_heading(client, toc_headings) if toc_headings else ("", "")

    if verbose:
        print(f"  LLM 识别统计章节标题: {stats_heading or '(未识别)'}"
              + (f"  (core: {core_subject})" if core_subject else ""))

    # ---- 第一轮：标题定位 + extract_section（格式无关） ----
    round1_ok = False
    round1_pass = False
    round1_missing = "提取失败"
    md_text = ""

    if stats_heading:
        try:
            doc = DocxDocument(docx_path)
            # 优先 TOC 锚点模式（对无 Heading 样式文档也能正确识别边界）
            if toc_headings:
                # core_subject 作为正文搜索回退（TOC 文本含编号前缀时正文可能只有纯标题）
                md_text = extract_section_toc(doc, stats_heading, toc_headings,
                                              search_text=core_subject)
            else:
                md_text = extract_section_func(doc, stats_heading)
            round1_ok = not md_text.strip().startswith("[错误]")
        except (ValueError, Exception) as e:
            md_text = f"[错误] 提取异常: {e}"
            round1_ok = False
    else:
        md_text = "[错误] LLM 未从 TOC 中识别到统计分析章节标题"
        round1_ok = False

    if round1_ok:
        round1_pass, round1_missing = _judge_quality(client, md_text)

    result["rounds"].append({
        "keyword": stats_heading or "(无)",
        "extracted": round1_ok,
        "pass": round1_pass,
        "missing": round1_missing,
    })

    if round1_pass:
        md_path = os.path.join(out_dir, f"{sname}_统计分析.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_text)
        result["md_text"] = md_text
        result["md_path"] = md_path
        result["heading"] = stats_heading
        result["pass"] = True
        result["status"] = "成功"
        result["missing"] = ""
        return result

    # ---- 第二轮：表格提取（仅方案 docx 内含统计表格时作为补充） ----
    # 方案通常是纯文字文档，表格搜索命中率极低。Round 1 已有章节内容时跳过。
    if round1_ok:
        md_text_table = _extract_from_tables(docx_path, _TABLE_ROW_KEYWORDS)
        round2_ok = md_text_table is not None
        round2_pass = False
        round2_missing = "表格中未找到统计相关内容"

        if round2_ok:
            round2_pass, round2_missing = _judge_quality(client, md_text_table)
    else:
        round2_ok = False
        round2_pass = False
        round2_missing = "跳过（章节提取未找到内容，不执行表格搜索）"

    result["rounds"].append({
        "keyword": "统计(表格)",
        "extracted": round2_ok,
        "pass": round2_pass,
        "missing": round2_missing,
    })

    if round2_pass:
        md_path = os.path.join(out_dir, f"{sname}_统计分析.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_text_table)
        result["md_text"] = md_text_table
        result["md_path"] = md_path
        result["heading"] = "统计(表格)"
        result["pass"] = True
        result["status"] = "成功"
        result["missing"] = ""
        return result

    # 两轮都没过——用更有信息量的错误原因
    result["md_text"] = md_text if round1_ok else (md_text_table or "")
    result["status"] = "待定"
    # 优先报 Round 1 的真实原因（LLM 质量判断），而非 Round 2 的 "表格中未找到"
    if round1_ok:
        result["missing"] = round1_missing
    else:
        result["missing"] = round1_missing
    return result


def _judge_quality(client: anthropic.Anthropic, md_text: str) -> tuple[bool, str]:
    """用 LLM 判断提取的 Markdown 是否满足质量要求。"""
    if md_text.strip().startswith("[错误]"):
        return False, "上游提取失败"

    resp = client.messages.create(
        model=API_MODEL,
        max_tokens=256,
        system=JUDGE_PROMPT,
        messages=[{"role": "user", "content": md_text}],
        tools=[{"name": "judge", "description": "判断结果", "input_schema": JUDGE_SCHEMA}],
        tool_choice={"type": "tool", "name": "judge"},
        thinking={"type": "disabled"},
    )
    for b in resp.content:
        if b.type == "tool_use" and b.name == "judge":
            return b.input.get("pass", False), b.input.get("missing", "")
    return False, "模型未输出"


# ============================================================================
# 第二阶段：LLM 结构化提取（Markdown → JSON）
# ============================================================================


def _call_structured_api(client: anthropic.Anthropic, md_text: str) -> dict | None:
    """单次 API 调用，返回结构化 dict 或 None。"""
    response = client.messages.create(
        model=API_MODEL,
        max_tokens=8192,
        temperature=0,
        system=STRUCTURED_SYSTEM_PROMPT_JSON,
        messages=[
            {
                "role": "user",
                "content": f"请逐字抄录以下统计分析章节中的关键信息，禁止改写、概括或换词：\n\n{md_text}",
            }
        ],
        thinking={"type": "disabled"},
    )

    full_text = ""
    for block in response.content:
        if block.type == "text":
            full_text += block.text
        elif block.type == "tool_use":
            inp = block.input
            if isinstance(inp, dict) and len(inp) > 0:
                return inp

    if not full_text.strip():
        return None

    # 清理并解析 JSON
    json_text = full_text.strip()
    json_text = re.sub(r"^```(?:json)?\s*\n?", "", json_text)
    json_text = re.sub(r"\n?```\s*$", "", json_text)

    try:
        result = json.loads(json_text)
        if isinstance(result, dict) and len(result) > 0:
            return result
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", full_text, re.DOTALL)
        if m:
            try:
                result = json.loads(m.group())
                if isinstance(result, dict) and len(result) > 0:
                    return result
            except json.JSONDecodeError:
                pass

    return None


def _ensure_population(client: anthropic.Anthropic, result: dict, md_text: str, verbose: bool) -> dict:
    """检查并补全分析人群。缺了就单独再问一次。"""
    plan = result.get("统计分析计划", result)
    if "分析人群" in plan:
        return result

    if verbose:
        print("  ⚠️ 缺少「分析人群」，补提取...")

    for attempt in range(3):
        resp = client.messages.create(
            model=API_MODEL, max_tokens=1024, temperature=0,
            system="从临床方案中找出 FAS/PPS/SS 三个分析集的定义或用途。返回纯JSON没有废话：\n{\"分析人群\": {\"FAS\": \"原文抄录...\", \"PPS\": \"原文抄录...\", \"SS\": \"原文抄录...\"}}\n如无完整定义，抄录用途说明并加「（原文未明确定义）」",
            messages=[{"role": "user", "content": md_text[:3000]}],
            thinking={"type": "disabled"},
        )
        text = "".join(b.text for b in resp.content if b.type == "text")
        if not text.strip():
            continue
        text = re.sub(r"^```(?:json)?\s*\n?", "", text.strip())
        text = re.sub(r"\n?```\s*$", "", text)
        try:
            pop = json.loads(text)
        except json.JSONDecodeError:
            m = re.search(r"\{.*\}", text, re.DOTALL)
            if m:
                try:
                    pop = json.loads(m.group())
                except json.JSONDecodeError:
                    continue
            else:
                continue
        if isinstance(pop, dict) and "分析人群" in pop:
            ordered = {"分析人群": pop["分析人群"]}
            ordered.update(plan)
            if verbose:
                print("  ✅ 分析人群已补全")
            return {"统计分析计划": ordered}

    if verbose:
        print("  ⚠️ 分析人群补提取失败")
    return result


def extract_structured(
    client: anthropic.Anthropic,
    md_text: str,
    max_retries: int = 3,
    verbose: bool = True,
) -> dict | None:
    """LLM 结构化提取统计分析要素。返回 dict 或 None。"""
    if md_text.strip().startswith("[错误]"):
        error_msg = md_text.strip().split("\n")[0]
        if verbose:
            print(f"  ⚠️ 源文件解析失败，跳过结构化提取: {error_msg}")
        return None

    for attempt in range(max_retries):
        result = _call_structured_api(client, md_text)
        if result is not None:
            result = _strip_empty(result)
            result = _ensure_population(client, result, md_text, verbose)
            return result
        if attempt < max_retries - 1:
            wait = (attempt + 1) * 2
            if verbose:
                print(f"  ⚠️ API 返回空/解析失败，{wait}s 后重试 ({attempt + 1}/{max_retries})...")

    if verbose:
        print(f"  ❌ 重试 {max_retries} 次后仍失败")
    return None


def _strip_empty(result: dict) -> dict:
    """去掉空值字段（递归处理嵌套结构）。"""
    if "统计分析计划" not in result:
        return result
    plan = result["统计分析计划"]
    cleaned = {}
    for k, v in plan.items():
        if v is None:
            continue
        if isinstance(v, dict):
            entries = v.get("条目", [])
            if not entries and not v.get("分析人群"):
                continue
        if isinstance(v, list):
            if k == "缺失数据处理":
                # 保留有处理方法的条目（适用范围可为空，标记为「未指定」）
                filtered = []
                for x in v:
                    if not isinstance(x, dict):
                        continue
                    if not x.get("处理方法"):
                        continue  # 完全空条目丢弃
                    if not x.get("适用范围"):
                        x["适用范围"] = "（原文未明确指定适用范围）"
                    filtered.append(x)
                if not filtered:
                    continue
                cleaned[k] = filtered
                continue
            if len(v) == 0:
                continue
        if isinstance(v, str) and not v.strip():
            continue
        cleaned[k] = v
    return {"统计分析计划": cleaned}


def print_summary(result: dict, elapsed: float) -> None:
    """打印结构化提取结果摘要。"""
    plan = _unwrap(result)
    for cat in ALL_CATEGORIES:
        val = plan.get(cat)
        if val is None:
            continue
        if isinstance(val, dict):
            entries = val.get("条目", [])
            pop = val.get("分析人群", "")
            if entries:
                print(f"  【{cat}】{pop}  {len(entries)} 项")
        elif isinstance(val, list):
            print(f"  【{cat}】 {len(val)} 项")
        elif isinstance(val, str):
            print(f"  【{cat}】 1 项")

    for cat in ["分析人群", "主要终点分析", "次要终点分析", "安全性分析", "缺失数据处理"]:
        val = plan.get(cat)
        if val is None:
            continue
        print(f"\n  ── {cat} ──")
        if isinstance(val, dict):
            if cat == "分析人群":
                for pop_name in ["FAS", "PPS", "SS"]:
                    if pop_name in val:
                        text = val[pop_name]
                        print(f"    {pop_name}: {text[:100]}{'...' if len(text)>100 else ''}")
            else:
                pop = val.get("分析人群", "")
                print(f"    分析人群: {pop}")
                for item in val.get("条目", [])[:8]:
                    print(f"    • {item[:100]}{'...' if len(item)>100 else ''}")
                if len(val.get("条目", [])) > 8:
                    print(f"    ... 共 {len(val.get('条目', []))} 项")
        elif isinstance(val, list):
            for item in val[:8]:
                if isinstance(item, dict):
                    print(f"    [{item.get('适用范围','')}] {item.get('处理方法','')[:100]}")
                else:
                    print(f"    • {str(item)[:100]}")
            if len(val) > 8:
                print(f"    ... 共 {len(val)} 项")


def fmt_line(i: int, total: int, name: str, status: str, keyword: str | None,
             rounds: list[dict] | None = None) -> str:
    """格式化单行输出。"""
    display = name if len(name) <= 52 else name[:49] + "..."

    mark = {"成功": "✅", "待定": "⏳"}.get(status, "❌")

    if keyword:
        detail = f"（关键词={keyword}）"
    elif rounds:
        attempts = " → ".join(
            f"{rd['keyword']}({'✓' if rd.get('pass') else '✗'})" for rd in rounds
        )
        detail = f"（尝试: {attempts}）"
    else:
        detail = ""

    return f"[{i:2d}/{total}] {mark} {display:<52} {status} {detail}"


# ============================================================================
# 主处理流程
# ============================================================================


def process_one(
    client: anthropic.Anthropic,
    docx_path: str,
    toc_dir: str,
    out_dir: str,
    save_md: bool = True,
    verbose: bool = True,
) -> dict:
    """
    处理单个方案文件的完整流程：docx → Markdown → 质量判断 → 结构化 JSON。

    返回:
        {
            "name": str,
            "status": "成功" | "待定" | "失败",
            "extraction_pass": bool,     # 章节提取质量判断
            "extraction_missing": str,
            "structured": dict | None,   # 结构化要素
            "json_path": str | None,     # JSON 保存路径
            "md_path": str | None,       # Markdown 保存路径
            "elapsed_phase1": float,     # 阶段一耗时
            "elapsed_phase2": float,     # 阶段二耗时
        }
    """
    base = os.path.splitext(os.path.basename(docx_path))[0]
    sname = safe_name(base)

    if verbose:
        print(f"\n{'━' * 65}")
        print(f"  📄 {base}")
        print(f"{'━' * 65}")

    # ---- 阶段一：提取统计分析章节 ----
    t0 = time.monotonic()
    chapter = extract_stats_chapter(client, docx_path, toc_dir, out_dir, verbose=verbose)
    elapsed_phase1 = time.monotonic() - t0

    result = {
        "name": base,
        "status": chapter["status"],
        "extraction_pass": chapter["pass"],
        "extraction_missing": chapter["missing"],
        "structured": None,
        "json_path": None,
        "md_path": None,
        "elapsed_phase1": elapsed_phase1,
        "elapsed_phase2": 0,
    }

    # 保存 Markdown（中间产物）
    if save_md and chapter["md_text"]:
        md_filename = f"{sname}_统计分析.md"
        md_path = os.path.join(out_dir, md_filename)
        # 如果是第二轮表格提取且覆盖了第一轮，则先写临时文件再覆盖
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(chapter["md_text"])
        result["md_path"] = md_path
        if verbose:
            print(f"  📝 Markdown 已保存 → {md_filename}  ({elapsed_phase1:.0f}s)")

    if not chapter["pass"] and verbose:
        print(f"  ⚠️ 章节提取未通过质量判断: {chapter['missing']}")

    # ---- 阶段二：结构化提取 ----
    if not chapter["md_text"] or chapter["md_text"].strip().startswith("[错误]"):
        result["status"] = "失败"
        # 把具体失败原因写入 status 和 missing 字段，方便上层诊断
        reason = chapter.get("missing", "") or chapter.get("status", "未知原因")
        if chapter["md_text"] and chapter["md_text"].strip().startswith("[错误]"):
            reason = chapter["md_text"].strip()
        result["extraction_missing"] = reason
        if verbose:
            print(f"  ❌ 无可提取的章节内容: {reason}")
        return result

    t0 = time.monotonic()
    structured = extract_structured(client, chapter["md_text"], verbose=verbose)
    elapsed_phase2 = time.monotonic() - t0
    result["elapsed_phase2"] = elapsed_phase2

    if structured:
        result["structured"] = structured
        result["status"] = "成功"

        # 保存 JSON
        json_name = f"{sname}_统计分析_分析提取.json"
        json_path = os.path.join(out_dir, json_name)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(structured, f, ensure_ascii=False, indent=2)
        result["json_path"] = json_path

        if verbose:
            print_summary(structured, elapsed_phase2)
            print(f"\n  💾 JSON 已保存 → {json_path}  ({elapsed_phase2:.0f}s)")
    else:
        if result["status"] == "成功":
            result["status"] = "结构化失败"
        if verbose:
            print(f"  ❌ 结构化提取失败 ({elapsed_phase2:.0f}s)")

    return result


def batch_process(
    client: anthropic.Anthropic,
    docx_files: list[str],
    toc_dir: str,
    out_dir: str,
    save_md: bool = True,
    verbose: bool = True,
) -> list[dict]:
    """批量处理多个方案文件。"""
    results = []

    for i, docx in enumerate(docx_files):
        r = process_one(client, docx, toc_dir, out_dir, save_md=save_md, verbose=verbose)
        results.append(r)

        # 每处理完一个就打印汇总行
        if verbose and len(docx_files) > 1:
            print(fmt_line(
                i + 1, len(docx_files), r["name"], r["status"],
                keyword=None,  # 简化显示
            ))

    return results


# ============================================================================
# 确保 TOC 存在
# ============================================================================


def ensure_toc_exists(client, scheme_dir: str, out_dir: str, docx_files: list[str]) -> bool:
    """检查并生成 TOC 文件（放在输出目录内）。如果 TOC 已存在则跳过。
    使用 LLM client 确保 TOC 层级正确（即使 OOXML 样式名缺失）。"""
    from 提取方案目录 import process_protocol, save_results

    missing_files = []
    for docx in docx_files:
        base = os.path.splitext(os.path.basename(docx))[0]
        sname = safe_name(base)
        if not find_toc_md(sname, out_dir):
            missing_files.append(docx)

    if not missing_files:
        return True

    print(f"📋 正在生成方案章节目录（TOC），共 {len(missing_files)} 个文件...")
    for docx in missing_files:
        try:
            r = process_protocol(docx, client=client, model=API_MODEL)
            save_results([r], out_dir)
        except Exception as e:
            print(f"❌ TOC 生成失败 ({os.path.basename(docx)}): {e}")
            return False
        base = os.path.splitext(os.path.basename(docx))[0]
        sname = safe_name(base)
        if not find_toc_md(sname, out_dir):
            print(f"❌ TOC 文件未生成 ({os.path.basename(docx)})")
            return False

    print("✅ TOC 生成完毕")
    return True


# ============================================================================
# CLI
# ============================================================================


def main():
    parser = argparse.ArgumentParser(
        description="临床试验方案统计分析要素提取 —— docx → 结构化 JSON（一体化）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python 统计分析提取_一体化.py --file "方案/xxx.docx"
  python 统计分析提取_一体化.py --file "方案/xxx.docx" --out-dir "我的输出"
  python 统计分析提取_一体化.py --dir "方案"
  python 统计分析提取_一体化.py --dry-run
        """,
    )
    parser.add_argument("--file", "-f", default=None, help="只处理单个 docx 文件")
    parser.add_argument("--dir", default=None, help=f"批量处理目录下所有 docx（默认: {DEFAULT_SCHEME_DIR}）")
    parser.add_argument("--out-dir", default=DEFAULT_OUT_DIR, help=f"输出目录，所有产物放一起（默认: {DEFAULT_OUT_DIR}）")
    parser.add_argument("--dry-run", action="store_true", help="仅列出方案文件，不执行")
    parser.add_argument("--quiet", "-q", action="store_true", help="静默模式，减少输出")
    args = parser.parse_args()

    # ---- 收集方案文件 ----
    if args.file:
        docx_path = args.file if os.path.isabs(args.file) else os.path.join(PROJECT_DIR, args.file)
        docx_files = [docx_path] if os.path.exists(docx_path) else []
        if not docx_files:
            print(f"❌ 文件不存在: {docx_path}")
            sys.exit(1)
    else:
        scheme_dir = args.dir or DEFAULT_SCHEME_DIR
        if not os.path.isabs(scheme_dir):
            scheme_dir = os.path.join(PROJECT_DIR, scheme_dir)
        docx_files = sorted(glob.glob(os.path.join(scheme_dir, "*.docx")))
        docx_files = [f for f in docx_files if not os.path.basename(f).startswith("~$")]

    if not docx_files:
        print("❌ 未找到 docx 方案文件")
        sys.exit(1)

    print("=" * 65)
    print(f"  临床试验方案统计分析要素提取（一体化）")
    print(f"  模型: {API_MODEL}")
    print(f"  方案数量: {len(docx_files)}")
    print(f"  输出目录: {args.out_dir}")
    print("=" * 65)

    if args.dry_run:
        for i, f in enumerate(docx_files, 1):
            print(f"  [{i:2d}] {os.path.basename(f)}")
        return

    # ---- 确保 TOC 存在（放在输出目录内） ----
    scheme_dir = os.path.dirname(docx_files[0]) if args.file else (args.dir or DEFAULT_SCHEME_DIR)
    if not os.path.isabs(scheme_dir):
        scheme_dir = os.path.join(PROJECT_DIR, scheme_dir)
    out_dir = args.out_dir if os.path.isabs(args.out_dir) else os.path.join(PROJECT_DIR, args.out_dir)

    os.makedirs(out_dir, exist_ok=True)

    # ---- 初始化客户端 ----
    client = anthropic.Anthropic(base_url=API_BASE_URL, auth_token=API_AUTH_TOKEN, timeout=180, max_retries=3)

    if not ensure_toc_exists(client, scheme_dir, out_dir, docx_files):
        sys.exit(1)

    # ---- 执行提取 ----
    pipeline_start = time.monotonic()
    results = batch_process(
        client, docx_files, out_dir, out_dir,
        save_md=True,
        verbose=not args.quiet,
    )
    total_elapsed = time.monotonic() - pipeline_start

    # ---- 汇总 ----
    success = sum(1 for r in results if r["structured"] is not None)
    fail = len(results) - success

    print(f"\n{'=' * 65}")
    print(f"  处理完毕")
    print(f"  总耗时: {total_elapsed:.0f}s")
    print(f"  结构化成功: {success}/{len(results)}")
    if fail:
        failed_names = [r["name"][:40] for r in results if r["structured"] is None]
        print(f"  失败: {', '.join(failed_names)}")
        # 打印每个失败文件的详细原因
        for r in results:
            if r["structured"] is None:
                reason = r.get("extraction_missing", "") or r.get("status", "未知")
                print(f"    [{r['name'][:50]}] 原因: {reason}")
    print(f"  输出目录: {out_dir}")
    print(f"{'=' * 65}")

    if success == 0:
        print("❌ 所有方案文件提取失败，请检查：1) docx 是否包含统计分析章节 2) 是否使用标准 Heading 样式 3) API Key 是否有效")
        sys.exit(1)


if __name__ == "__main__":
    main()
