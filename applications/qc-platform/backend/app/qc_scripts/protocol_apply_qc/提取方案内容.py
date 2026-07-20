"""
从 Word 文档中按标题提取章节内容为 Markdown。

用法:
    python 提取方案内容.py <docx路径> "<标题文本>"
    python 提取方案内容.py <docx路径> "<标题文本>" --toc <toc.md>  # TOC锚点模式

提取逻辑:
  A) TOC锚点模式（推荐）:
     TOC中匹配目标标题 → TOC中找下一个同级标题作为边界
     → 纯文本搜索在正文中定位起点终点 → extract_to_markdown
     完全不依赖Word样式，格式无关。

  B) 纯样式模式（无TOC时回退）:
     docx段落中搜索标题 → get_heading_level确定边界
     → extract_to_markdown
"""

import argparse
import re
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table


# ---------------------------------------------------------------------------
# 标题识别
# ---------------------------------------------------------------------------

def get_heading_level(para) -> int | None:
    """返回段落的标题级别（1-9），非标题段落返回 None。"""
    if para.style and para.style.name.startswith("Heading"):
        try:
            return int(para.style.name.split()[-1])
        except (ValueError, IndexError):
            return None
    style_name = (para.style.name if para.style else "")
    for prefix in ("Heading", "heading", "标题", "标题 "):
        if style_name.startswith(prefix):
            try:
                return int(style_name[len(prefix):].strip())
            except ValueError:
                pass
    return None


# ---------------------------------------------------------------------------
# Markdown 渲染
# ---------------------------------------------------------------------------

def _runs_to_md(para) -> str:
    """将段落的 runs 转为 Markdown 内联格式。"""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


def _table_to_md(table) -> str:
    """将 docx Table 转为 Markdown 表格。"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ").replace("|", "\\|")
                 for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")
    lines = [f"| {' | '.join(rows[0])} |",
             f"| {' | '.join(['---'] * max_cols)} |"]
    for row in rows[1:]:
        lines.append(f"| {' | '.join(row)} |")
    return "\n".join(lines)


def _has_image(para) -> bool:
    return bool(para._element.findall('.//' + qn('w:drawing')))


# ---------------------------------------------------------------------------
# 正文段落搜索（语言无关，纯文本匹配）
# ---------------------------------------------------------------------------

def _toc_entry_para(para) -> bool:
    """判断段落是否为目录条目（标题后跟 tab + 页码）。"""
    return bool(re.search(r'\t\s*\d+\s*$', para.text))


def _is_toc_style(para) -> bool:
    """判断段落样式是否为 TOC 目录样式（toc 1 / TOC 2 等）。"""
    style_name = (para.style.name if para.style else '') or ''
    return style_name.lower().startswith('toc')


def _normalize_ws(s: str) -> str:
    """空白字符归一化：连续空白 → 单个空格。"""
    return re.sub(r'\s+', ' ', s).strip()


def _strip_number_prefix(text: str) -> str:
    """递归去除标题开头的编号前缀，覆盖阿拉伯数字、中文数字、括号编号。
    "7. 监查计划" → "监查计划"
    "（一）申办者名称" → "申办者名称"
    "13.6 统计分析" → "统计分析"
    "六、 统计学考虑" → "统计学考虑"
    """
    changed = True
    while changed:
        changed = False
        # 括号编号：（一）/ (一) / （1）/ (1)
        m = re.match(r'^[（(]\s*[\d一二三四五六七八九十]+\s*[）)]\s*', text)
        if m:
            text = text[m.end():]
            changed = True
            continue
        # 阿拉伯数字编号：1. / 13.6 / 6.1. / 7．
        m = re.match(r'^(\d+)[\.\s、．]+(?=\S)', text)
        if m:
            text = text[m.end():]
            changed = True
            continue
        # 中文数字编号：一、/ 六
        m = re.match(r'^[一二三四五六七八九十]+[、\s]+(?=\S)', text)
        if m:
            text = text[m.end():]
            changed = True
    return text.strip()


def _find_para_by_text(doc: Document, text: str, start_from: int = 0,
                       heading_only: bool = False) -> int:
    """在段落中搜索包含 text 的段落索引，跳过目录条目。
    匹配时做空白归一化——TOC 空格 vs 正文 tab 都能命中。

    如果 heading_only=True，仅匹配短段落（大概率是标题而非正文），
    避免正文中出现相同词汇的误匹配。"""
    text_norm = _normalize_ws(text)
    for i in range(start_from, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        # 跳过 TOC 样式段落（toc 1 / TOC 2 等），这些是目录条目不是正文
        if _is_toc_style(para):
            continue
        if not heading_only and _toc_entry_para(para):
            continue
        if heading_only and len(para.text.strip()) > 60:
            continue
        if text in para.text:
            return i
        if _normalize_ws(para.text) == text_norm:
            return i
    return -1


def extract_to_markdown(
    doc: Document,
    start_index: int,
    stop_index: int | None,
) -> str:
    """在 XML body 元素级别遍历，将指定段落范围转为 Markdown。"""
    body = doc.element.body
    paragraphs = doc.paragraphs
    start_element = paragraphs[start_index]._element
    stop_element = paragraphs[stop_index]._element if stop_index is not None else None

    elem_to_para = {p._element: p for p in paragraphs}

    lines = []
    started = False
    prev_empty = False

    for child in body:
        if not started:
            if child is start_element:
                started = True
                para = elem_to_para[child]
                level = get_heading_level(para)
                text = para.text.strip()
                if level and level <= 6:
                    lines.append("#" * level + " " + text)
                else:
                    lines.append(text)
                lines.append("")
            continue

        if stop_element is not None and child is stop_element:
            break

        if child.tag == qn("w:p"):
            para = elem_to_para.get(child)
            if para is None:
                continue

            text = para.text.strip()

            # 子标题（仅识别 Word 样式标题，不推断）
            level = get_heading_level(para)
            if level is not None and level <= 6 and len(text) < 120:
                if prev_empty or not lines or lines[-1] == "":
                    lines.append("#" * level + " " + text)
                else:
                    lines.append("")
                    lines.append("#" * level + " " + text)
                lines.append("")
                prev_empty = True
                continue

            if not text:
                if _has_image(para):
                    lines.append("> 📷 *[图片]*")
                    lines.append("")
                    prev_empty = True
                elif not prev_empty:
                    prev_empty = True
                continue

            md_text = _runs_to_md(para)
            lines.append(md_text)
            prev_empty = False

        elif child.tag == qn("w:tbl"):
            if not prev_empty:
                lines.append("")
            table = Table(child, doc)
            md_table = _table_to_md(table)
            if md_table.strip():
                lines.append(md_table)
                lines.append("")
                prev_empty = True

    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines)


def extract_section(doc: Document, heading_text: str) -> str:
    """提取指定标题及其下所有内容（纯样式模式）。
    依赖 Word Heading 样式做边界检测，无样式时边界失效。
    """
    # 查找起点
    for i, para in enumerate(doc.paragraphs):
        if heading_text in para.text:
            level = get_heading_level(para)
            if level is not None:
                start_index, start_level = i, level
                break
    else:
        # 回退：不要求 heading 样式
        for i, para in enumerate(doc.paragraphs):
            if heading_text in para.text.strip():
                start_index, start_level = i, 1
                break
        else:
            raise ValueError(f"未在文档中找到标题: {heading_text}")

    # 查找终点
    stop_index = None
    for i in range(start_index + 1, len(doc.paragraphs)):
        level = get_heading_level(doc.paragraphs[i])
        if level is not None and level <= start_level:
            stop_index = i
            break

    return extract_to_markdown(doc, start_index, stop_index)


def extract_section_toc(
    doc: Document,
    heading_text: str,
    toc_headings: list[dict],
    search_text: str = "",
) -> str:
    """TOC锚点模式提取章节（语言无关，不依赖Word样式）。

    search_text: LLM 返回的纯标题文本（无编号前缀）。
      当 heading_text 在正文中搜不到时作为回退。
    """
    # 1. 在 TOC 中匹配目标标题
    match_idx = None
    match_level = None
    for i, h in enumerate(toc_headings):
        if heading_text in h["text"]:
            match_idx = i
            match_level = h["level"]
            break

    if match_idx is None:
        raise ValueError(f"TOC 中未找到包含「{heading_text}」的标题")

    target_text = toc_headings[match_idx]["text"]

    # 2. 在 TOC 中向后找边界标题（第一个 level <= match_level 的条目，即下一个同级或更高级章节）
    boundary_text = None
    boundary_level = None
    for i in range(match_idx + 1, len(toc_headings)):
        if toc_headings[i]["level"] <= match_level:
            boundary_text = toc_headings[i]["text"]
            boundary_level = toc_headings[i]["level"]
            break

    # 3. 向前找最近的、能在正文中定位到的 TOC 锚点作为搜索起点
    search_from = 0
    for i in range(match_idx - 1, -1, -1):
        prev_idx = _find_para_by_text(doc, toc_headings[i]["text"], heading_only=True)
        if prev_idx >= 0:
            search_from = prev_idx
            break

    # 4. 在锚点之后找目标标题的正文位置（只搜短段落，避免正文中包含相同词汇）
    #    回退链：精确TOC文本 → 全文回退 → LLM纯标题文本 → 去编号前缀
    start_index = _find_para_by_text(doc, target_text, search_from, heading_only=True)
    if start_index < 0:
        start_index = _find_para_by_text(doc, target_text, heading_only=True)
    if start_index < 0 and search_text and search_text != target_text:
        start_index = _find_para_by_text(doc, search_text, heading_only=True)
    # 兜底：TOC 标题带编号前缀（如"7. 监查计划"）但正文标题不带（如"监查计划"）
    if start_index < 0:
        start_index = _find_para_by_text(doc, _strip_number_prefix(target_text),
                                         heading_only=True)
    if start_index < 0 and search_text and search_text != target_text:
        start_index = _find_para_by_text(doc, _strip_number_prefix(search_text),
                                         heading_only=True)
    if start_index < 0:
        detail = f"，搜索文本「{search_text}」" if search_text else ""
        raise ValueError(f"TOC 匹配到「{target_text}」{detail}，但正文中均未找到")

    # 5. 定位边界标题的正文位置
    #    回退链：TOC精确文本 → 全文回退 → 去编号前缀
    stop_index = None
    if boundary_text is not None:
        stop_index = _find_para_by_text(doc, boundary_text, start_index + 1, heading_only=True)
        if stop_index < 0:
            stop_index = _find_para_by_text(doc, boundary_text, heading_only=True)
        if stop_index < 0:
            stop_index = _find_para_by_text(doc, _strip_number_prefix(boundary_text),
                                            start_index + 1, heading_only=True)
        if stop_index < 0:
            stop_index = _find_para_by_text(doc, _strip_number_prefix(boundary_text),
                                            heading_only=True)

    return extract_to_markdown(doc, start_index, stop_index)


# ---------------------------------------------------------------------------
# TOC 解析（供上游调用）
# ---------------------------------------------------------------------------

def parse_toc(toc_path: str) -> list[dict]:
    """解析 toc.md 为有序标题列表 [{level, text, page}, ...]。
    格式: - **标题文本** `[页码]`
    """
    headings = []
    with open(toc_path, "r", encoding="utf-8") as f:
        for line in f:
            m = re.match(r"^(\s*)-\s+\*\*(.+?)\*\*(?:\s*`\[(\d+)\]`)?", line)
            if m:
                indent = len(m.group(1))
                level = indent // 4 + 1
                text = m.group(2).strip()
                # 清洗尾部的目录点和页码（PAGEREF 字段解析失败时上游会混入）
                text = re.sub(r'\s*\.{3,}.*$', '', text).strip()
                page = m.group(3)
                headings.append({"level": level, "text": text, "page": page})
    return headings


def toc_to_text(toc_headings: list[dict]) -> str:
    """将 TOC 标题列表转为人可读的层级文本，供 LLM 分析。"""
    lines = []
    for h in toc_headings:
        indent = "  " * (h["level"] - 1)
        page_str = f" [p.{h['page']}]" if h.get("page") else ""
        lines.append(f"{indent}{h['text']}{page_str}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="按标题提取 Word 文档中章节为 Markdown",
    )
    parser.add_argument("docx_path", help="Word 文档路径 (.docx)")
    parser.add_argument("heading", help="要提取的标题文本（如 '11 统计分析结果'）")
    parser.add_argument("--toc", "-t", help="TOC 目录文件路径，启用TOC锚点模式（推荐）")
    parser.add_argument("--to-md", "-m", help="输出为 Markdown 文件（默认 stdout）")
    args = parser.parse_args()

    try:
        doc = Document(args.docx_path)
    except FileNotFoundError:
        print(f"[错误] 文件不存在: {args.docx_path}", file=sys.stderr)
        sys.exit(1)

    try:
        if args.toc:
            toc_headings = parse_toc(args.toc)
            result = extract_section_toc(doc, args.heading, toc_headings)
        else:
            result = extract_section(doc, args.heading)
    except ValueError as e:
        print(f"[错误] {e}", file=sys.stderr)
        sys.exit(1)

    if args.to_md:
        with open(args.to_md, "w", encoding="utf-8") as f:
            f.write(result)
        print(f"[完成] → {args.to_md}")
    else:
        print(result)


if __name__ == "__main__":
    main()
