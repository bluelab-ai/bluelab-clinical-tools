#!/usr/bin/env python3
"""
临床试验方案章节目录提取工具

从方案 docx 文件中提取完整的章节目录（标题层级树），
每个方案输出其内部的章节结构，如:
  第一章 试验概述
    1.1 试验背景
    1.2 试验目的
    ...

支持两种 docx 标题标记方式:
  - Heading 1/2/3 段落样式
  - outlineLvl XML 属性

用法:
  python3 extract_toc.py                              # 终端展示所有方案的目录
  python3 extract_toc.py --dir /path/to/project        # 换项目根目录
  python3 extract_toc.py --file 方案/xxx.docx           # 只提取单个方案
  python3 extract_toc.py --json                        # JSON 输出
  python3 extract_toc.py --save output_dir             # 保存到文件夹
"""

import os
import sys
import json
import re
import argparse
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# OOXML 命名空间
_NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_ns = {'w': _NS_W}
nsmap = {'w': _NS_W}

# ═══════════════════════════════════════════════════════════════════════════
# TOC 段落检测（语言无关，与 extract_tables.py 一致）
# ═══════════════════════════════════════════════════════════════════════════


def _is_toc_paragraph(elem):
    """通过 XML 字段码 + 点线分隔符检测 TOC 段落（语言/样式名无关）。"""
    fld = [fc.get(f'{{{nsmap["w"]}}}fldCharType') for fc in elem.findall(f'.//{{{nsmap["w"]}}}fldChar')]
    has_fld = 'begin' in fld and 'separate' in fld
    has_dot = any(t.get(f'{{{nsmap["w"]}}}leader') == 'dot' for t in elem.findall(f'.//{{{nsmap["w"]}}}tab'))
    return has_fld and has_dot


def _toc_text(elem):
    """提取 TOC 段落纯文本（跳过 instrText 字段指令）。"""
    return ''.join(t.text or '' for t in elem.iter(f'{{{nsmap["w"]}}}t')
                   if etree.QName(t.getparent()).localname != 'instrText').strip()


def _p_style(elem):
    """获取段落的样式名。"""
    pPr = elem.find(f'{{{nsmap["w"]}}}pPr')
    if pPr is None: return 'NONE'
    ps = pPr.find(f'{{{nsmap["w"]}}}pStyle')
    return ps.get(f'{{{nsmap["w"]}}}val', 'NONE') if ps is not None else 'NONE'


def _toc_style_level(name):
    """从 TOC 样式名推断层级（语言无关，覆盖 TOC/目录/TM/TDC）。"""
    for pat in [r'^TOC\s*(\d+)$', r'^(目录|目次)\s*(\d+)$', r'^TM\s*(\d+)$', r'^TDC\s*(\d+)$']:
        m = re.match(pat, name, re.IGNORECASE)
        if m: return int(m.group(1) if len(m.groups()) == 1 else m.group(2))
    m = re.match(r'^TOC\s*Heading$', name, re.IGNORECASE)
    return 1 if m else None


def _extract_toc_entries(doc):
    """从文档中提取 TOC 条目，返回 [(level, heading_text, page_number), ...]。

    通过 XML 字段码检测 TOC 段落（语言/样式名无关），
    解析 PAGEREF 字段获取页码。
    """
    entries = []
    raw_entries: list[tuple[int | None, str, str | None]] = []  # (level_or_none, text, page)
    body = doc.element.body

    for p in body.iter(f'{{{nsmap["w"]}}}p'):
        if not _is_toc_paragraph(p):
            continue

        style_name = _p_style(p)
        level = _toc_style_level(style_name)
        # ── FLD 栈解析：分离标题文本和 PAGEREF 页码 ──
        runs = p.findall(f'.//{{{nsmap["w"]}}}r')
        field_stack = []       # [(field_type, is_in_result), ...]
        heading_parts = []
        page_parts = []

        for run in runs:
            fldChar = run.find(f'{{{nsmap["w"]}}}fldChar')
            instrText = run.find(f'{{{nsmap["w"]}}}instrText')
            run_text = ''.join(t.text or '' for t in run.findall(f'.//{{{nsmap["w"]}}}t'))

            if fldChar is not None:
                ct = fldChar.get(f'{{{nsmap["w"]}}}fldCharType', '')
                if ct == 'begin':
                    field_stack.append(['unknown', False])
                    continue
                elif ct == 'separate':
                    if field_stack:
                        field_stack[-1][1] = True
                    continue
                elif ct == 'end':
                    if field_stack:
                        field_stack.pop()
                    continue

            if instrText is not None and instrText.text:
                if field_stack and field_stack[-1][0] == 'unknown':
                    txt = instrText.text.strip()
                    if 'PAGEREF' in txt:
                        field_stack[-1][0] = 'PAGEREF'
                    elif 'HYPERLINK' in txt:
                        field_stack[-1][0] = 'HYPERLINK'
                continue

            if not run_text:
                continue

            in_pageref_result = any(
                ft == 'PAGEREF' and in_res for ft, in_res in field_stack
            )
            if in_pageref_result:
                page_parts.append(run_text)
            else:
                in_code = any(not in_res for _, in_res in field_stack)
                if not in_code:
                    heading_parts.append(run_text)

        heading = ''.join(heading_parts).strip()
        # 清洗尾部的目录点和页码（PAGEREF 字段可能解析失败，点号>=3视为目录引导符）
        heading = re.sub(r'\s*\.{3,}.*$', '', heading).strip()
        page = ''.join(page_parts).strip()

        # 排除 TOC 标题自身（"目录""Table of Contents""目 录"）
        if heading and heading not in ('目录', 'Table of Contents', '目 录', 'Contents'):
            raw_entries.append((level, heading, page if page else None))

    # ── 返回条目，level 可能为 None（表示需要调用方通过 assign_toc_levels 补齐层级）──
    for level, heading, page in raw_entries:
        entries.append((level, heading, page))
    return entries

    return entries


# ═══════════════════════════════════════════════════════════════════════════
# TOC 层级赋值（纯代码实现，格式无关，语言无关）
# ═══════════════════════════════════════════════════════════════════════════


def _numbering_signature(text: str) -> str:
    """提取标题编号前缀的格式签名。

    归一化所有计数符号：N=阿拉伯数字 C=中文数字 R=罗马数字 L=字母
    保留标点结构（、. - （ ）等）。

    同壳 = 同编号惯例 = 同层级。例：
      「六、」「七、」→ C、  「（一）」「（五）」→ （C）
      「1.1」「3.2」→ N.N    「IV.」「V.」→ R.
    """
    s = text.strip()
    s = re.sub(r'\b[IVXLCDM]+\b', 'R', s)           # 罗马数字
    s = re.sub(r'[一二三四五六七八九十百千]+', 'C', s)  # 中文数字
    s = re.sub(r'\d+', 'N', s)                       # 阿拉伯数字
    s = re.sub(r'[a-zA-Z]', 'L', s)                  # 字母
    m = re.match(r'^([NCRL\(\)（）\[\]【】、，,.\s\-—…·]+)', s)
    if m:
        return re.sub(r'\s+', '', m.group(1))
    return ''


def assign_toc_levels(entries: list[tuple[int | None, str, str | None]],
                      client=None, model: str = "") -> list[tuple[int, str, str | None]]:
    """给 TOC 条目赋层级。纯代码实现——按编号壳分组，看首次出现时的嵌套关系排层级。

    不再依赖 LLM，完全确定性。client 和 model 参数保留以兼容旧调用方，但忽略。
    """
    if all(lv is not None for lv, _, _ in entries):
        return [(lv, t, p) for lv, t, p in entries]

    texts = [t for _, t, _ in entries]
    if not texts:
        return []

    # Pass 1: 每个条目提取编号壳 + 记录每种壳的首次出现位置
    shells: list[str] = []
    shell_first_idx: dict[str, int] = {}
    for i, t in enumerate(texts):
        sig = _numbering_signature(t)
        shells.append(sig)
        if sig and sig not in shell_first_idx:
            shell_first_idx[sig] = i

    # Pass 2: 按首次出现顺序排层级
    # 每种壳第一次出现时，向前找最近的不同壳 → 那就是父壳
    shell_level: dict[str, int] = {}
    sorted_shells = sorted(shell_first_idx.items(), key=lambda x: x[1])

    for sig, first_idx in sorted_shells:
        if not sig:
            shell_level[sig] = 1
            continue
        # 向前找最近的不同壳
        parent_sig = ""
        for j in range(first_idx - 1, -1, -1):
            prev = shells[j]
            if prev and prev != sig:
                parent_sig = prev
                break
        if not parent_sig or parent_sig not in shell_level:
            shell_level[sig] = 1
        else:
            shell_level[sig] = shell_level[parent_sig] + 1

    # Pass 3: 逐条赋值，限制最多 4 级
    levels = [shell_level.get(s, 1) for s in shells]
    max_lv = max(levels) if levels else 1
    if max_lv > 4:
        scale = 4 / max_lv
        levels = [max(1, round(l * scale)) for l in levels]

    return [(levels[i], t, p) for i, (_, t, p) in enumerate(entries)]


def get_heading_level(paragraph):
    """
    判断段落是否为标题，返回标题层级 (1-9) 或 None。

    检测方式（按优先级）:
      1. Heading N 样式        → 返回 N    (Word 标准)
      2. outlineLvl XML 属性   → 返回 ol+1 (大纲级别)
      3. 中文数字样式 "1级/2级/3级/4级" → 返回数字
      4. 中文大写样式 "一级/二级/三级"    → 返回数字
      5. 都不是 → 返回 None
    """
    import re

    style_name = paragraph.style.name if paragraph.style else ''
    text = paragraph.text.strip()
    if not text:
        return None

    level = None

    # 排除 TOC/目录类样式
    if 'TOC' in style_name or 'toc' in style_name.lower():
        return None

    # 方式一: Heading N 样式
    if style_name.startswith('Heading') or style_name.startswith('heading'):
        try:
            level = int(style_name.split()[-1])
        except ValueError:
            level = 1

    # 方式二: outlineLvl (如果没有 Heading 样式则用这个)
    # OOXML outlineLvl: 0=Level1 ... 8=Level9, 9=BodyText (正文，排除)
    if level is None:
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            ol = pPr.find(qn('w:outlineLvl'))
            if ol is not None:
                ol_val = int(ol.get(qn('w:val'), '0'))
                if ol_val <= 8:  # 0-8 是标题，9 是正文
                    level = ol_val + 1

    # 方式三: 中文数字样式 "1级"/"2级"/"2级标题" 等
    if level is None:
        m = re.match(r'^(\d+)级', style_name.strip())
        if m:
            level = int(m.group(1))

    # 方式四: 中文大写样式 "一级"/"二级"/"一级标题" 等
    if level is None:
        CN_NUM = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                  '六': 6, '七': 7, '八': 8, '九': 9}
        m = re.match(r'^([一二三四五六七八九])级', style_name.strip())
        if m:
            level = CN_NUM.get(m.group(1))

    # 方式五: "标题 N" 样式 (中文 Word 常见)
    if level is None:
        m = re.match(r'^标题\s*(\d+)$', style_name.strip())
        if m:
            level = int(m.group(1))

    return level


def extract_headings(docx_path):
    """
    从 docx 文件中提取所有标题，返回:
      [
        {"level": 1, "text": "第一章 试验概述", "para_index": 5},
        {"level": 2, "text": "1.1 试验背景", "para_index": 10},
        ...
      ]
    """
    doc = Document(docx_path)
    headings = []
    for i, p in enumerate(doc.paragraphs):
        level = get_heading_level(p)
        if level is not None:
            headings.append({
                "level": level,
                "text": p.text.strip(),
                "para_index": i,
            })
    return headings


def build_toc_tree(headings):
    """
    将平铺的标题列表转为树形结构。

    返回:
      [
        {
          "text": "第一章 试验概述",
          "level": 1,
          "para_index": 5,
          "children": [
            {"text": "1.1 试验背景", "level": 2, "para_index": 10, "children": []},
            ...
          ]
        },
        ...
      ]
    """
    root = []          # 最终树根节点列表
    stack = []         # 路径栈 [(level, node), ...]

    for h in headings:
        node = {"text": h["text"], "level": h["level"],
                "para_index": h["para_index"], "children": [],
                "page": h.get("page")}

        # 弹出栈中所有层级 >= 当前标题的节点
        while stack and stack[-1][0] >= h["level"]:
            stack.pop()

        if not stack:
            # 顶层标题
            root.append(node)
        else:
            # 挂到最近的父节点
            parent = stack[-1][1]
            parent["children"].append(node)

        stack.append((h["level"], node))

    return root


def format_toc_tree(tree, indent=0):
    """将标题树格式化为文本行（含页码）"""
    lines = []
    MARKERS = {1: "■ ", 2: "◆ ", 3: "▸ ", 4: "· ", 5: "- "}

    for node in tree:
        prefix = "    " * indent
        marker = MARKERS.get(node["level"], "- ")
        page_str = f"  [{node['page']}]" if node.get("page") else ""
        lines.append(f"{prefix}{marker}{node['text']}{page_str}")
        if node["children"]:
            lines.extend(format_toc_tree(node["children"], indent + 1))
    return lines


def format_toc_flat(headings):
    """将平铺标题列表格式化为带缩进的文本行（含页码）"""
    lines = []
    for h in headings:
        indent = "    " * (h["level"] - 1)
        prefix = {1: "■ ", 2: "◆ ", 3: "▸ ", 4: "· "}.get(h["level"], "- ")
        page_str = f"  [{h['page']}]" if h.get("page") else ""
        lines.append(f"{indent}{prefix}{h['text']}{page_str}")
    return lines


def toc_to_markdown(tree, indent=0):
    """将标题树格式化为 Markdown 无序列表（含页码）"""
    lines = []
    for node in tree:
        prefix = "    " * indent + "- "
        page_str = f" `[{node['page']}]`" if node.get("page") else ""
        lines.append(f"{prefix}**{node['text']}**{page_str}")
        if node["children"]:
            lines.extend(toc_to_markdown(node["children"], indent + 1))
    return lines


def process_protocol(docx_path, client=None, model: str = ""):
    """处理单个方案文件，返回目录信息。

    TOC 优先：通过 XML 字段码 + 点线分隔符检测自动目录（语言/样式名无关），
    直接提取标题文本 + 页码。层级通过编号壳分组推断（纯代码，格式无关）。
    """
    basename = os.path.basename(docx_path)
    name = basename
    if name.endswith('.docx'):
        name = name[:-5]

    doc = Document(docx_path)
    toc_entries = _extract_toc_entries(doc)  # [(level_or_none, heading_text, page), ...]

    if toc_entries and len(toc_entries) >= 5:
        # 字段码检测到足够多的条目 → 真正的自动目录
        # 层级缺失时用纯代码推断（编号壳分组 + 嵌套关系）
        if any(lv is None for lv, _, _ in toc_entries):
            toc_entries = assign_toc_levels(toc_entries)

        headings = []
        for i, (level, text, page) in enumerate(toc_entries):
            headings.append({
                "level": level or 1,
                "text": text,
                "para_index": i,
                "page": page,
            })
    else:
        # 回退：段落样式检测
        headings = extract_headings(docx_path)
        try:
            fallback_entries = _extract_toc_entries(doc)
            toc_page_map = {text: page for _, text, page in fallback_entries if page}
        except Exception:
            toc_page_map = {}
        if toc_page_map:
            pages_matched = 0
            for h in headings:
                h_text = h["text"]
                page = toc_page_map.get(h_text)
                if page is None:
                    for toc_text, toc_page in toc_page_map.items():
                        if h_text in toc_text or toc_text in h_text:
                            page = toc_page
                            break
                h["page"] = page
                if page:
                    pages_matched += 1

    tree = build_toc_tree(headings)

    # 统计
    level_counts = {}
    for h in headings:
        level_counts[h["level"]] = level_counts.get(h["level"], 0) + 1

    pages_matched = sum(1 for h in headings if h.get("page"))

    return {
        "name": name,
        "source_file": docx_path,
        "total_headings": len(headings),
        "level_counts": level_counts,
        "max_depth": max(level_counts.keys()) if level_counts else 0,
        "headings_flat": headings,
        "toc_tree": tree,
        "has_toc_pages": pages_matched > 0,
        "pages_matched": pages_matched,
    }


def print_single_protocol(result, show_tree=True):
    """终端打印单个方案的章节目录"""
    print(f"\n{'═' * 70}")
    print(f"📋 {result['name']}")
    print(f"{'═' * 70}")
    print(f"标题总数: {result['total_headings']}")
    page_info = f"  页码匹配: {result.get('pages_matched', 0)}/{result['total_headings']}" \
        if result.get('has_toc_pages') else "  页码: 无（文档未使用自动目录）"
    print(f"层级分布: {dict(sorted(result['level_counts'].items()))}{page_info}")
    print(f"最大深度: {result['max_depth']} 级")
    print()

    if show_tree and result["toc_tree"]:
        print("目录树:")
        for line in format_toc_tree(result["toc_tree"]):
            print(f"  {line}")
    elif result["headings_flat"]:
        print("标题列表:")
        for line in format_toc_flat(result["headings_flat"]):
            print(f"  {line}")


def print_all_protocols(results):
    """终端打印所有方案的章节目录"""
    print(f"\n{'█' * 70}")
    print(f"临床试验方案章节目录  · 共 {len(results)} 个方案")
    print(f"{'█' * 70}")

    for r in results:
        print_single_protocol(r, show_tree=True)


def save_results(results, save_dir):
    """保存所有结果到文件夹"""
    os.makedirs(save_dir, exist_ok=True)

    # 1. 每个方案单独一个目录树文本
    for i, r in enumerate(results):
        safe_name = r["name"].replace('/', '_').replace('\\', '_')
        safe_name = safe_name[:60]  # 限制长度

        # 文本版目录树
        txt_path = os.path.join(save_dir, f"{i+1:02d}_{safe_name}.txt")
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(f"{r['name']}\n")
            f.write(f"{'─' * 60}\n")
            f.write(f"标题总数: {r['total_headings']}  ·  ")
            f.write(f"层级: {dict(sorted(r['level_counts'].items()))}")
            if r.get('has_toc_pages'):
                f.write(f"  ·  页码匹配: {r.get('pages_matched', 0)}/{r['total_headings']}")
            f.write("\n\n")
            for line in format_toc_tree(r["toc_tree"]):
                f.write(line + "\n")

        # Markdown 版
        md_path = os.path.join(save_dir, f"{i+1:02d}_{safe_name}.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# {r['name']}\n\n")
            f.write(f"**标题总数:** {r['total_headings']}  ")
            f.write(f"**最大深度:** {r['max_depth']} 级  ")
            f.write(f"**层级分布:** {dict(sorted(r['level_counts'].items()))}")
            if r.get('has_toc_pages'):
                f.write(f"  \n**页码匹配:** {r.get('pages_matched', 0)}/{r['total_headings']}")
            f.write("\n\n")
            f.write("## 章节目录\n\n")
            for line in toc_to_markdown(r["toc_tree"]):
                f.write(line + "\n")

    # 2. 汇总 JSON
    json_data = []
    for r in results:
        json_data.append({
            "name": r["name"],
            "source_file": r["source_file"],
            "total_headings": r["total_headings"],
            "level_counts": {str(k): v for k, v in r["level_counts"].items()},
            "max_depth": r["max_depth"],
            "has_toc_pages": r.get("has_toc_pages", False),
            "pages_matched": r.get("pages_matched", 0),
            "headings": [
                {"level": h["level"], "text": h["text"],
                 "page": h.get("page")}
                for h in r["headings_flat"]
            ],
        })
    json_path = os.path.join(save_dir, "所有方案目录汇总.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)

    # 3. 汇总总目（所有方案的一级标题对比）
    overview_path = os.path.join(save_dir, "方案目录总览.txt")
    with open(overview_path, 'w', encoding='utf-8') as f:
        f.write("方案章节目录总览  ·  一级标题对比\n")
        f.write("=" * 60 + "\n\n")
        for r in results:
            f.write(f"【{r['name']}】\n")
            f.write(f"  标题总数: {r['total_headings']}  ·  "
                    f"深度: {r['max_depth']} 级\n")
            for line in format_toc_tree(r["toc_tree"]):
                f.write(f"  {line}\n")
            f.write("\n")

    print(f"结果已保存到: {save_dir}/", file=sys.stderr)
    print(f"  {len(results):02d}_方案名.txt ×{len(results)}   - 每个方案的目录树", file=sys.stderr)
    print(f"  {len(results):02d}_方案名.md ×{len(results)}   - 每个方案的 Markdown 目录", file=sys.stderr)
    print(f"  所有方案目录汇总.json        - 结构化 JSON", file=sys.stderr)
    print(f"  方案目录总览.txt             - 全部方案总览", file=sys.stderr)


def find_docx_files(root_dir, tag="方案"):
    """在项目根目录下找方案源文件目录，返回 docx 文件路径列表"""
    # 先精确找名为 tag 的目录
    for dirpath, dirnames, filenames in os.walk(root_dir):
        dirnames[:] = [d for d in dirnames if not d.startswith('.') and d != '__pycache__']
        if os.path.basename(dirpath) == tag:
            docx_files = sorted([
                os.path.join(dirpath, f) for f in filenames
                if f.endswith('.docx') and not f.startswith('~$')
            ])
            if docx_files:
                return docx_files

    # 宽松匹配
    for dirpath, dirnames, filenames in os.walk(root_dir):
        dirnames[:] = [d for d in dirnames if not d.startswith('.') and d != '__pycache__']
        if tag in os.path.basename(dirpath):
            docx_files = sorted([
                os.path.join(dirpath, f) for f in filenames
                if f.endswith('.docx') and not f.startswith('~$')
            ])
            if docx_files:
                return docx_files

    # 回退：直接在 root_dir 下找 docx 文件（适配 web 上传等场景）
    docx_files = sorted([
        os.path.join(root_dir, f) for f in os.listdir(root_dir)
        if f.endswith('.docx') and not f.startswith('~$')
    ])
    if docx_files:
        return docx_files

    return []


def main():
    parser = argparse.ArgumentParser(
        description="提取临床试验方案 docx 的章节目录",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python3 extract_toc.py                          扫描方案目录下所有 docx
  python3 extract_toc.py --dir ../my_project      换项目根目录
  python3 extract_toc.py --file 方案/xxx.docx      只提取单个方案
  python3 extract_toc.py --json                   JSON 输出
  python3 extract_toc.py --save output/toc        保存到文件夹
  python3 extract_toc.py --tag 方案                自定义源目录名
        """,
    )
    parser.add_argument('--dir', type=str, default=None,
                        help='项目根目录 (默认: 当前目录)')
    parser.add_argument('--file', type=str, default=None,
                        help='只处理单个 docx 文件')
    parser.add_argument('--json', action='store_true',
                        help='JSON 格式输出')
    parser.add_argument('--save', type=str, default=None,
                        help='保存结果到指定文件夹')
    parser.add_argument('--tag', type=str, default='方案',
                        help='方案源文件目录名 (默认: 方案)')

    args = parser.parse_args()
    root = os.path.abspath(args.dir) if args.dir else os.getcwd()

    # 确定要处理的文件
    if args.file:
        docx_paths = [args.file]
    else:
        docx_paths = find_docx_files(root, tag=args.tag)

    if not docx_paths:
        print("未找到任何方案 docx 文件。", file=sys.stderr)
        sys.exit(1)

    print(f"🔍 扫描: {os.path.dirname(docx_paths[0])}", file=sys.stderr)
    print(f"   共 {len(docx_paths)} 个方案文件\n", file=sys.stderr)

    # 提取
    results = []
    for i, path in enumerate(docx_paths):
        name = os.path.basename(path)
        print(f"  [{i+1}/{len(docx_paths)}] {name[:55]}", file=sys.stderr)
        try:
            r = process_protocol(path)
            results.append(r)
        except Exception as e:
            print(f"    ❌ 失败: {e}", file=sys.stderr)

    if not results:
        print("所有方案处理失败。", file=sys.stderr)
        return

    # JSON 输出
    if args.json:
        json_data = []
        for r in results:
            json_data.append({
                "name": r["name"],
                "source_file": r["source_file"],
                "total_headings": r["total_headings"],
                "level_counts": {str(k): v for k, v in r["level_counts"].items()},
                "max_depth": r["max_depth"],
                "has_toc_pages": r.get("has_toc_pages", False),
                "pages_matched": r.get("pages_matched", 0),
                "headings": [
                    {"level": h["level"], "text": h["text"],
                     "page": h.get("page")}
                    for h in r["headings_flat"]
                ],
            })
        print(json.dumps(json_data, ensure_ascii=False, indent=2))
        return

    # 终端打印
    print_all_protocols(results)

    # 保存
    if args.save:
        save_results(results, args.save)


if __name__ == '__main__':
    main()
