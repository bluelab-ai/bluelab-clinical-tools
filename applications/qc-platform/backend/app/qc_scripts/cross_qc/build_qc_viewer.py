#!/usr/bin/env python3
"""
Build an interactive QC viewer HTML page (v2).
Extracts tables from DOCX, parses QC markdown, maps them via mapping JSON,
reads Excel original tables, and generates a self-contained HTML file with
left-right split panel layout.

Usage:
    python3 build_qc_viewer_v2.py \
        --docx 表格附件.docx \
        --md QC报告-汇总.md \
        --mapping 表格-清单-映射表.json \
        --excel-dir 表格/ \
        --output qc-viewer.html

Key changes from v1:
  - Left panel shows ALL tables from DOCX (not just mapped ones)
  - QC status driven by mapping JSON (是否QC=="是" + 关键字匹配/人工指定)
  - "原表格" view uses Excel files instead of DOCX-extracted HTML
  - TOC from QC report drives Pair → conclusion matching
  - --mapping is required (no hardcoded fallback)
"""

import argparse
import json
import re
import html as html_mod
from pathlib import Path
from docx import Document
import markdown

BASE_DIR = Path(__file__).parent

# ── Emoji-to-class mapping (for parsing QC report TOC lines) ──────────────
_EMOJI_TO_CLASS = {
    "🟠": "major",
    "🟡": "minor",
    "🔵": "suggestion",
    "🟣": "pending",
    "✅": "ok",
    "⬜": "none",
}
_CLASS_LABEL = {
    "major": "🟠 Major",
    "minor": "🟡 Minor",
    "suggestion": "🔵 Suggestion",
    "pending": "🟣 待人工",
    "ok": "✅ 无问题",
    "none": "⬜ 未被质控",
}

# Runtime globals set by main()
TABLE_IDX_TO_PAIR: dict[int, int] = {}
PAIR_CONCLUSIONS: dict[int, tuple[str, str]] = {}
EXCEL_HTML_CACHE: dict[int, str] = {}


# ═══════════════════════════════════════════════════════════════════════════
# NEW: Excel → HTML conversion
# ═══════════════════════════════════════════════════════════════════════════

def _excel_col_letter(n: int) -> str:
    """Convert 1-based column index to Excel column letter(s)."""
    result = ""
    while n > 0:
        n, rem = divmod(n - 1, 26)
        result = chr(65 + rem) + result
    return result


def extract_excel_table_html(excel_path: Path, sheet_index: int = 0) -> str:
    """Convert an Excel (.xlsx) sheet to an HTML <table> string.

    Handles merged cells via rowspan/colspan and preserves basic formatting.
    """
    import openpyxl

    wb = openpyxl.load_workbook(str(excel_path), data_only=True)
    if sheet_index >= len(wb.sheetnames):
        wb.close()
        return '<p class="excel-error">Sheet 不存在</p>'

    ws = wb[wb.sheetnames[sheet_index]]

    # Build merged-cell map: (row, col) → (rowspan, colspan, is_hidden)
    # is_hidden = True for cells swallowed by a merge (skip rendering)
    merge_map: dict[tuple[int, int], tuple[int, int, bool]] = {}
    for merge_range in ws.merged_cells.ranges:
        rmin = merge_range.min_row
        rmax = merge_range.max_row
        cmin = merge_range.min_col
        cmax = merge_range.max_col
        rowspan = rmax - rmin + 1
        colspan = cmax - cmin + 1
        # Top-left cell of the merge
        merge_map[(rmin, cmin)] = (rowspan, colspan, False)
        # All other cells in the merge range are hidden
        for r in range(rmin, rmax + 1):
            for c in range(cmin, cmax + 1):
                if (r, c) == (rmin, cmin):
                    continue
                merge_map[(r, c)] = (1, 1, True)

    # Determine max data dimensions
    max_row = ws.max_row or 1
    max_col = ws.max_column or 1

    parts = ['<table class="excel-table">']

    for ri in range(1, max_row + 1):
        row_cells = []
        for ci in range(1, max_col + 1):
            key = (ri, ci)

            # Check if this cell is swallowed by a merge
            if key in merge_map:
                rs, cs, hidden = merge_map[key]
                if hidden:
                    continue  # skip – rendered by the top-left merge cell
            else:
                rs, cs = 1, 1  # no merge

            cell = ws.cell(row=ri, column=ci)
            cell_text = str(cell.value) if cell.value is not None else ""
            cell_text = cell_text.strip().replace("\n", "<br>")
            if not cell_text:
                cell_text = "&nbsp;"

            tag = "th" if ri == 1 else "td"

            attrs = []
            if rs > 1:
                attrs.append(f'rowspan="{rs}"')
            if cs > 1:
                attrs.append(f'colspan="{cs}"')

            attr_str = (" " + " ".join(attrs)) if attrs else ""
            row_cells.append(f"<{tag}{attr_str}>{cell_text}</{tag}>")

        if row_cells:
            parts.append("<tr>" + "".join(row_cells) + "</tr>")

    parts.append("</table>")
    wb.close()
    return "\n".join(parts)


def match_excel_to_tables(
    excel_dir: Path, tables: list[dict]
) -> dict[int, str]:
    """Match Excel files to DOCX tables by table number prefix.

    Excel filenames are like: 01-表 11.1.1.1 各中心...xlsx
    The numeric prefix (01) equals the 表格编号 (1-based) from the mapping JSON.

    Returns: {table_docx_index: excel_html_string}
    """
    result: dict[int, str] = {}

    # Build lookup: 表格编号 (1-based) → excel file path
    num_to_path: dict[int, Path] = {}
    for fpath in sorted(excel_dir.glob("*.xlsx")):
        try:
            num = int(fpath.name.split("-")[0])
            num_to_path[num] = fpath
        except (ValueError, IndexError):
            continue

    print(f"   📂 找到 {len(num_to_path)} 个 Excel 文件（表格/）")

    for t in tables:
        # DOCX table index is 0-based, 表格编号 is 1-based
        table_num = t["index"] + 1
        if table_num in num_to_path:
            try:
                result[t["index"]] = extract_excel_table_html(
                    num_to_path[table_num]
                )
            except Exception as e:
                result[t["index"]] = (
                    f'<p class="excel-error">Excel 读取失败: {html_mod.escape(str(e))}</p>'
                )
        else:
            # Fallback: use DOCX-extracted table HTML
            result[t["index"]] = t.get("html", "<p>无表格数据</p>")

    return result


# ═══════════════════════════════════════════════════════════════════════════
# DOCX extraction (reused from v1)
# ═══════════════════════════════════════════════════════════════════════════

def extract_table_html(table) -> str:
    """Convert a python-docx Table to an HTML <table> string."""
    parts = ['<table class="docx-table">']
    for ri, row in enumerate(table.rows):
        tag = "th" if ri == 0 else "td"
        parts.append("<tr>")
        for cell in row.cells:
            cell_text = cell.text.strip().replace("\n", "<br>")
            if not cell_text:
                cell_text = "&nbsp;"
            parts.append(f"<{tag}>{cell_text}</{tag}>")
        parts.append("</tr>")
    parts.append("</table>")
    return "\n".join(parts)


def extract_tables_from_docx(docx_path: Path) -> list[dict]:
    """Extract all tables from the DOCX with their titles.

    Iterates the document body in XML order, interleaving paragraphs and
    tables.  Table titles are matched from the nearest preceding paragraph
    whose text starts with a table-number pattern (e.g. 表 11.1.1.8).
    """
    doc = Document(str(docx_path))

    # Build O(1) lookup: XML element id → Paragraph object (avoids O(n²) in v1)
    para_by_eid = {id(p._element): p for p in doc.paragraphs}

    # First pass: collect interleaved paragraphs and tables in document order
    elements: list[tuple[str, object]] = []  # ("para", text) | ("table", index)
    table_idx = 0
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            p = para_by_eid.get(id(element))
            if p:
                text = p.text.strip()
                if text:
                    elements.append(("para", text))
        elif tag == "tbl":
            if table_idx < len(doc.tables):
                elements.append(("table", table_idx))
                table_idx += 1

    # Second pass: match each table with the nearest preceding title paragraph
    _table_title_re = re.compile(r"表\s+[\d.]+\s+(.+)")
    _table_num_re = re.compile(r"(表\s+[\d.]+)")

    tables = []
    for i, (etype, idx) in enumerate(elements):
        if etype != "table":
            continue
        table = doc.tables[idx]

        # Search backwards for the last table-caption paragraph before this table
        title = None
        table_number = f"表 {idx + 1}"
        for j in range(i - 1, max(0, i - 6) - 1, -1):
            if elements[j][0] == "para":
                t = elements[j][1]
                m = _table_title_re.match(t)
                if m:
                    title = t
                    tn = _table_num_re.match(t)
                    if tn:
                        table_number = tn.group(1)
                    break

        if title is None:
            # Fallback: use first-row first-cell as hint
            first_cell = (
                table.rows[0].cells[0].text.strip()[:50] if table.rows else ""
            )
            title = f"{table_number} ({first_cell}...)"
        else:
            short_match = _table_title_re.match(title)
            short_name = short_match.group(1) if short_match else title

        tables.append({
            "index": idx,
            "table_number": table_number,
            "title": title,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "html": extract_table_html(table),
        })

    return tables


# ═══════════════════════════════════════════════════════════════════════════
# QC Markdown parsing (reused from v1)
# ═══════════════════════════════════════════════════════════════════════════

def parse_conclusions_from_md(md_text: str) -> dict[int, tuple[str, str]]:
    """从 QC 报告的 TOC 行自动提取 Pair 结论。

    TOC 行格式:  - [Pair 1: 方案偏离情况](#pair-1)  ✅ 无问题
    """
    conclusions: dict[int, tuple[str, str]] = {}
    toc_pattern = re.compile(
        r"- \[Pair (\d+):[^\]]+\]\(#[^)]+\)\s+"
        r"([\U0001F000-\U0001FFFF✅⬜🟠🟡🔴🔵])\s+(\w+)"
    )
    for m in toc_pattern.finditer(md_text):
        pair_num = int(m.group(1))
        emoji = m.group(2)
        cls = _EMOJI_TO_CLASS.get(emoji, "none")
        label = _CLASS_LABEL.get(cls, "⬜ 未被质控")
        conclusions[pair_num] = (cls, label)
    return conclusions


def parse_qc_markdown(md_path: Path) -> dict:
    """Parse the QC markdown into sections: overview + per-pair content."""
    text = md_path.read_text(encoding="utf-8")

    pair_pattern = re.compile(
        r'## <a id="pair-(\d+)"></a>(Pair \d+:.*?)\n\n(.*?)'
        r"(?=\n## <a id=\"pair-|\Z)",
        re.DOTALL,
    )

    pairs = {}
    for m in pair_pattern.finditer(text):
        pair_num = int(m.group(1))
        pair_title = m.group(2).strip()
        pair_content = m.group(3).strip()
        pairs[pair_num] = {
            "title": pair_title,
            "content_md": pair_content,
            "content_html": markdown.markdown(
                pair_content, extensions=["tables", "fenced_code"]
            ),
        }

    # Extract overview section (everything before the first pair)
    first_pair_match = re.search(r'\n## <a id="pair-1"></a>', text)
    if first_pair_match:
        overview_md = text[: first_pair_match.start()].strip()
    else:
        overview_md = text

    overview_html = markdown.markdown(
        overview_md, extensions=["tables", "fenced_code"]
    )

    date_match = re.search(r"\*\s*\*核查日期\*\s*\*:\s*(.+)", overview_md)
    qc_date = date_match.group(1).strip() if date_match else "—"

    return {
        "overview_html": overview_html,
        "pairs": pairs,
        "date": qc_date,
    }


# ═══════════════════════════════════════════════════════════════════════════
# Mapping JSON → Pair ↔ Table (reused from v1)
# ═══════════════════════════════════════════════════════════════════════════

def parse_mapping_from_json(mapping_path: str) -> list[dict]:
    """加载映射表 JSON，返回 entries 列表。"""
    with open(mapping_path, encoding="utf-8") as f:
        return json.load(f)


def build_pair_table_mapping(
    mapping_entries: list[dict],
    tables: list[dict],
) -> tuple[dict[int, int], dict[int, tuple[str, str]], dict[int, int]]:
    """从映射表 JSON 构建 Pair ↔ DOCX 表格索引映射。

    Pair 编号: 严格按映射表 JSON 顺序，对 是否QC=="是" 且
               匹配方法 in (关键字匹配, 人工指定) 的条目依次分配 1, 2, 3, ...
              （与 tfl_qc_workflow._build_qc_pairs 完全一致）
    """
    pair_to_idx: dict[int, int] = {}
    table_idx_to_pair: dict[int, int] = {}
    pair_conclusions: dict[int, tuple[str, str]] = {}
    pid = 0

    for entry in mapping_entries:
        qc_flag = entry.get("是否QC", "是")
        match_method = entry.get("匹配方法", "")
        if not (qc_flag == "是" and match_method in ("关键字匹配", "人工指定")):
            continue

        table_num = entry.get("表格编号", 0)
        if not isinstance(table_num, int) or table_num <= 0:
            continue

        pid += 1
        docx_idx = table_num - 1  # 表格编号 1-based → docx 索引 0-based
        if 0 <= docx_idx < len(tables):
            pair_to_idx[pid] = docx_idx
            table_idx_to_pair[docx_idx] = pid
            pair_conclusions[pid] = ("none", "⬜ 未被质控")

    print(f"   📏 QC Pair: {pid}  →  全部映射: {len(pair_to_idx)}")

    return pair_to_idx, pair_conclusions, table_idx_to_pair


# ═══════════════════════════════════════════════════════════════════════════
# HTML generation
# ═══════════════════════════════════════════════════════════════════════════

def build_html(
    tables: list[dict],
    qc_data: dict,
    excel_html_map: dict[int, str],
) -> str:
    """Generate the complete self-contained HTML page."""

    # ── Logo base64 ──
    logo_b64 = ""
    try:
        from PIL import Image
        import base64 as b64_mod
        import io
        logo_path = BASE_DIR.parent.parent.parent.parent / "frontend" / "public" / "logo.png"
        if logo_path.exists():
            logo = Image.open(str(logo_path))
            h = 120
            w = int(logo.width * h / logo.height)
            logo_small = logo.resize((w, h), Image.LANCZOS)
            buf = io.BytesIO()
            logo_small.save(buf, format="PNG", optimize=True)
            logo_b64 = (
                "data:image/png;base64,"
                + b64_mod.b64encode(buf.getvalue()).decode()
            )
    except Exception:
        pass

    # ── Build table list data for left panel ──
    table_list = []
    for t in tables:
        idx = t["index"]
        pair_num = TABLE_IDX_TO_PAIR.get(idx)
        if pair_num:
            conclusion_class, conclusion_label = PAIR_CONCLUSIONS.get(
                pair_num, ("ok", "✅ 无问题")
            )
            qc_status = "qc"
        else:
            pair_num = None
            conclusion_class = "none"
            conclusion_label = "⬜ 未被质控"
            qc_status = "no-qc"

        # Extract short name from title
        short_match = re.match(r"表\s+[\d.]+\s+(.+)", t["title"])
        short_name = short_match.group(1) if short_match else t["title"]

        table_list.append({
            "index": idx,
            "tableNumber": t["table_number"],
            "title": t["title"],
            "shortName": short_name,
            "rows": t["rows"],
            "cols": t["cols"],
            "pairNum": pair_num,
            "qcStatus": qc_status,
            "conclusionClass": conclusion_class,
            "conclusionLabel": conclusion_label,
        })

    # ── Build pair data ──
    pairs_data = {}
    for pair_num, pair_info in qc_data["pairs"].items():
        conclusion_class, conclusion_label = PAIR_CONCLUSIONS.get(
            pair_num, ("ok", "✅ 无问题")
        )
        pairs_data[str(pair_num)] = {
            "title": pair_info["title"],
            "html": pair_info["content_html"],
            "conclusionClass": conclusion_class,
            "conclusionLabel": conclusion_label,
        }

    # ── QC pair HTML blocks (hidden) ──
    pair_html_blocks = []
    for pair_num in sorted(qc_data["pairs"].keys()):
        info = qc_data["pairs"][pair_num]
        conclusion_class, conclusion_label = PAIR_CONCLUSIONS.get(
            pair_num, ("ok", "✅ 无问题")
        )
        short_title = info["title"].replace(f"Pair {pair_num}: ", "")
        pair_html_blocks.append(
            f'<div class="qc-pair-block" id="qc-pair-{pair_num}" data-pair="{pair_num}">\n'
            f'  <h2 class="qc-pair-title">Pair {pair_num}: {html_mod.escape(short_title)}</h2>\n'
            f'  <div class="qc-conclusion-badge badge-{conclusion_class}">{conclusion_label}</div>\n'
            f'  {info["content_html"]}\n'
            f"</div>"
        )

    # ── Excel table HTML blocks (hidden) ──
    excel_html_blocks = []
    for t in tables:
        idx = t["index"]
        excel_html = excel_html_map.get(idx, t.get("html", "<p>无表格数据</p>"))
        pair_num = TABLE_IDX_TO_PAIR.get(idx)
        pair_attr = f' data-pair="{pair_num}"' if pair_num else ""
        excel_html_blocks.append(
            f'<div class="excel-block" id="excel-block-{idx}" data-table-index="{idx}"{pair_attr}>\n'
            f'  <div class="table-block-title">{html_mod.escape(t["title"])}</div>\n'
            f'  {excel_html}\n'
            f"</div>"
        )

    # ── Statistics ──
    conclusion_counts: dict[str, int] = {}
    for _pair_num, (cls, _label) in PAIR_CONCLUSIONS.items():
        conclusion_counts[cls] = conclusion_counts.get(cls, 0) + 1

    qc_count = len(PAIR_CONCLUSIONS)
    noqc_count = len(tables) - qc_count
    total_tables = len(tables)

    # ── JSON payloads ──
    table_list_json = json.dumps(table_list, ensure_ascii=False)
    pairs_data_json = json.dumps(pairs_data, ensure_ascii=False)
    table_idx_to_pair_json = json.dumps(TABLE_IDX_TO_PAIR)
    pair_conclusions_json = json.dumps(
        {str(k): list(v) for k, v in PAIR_CONCLUSIONS.items()}
    )

    # ── Stats HTML ──
    stats_parts = []
    stats_parts.append(
        f'<span class="stat-total">总表格: {total_tables}</span>'
    )
    if qc_count > 0:
        stats_parts.append(
            f'<span class="stat-qc">已质控: {qc_count}</span>'
        )
    if noqc_count > 0:
        stats_parts.append(
            f'<span class="stat-noqc">未质控: {noqc_count}</span>'
        )
    for cls, label in [
        ("major", "Major"),
        ("minor", "Minor"),
        ("suggestion", "Suggestion"),
        ("pending", "待人工"),
        ("ok", "无问题"),
    ]:
        count = conclusion_counts.get(cls, 0)
        if count > 0:
            stats_parts.append(
                f'<span class="stat-{cls}">{label}: {count}</span>'
            )

    stats_html = "\n    ".join(stats_parts)

    # ── CSS (reused from v1 with additions for excel-table) ──
    css = """
/* ── Reset & Base ── */
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
html, body { height: 100%; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; }
body { display: flex; flex-direction: column; overflow: hidden; background: #f8fafc; }

/* ── Header ── */
.top-bar {
  display: flex; align-items: center; gap: 16px;
  padding: 16px 24px;
  background: #fff; border-bottom: 1px solid rgba(226,232,240,0.7);
  flex-shrink: 0;
}
.top-bar .logo { display: flex; align-items: center; gap: 10px; }
.top-bar .logo img { height: 28px; width: auto; vertical-align: middle; }
.top-bar h1 { font-size: 16px; font-weight: 700; color: #0f172a; }
.top-bar .subtitle { font-size: 12px; color: #64748b; margin-left: 12px; }
.top-bar .stats { display: flex; gap: 10px; margin-left: auto; font-size: 12px; flex-wrap: wrap; }
.top-bar .stats span { padding: 4px 12px; border-radius: 9999px; font-weight: 600; white-space: nowrap; }
.stat-total  { background: #f1f5f9; color: #475569; }
.stat-qc    { background: #dbeafe; color: #1e40af; }
.stat-noqc  { background: #f1f5f9; color: #64748b; }
.stat-pending   { background: #f3e8ff; color: #6b21a8; }
.stat-major      { background: #ffedd5; color: #9a3412; }
.stat-minor      { background: #fef9c3; color: #854d0e; }
.stat-suggestion { background: #dbeafe; color: #1e40af; }
.stat-ok         { background: #dcfce7; color: #166534; }

/* ── Main Layout ── */
.main-container { display: flex; flex: 1; overflow: hidden; }
@media (max-width: 900px) {
  .main-container { flex-direction: column; }
  .left-panel { max-height: 40vh !important; }
}

/* ── Left Panel ── */
.left-panel {
  width: 420px; min-width: 360px; flex-shrink: 0;
  display: flex; flex-direction: column;
  border-right: 1px solid rgba(226,232,240,0.7);
  background: #fff;
}
.left-panel .search-box { padding: 16px; background: #fff; }
.left-panel .search-box input {
  width: 100%; padding: 10px 14px;
  border: 1px solid #e2e8f0; border-radius: 12px;
  font-size: 14px; outline: none; background: #f8fafc;
  transition: border-color 0.2s, box-shadow 0.2s;
}
.left-panel .search-box input:focus {
  border-color: #2563eb; box-shadow: 0 0 0 3px rgba(37,99,235,0.1);
}

.table-list { flex: 1; overflow-y: auto; padding: 0 8px 8px; }
.table-list-item {
  display: flex; align-items: center; gap: 10px;
  padding: 10px 14px; margin: 2px 0;
  cursor: pointer; border-radius: 12px;
  transition: all 0.15s;
}
.table-list-item:hover { background: #eff6ff; }
.table-list-item.active { background: #eff6ff; border: 1px solid #bfdbfe; }
.table-list-item.qc-highlight { background: #fefce8; border: 1px solid #fde68a; }

.table-list-item .qc-badge {
  font-size: 11px; padding: 2px 8px; border-radius: 9999px;
  white-space: nowrap; flex-shrink: 0; font-weight: 600;
}
.badge-major     { background: #ffedd5; color: #9a3412; }
.badge-minor     { background: #fef9c3; color: #854d0e; }
.badge-suggestion{ background: #dbeafe; color: #1e40af; }
.badge-pending   { background: #f3e8ff; color: #6b21a8; }
.badge-ok        { background: #dcfce7; color: #166534; }
.badge-none      { background: #f1f5f9; color: #94a3b8; }

.table-list-item .table-idx { font-size: 12px; color: #94a3b8; min-width: 28px; font-weight: 600; }
.table-list-item .table-name {
  flex: 1; font-size: 13px; line-height: 1.4; color: #334155;
  overflow: hidden; text-overflow: ellipsis;
  display: -webkit-box; -webkit-line-clamp: 2; -webkit-box-orient: vertical;
}

/* ── Right Panel ── */
.right-panel { flex: 1; display: flex; flex-direction: column; overflow: hidden; }
.right-panel .panel-header {
  padding: 14px 24px; border-bottom: 1px solid rgba(226,232,240,0.7);
  background: #fff; display: flex; align-items: center; gap: 12px;
  flex-shrink: 0;
}
.right-panel .panel-header h2 { font-size: 15px; font-weight: 700; color: #1e293b; flex: 1; }
.right-panel .panel-header .view-toggle { display: flex; gap: 4px; }
.right-panel .panel-header .view-toggle button {
  padding: 8px 16px; border: 1px solid #e2e8f0; border-radius: 12px;
  background: #fff; cursor: pointer; font-size: 13px; font-weight: 500; color: #64748b;
  transition: all 0.15s;
}
.right-panel .panel-header .view-toggle button.active {
  background: #2563eb; color: #fff; border-color: #2563eb;
}
.right-panel .panel-header .view-toggle button:hover:not(.active) { background: #f1f5f9; }

.content-area { flex: 1; overflow-y: auto; padding: 24px; background: #f8fafc; }
.content-area .empty-state {
  display: flex; flex-direction: column; align-items: center;
  justify-content: center; height: 100%; color: #94a3b8;
}
.content-area .empty-state .icon { font-size: 64px; margin-bottom: 16px; opacity: 0.4; }
.content-area .empty-state p { font-size: 15px; }

.content-area .no-qc-notice {
  display: flex; flex-direction: column; align-items: center;
  justify-content: center; padding: 60px 20px; text-align: center;
  background: #fff; border-radius: 16px; border: 1px solid rgba(226,232,240,0.7);
}
.content-area .no-qc-notice .icon { font-size: 56px; margin-bottom: 16px; opacity: 0.5; }
.content-area .no-qc-notice h3 { font-size: 17px; color: #64748b; margin-bottom: 6px; font-weight: 700; }
.content-area .no-qc-notice p { font-size: 13px; color: #94a3b8; }

/* ── QC Content Styling ── */
.qc-pair-block {
  scroll-margin-top: 20px; padding: 24px;
  background: #fff; border-radius: 16px;
  border: 1px solid rgba(226,232,240,0.7);
  margin-bottom: 20px;
  transition: box-shadow 0.3s, border-color 0.3s;
}
.qc-pair-block.highlight {
  box-shadow: 0 0 0 3px #2563eb, 0 4px 20px rgba(37,99,235,0.15);
  border-color: #2563eb;
}
.qc-pair-block .qc-pair-title { font-size: 17px; margin-bottom: 8px; color: #0f172a; font-weight: 700; }
.qc-pair-block .qc-conclusion-badge {
  display: inline-block; padding: 4px 14px; border-radius: 9999px;
  font-size: 13px; font-weight: 600; margin-bottom: 20px;
}

.qc-pair-block table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13px; }
.qc-pair-block table th, .qc-pair-block table td {
  border: 1px solid #e2e8f0; padding: 8px 12px; text-align: left; vertical-align: top;
}
.qc-pair-block table th { background: #f8fafc; font-weight: 600; color: #475569; }
.qc-pair-block table tr:nth-child(even) td { background: #f8fafc; }
.qc-pair-block h1, .qc-pair-block h2, .qc-pair-block h3,
.qc-pair-block h4, .qc-pair-block h5, .qc-pair-block h6 { margin: 16px 0 8px; color: #1e293b; }
.qc-pair-block h3 { font-size: 15px; font-weight: 700; }
.qc-pair-block h4 { font-size: 14px; font-weight: 600; }
.qc-pair-block p { margin: 8px 0; line-height: 1.7; font-size: 14px; color: #334155; }
.qc-pair-block blockquote {
  border-left: 3px solid #2563eb; padding: 10px 16px;
  margin: 12px 0; background: #eff6ff; font-size: 13px; border-radius: 0 8px 8px 0;
}
.qc-pair-block code { background: #f1f5f9; padding: 2px 6px; border-radius: 4px; font-size: 13px; color: #1e293b; }
.qc-pair-block pre { background: #f8fafc; padding: 14px; border-radius: 12px; overflow-x: auto; font-size: 13px; margin: 12px 0; }
.qc-pair-block hr { border: none; border-top: 1px solid #e2e8f0; margin: 20px 0; }
.qc-pair-block ul, .qc-pair-block ol { margin: 8px 0; padding-left: 24px; }
.qc-pair-block li { line-height: 1.7; font-size: 14px; color: #334155; }
.qc-pair-block a { color: #2563eb; text-decoration: none; }
.qc-pair-block a:hover { text-decoration: underline; }

/* ── Excel Table View ── */
.excel-table { border-collapse: collapse; width: 100%; font-size: 12px; margin: 12px 0; }
.excel-table th, .excel-table td {
  border: 1px solid #e2e8f0; padding: 6px 10px; text-align: left; vertical-align: top;
}
.excel-table th { background: #dbeafe; font-weight: 600; color: #1e40af; }
.excel-table tr:nth-child(even) td { background: #f8fafc; }

.excel-block { margin-bottom: 24px; }
.excel-block .table-block-title {
  font-size: 14px; font-weight: 600; margin-bottom: 8px; color: #1e293b;
  padding: 10px 14px; background: #eff6ff; border-radius: 12px;
}
.excel-block.highlight-table {
  box-shadow: 0 0 0 3px #2563eb, 0 4px 20px rgba(37,99,235,0.15);
  border-radius: 12px;
}

.excel-error {
  padding: 20px; text-align: center; color: #ef4444;
  background: #fef2f2; border-radius: 12px;
}

/* ── Overview Container ── */
.overview-container { max-width: 860px; }
.overview-container h1 { font-size: 20px; margin-bottom: 12px; color: #0f172a; }
.overview-container h2 {
  font-size: 17px; margin: 24px 0 10px; padding-bottom: 8px; color: #1e293b;
  border-bottom: 1px solid #e2e8f0;
}
.overview-container table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13px; }
.overview-container table th, .overview-container table td {
  border: 1px solid #e2e8f0; padding: 8px 12px; text-align: left; vertical-align: top;
}
.overview-container table th { background: #f8fafc; font-weight: 600; color: #475569; }
.overview-container table tr:nth-child(even) td { background: #f8fafc; }
.overview-container p { margin: 8px 0; line-height: 1.7; font-size: 14px; color: #334155; }
.overview-container ul, .overview-container ol { margin: 8px 0; padding-left: 24px; }
.overview-container li { line-height: 1.8; font-size: 14px; color: #334155; }
.overview-container a { color: #2563eb; text-decoration: none; }
.overview-container a:hover { text-decoration: underline; }
.overview-container hr { border: none; border-top: 1px solid #e2e8f0; margin: 20px 0; }
.overview-container blockquote {
  border-left: 3px solid #2563eb; padding: 10px 16px;
  margin: 12px 0; background: #eff6ff; border-radius: 0 8px 8px 0;
}
.overview-container code { background: #f1f5f9; padding: 2px 6px; border-radius: 4px; font-size: 13px; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #cbd5e1; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #94a3b8; }

/* ═══════════════════════════════════════════════════════════════════════════
   Cover Page
   ═══════════════════════════════════════════════════════════════════════════ */
.cover-overlay {
  position: fixed; inset: 0; z-index: 9999;
  display: flex; align-items: center; justify-content: center;
  background: #fff;
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
  transition: opacity 0.5s, visibility 0.5s;
}
.cover-overlay.dismissed { opacity: 0; visibility: hidden; pointer-events: none; }

.cover-overlay::before {
  content: ''; position: absolute; top: 0; left: 0; right: 0; height: 320px;
  background: linear-gradient(180deg, #eff6ff 0%, rgba(239,246,255,0.3) 70%, transparent 100%);
  pointer-events: none;
}

.cover-card {
  position: relative; z-index: 1;
  text-align: center; max-width: 720px;
  padding: 48px 64px 60px;
}
.cover-logo {
  margin-bottom: 56px;
  display: inline-block;
}
.cover-logo img {
  height: 140px; width: auto;
  filter: drop-shadow(0 6px 16px rgba(37,99,235,0.08));
}
.cover-label {
  font-size: 13px; font-weight: 600; color: #2563eb;
  letter-spacing: 0.22em; text-transform: uppercase; margin-bottom: 24px;
}
.cover-h1 {
  font-size: 38px; font-weight: 800; color: #0f172a;
  line-height: 1.3; margin-bottom: 12px;
  letter-spacing: -0.015em;
}
.cover-h2 {
  font-size: 15px; color: #64748b; font-weight: 400; margin-bottom: 56px;
}

.cover-enter-btn {
  display: inline-flex; align-items: center; gap: 10px;
  padding: 16px 56px;
  background: #2563eb; color: #fff;
  border: none; border-radius: 14px;
  font-size: 16px; font-weight: 700; cursor: pointer;
  letter-spacing: 0.02em;
  transition: all 0.2s;
  box-shadow: 0 4px 20px rgba(37,99,235,0.30);
  font-family: inherit;
}
.cover-enter-btn:hover {
  background: #1d4ed8; transform: translateY(-1px);
  box-shadow: 0 8px 32px rgba(37,99,235,0.40);
}
.cover-enter-btn:active { transform: translateY(0); }

.cover-footer {
  margin-top: 48px; font-size: 12px; color: #94a3b8;
}
.cover-footer span { color: #64748b; font-weight: 500; }
"""

    # ── Build HTML page ──
    html_page = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>表格清单一致性质控报告</title>
<style>
{css}
</style>
</head>
<body>

<!-- ═══════════════════════════════════════════════════════════════════════════
     Cover Page
     ═══════════════════════════════════════════════════════════════════════════ -->
<div class="cover-overlay" id="coverOverlay">
  <div class="cover-card">
    <div class="cover-logo">
      {f'<img src="{logo_b64}" alt="Logo" />' if logo_b64 else '''
      <svg width="88" height="88" viewBox="0 0 88 88" fill="none">
        <rect x="14" y="8" width="60" height="72" rx="10" fill="#eff6ff" stroke="#2563eb" stroke-width="2.5"/>
        <path d="M40 26v20m0 0l-8-8m8 8l8-8" stroke="#2563eb" stroke-width="3" stroke-linecap="round" stroke-linejoin="round"/>
        <circle cx="44" cy="58" r="12" fill="#2563eb" opacity="0.12" stroke="#2563eb" stroke-width="2"/>
        <path d="M39 58l3 3 7-7" stroke="#16a34a" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"/>
      </svg>'''}
    </div>

    <div class="cover-label">Clinical Trial Quality Control Report</div>

    <div class="cover-h1">表格清单一致性质控</div>
    <div class="cover-h2">Table-Listing Cross-Validation Quality Control Report</div>

    <button class="cover-enter-btn" onclick="dismissCover()">
      <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round">
        <polyline points="9 18 15 12 9 6"/>
      </svg>
      查看报告
    </button>

    <div class="cover-footer">
      由 <span>TFL QC Platform</span> 自动生成 · {qc_data.get('date', '—')}
    </div>
  </div>
</div>

<script>
function dismissCover() {{
  document.getElementById('coverOverlay').classList.add('dismissed');
}}
</script>

<!-- Header -->
<header class="top-bar">
  <div class="logo">
    {f'<img src="{logo_b64}" alt="Logo" style="height:28px;width:auto;" />' if logo_b64 else '''
    <svg viewBox="0 0 80 80" fill="none" width="28" height="28">
      <circle cx="40" cy="40" r="36" fill="#2563eb" opacity="0.12"/>
      <path d="M28 22h16l10 10v20a4 4 0 0 1-4 4H28a4 4 0 0 1-4-4V26a4 4 0 0 1 4-4z" fill="none" stroke="#2563eb" stroke-width="3" stroke-linecap="round" stroke-linejoin="round"/>
      <path d="M44 22v10h10" fill="none" stroke="#2563eb" stroke-width="3" stroke-linecap="round" stroke-linejoin="round"/>
      <circle cx="48" cy="52" r="10" fill="#2563eb" opacity="0.15" stroke="#2563eb" stroke-width="2"/>
      <path d="M44 52l3 2.5 5-5" fill="none" stroke="#2563eb" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"/>
    </svg>'''}
    <h1>表格清单一致性质控报告</h1>
  </div>
  <div class="stats">
    {stats_html}
  </div>
</header>

<!-- Main Container -->
<div class="main-container">

  <!-- Left Panel -->
  <div class="left-panel" id="leftPanel">
    <div class="search-box">
      <input type="text" id="searchInput" placeholder="搜索表格编号或名称...">
    </div>
    <div class="table-list" id="tableList"></div>
  </div>

  <!-- Right Panel -->
  <div class="right-panel" id="rightPanel">
    <div class="panel-header">
      <h2 id="panelTitle">质控报告概览</h2>
      <div class="view-toggle">
        <button id="btnQcView" class="active" onclick="switchView('qc')">质控报告</button>
        <button id="btnTableView" onclick="switchView('table')">原表格</button>
      </div>
    </div>
    <div class="content-area" id="contentArea">
      <div class="overview-container" id="overviewContainer">
        <div style="background:#eff6ff;padding:12px 16px;border-radius:12px;margin-bottom:16px;font-size:13px;color:#1e40af;">
          💡 <strong>使用提示：</strong>将鼠标悬停在左侧表格上预览质控结果，点击可固定定位。
          按 <kbd style="background:#dbeafe;padding:1px 6px;border-radius:3px;">Esc</kbd> 返回此概览。
        </div>
        {qc_data['overview_html']}
      </div>
    </div>
  </div>

</div>

<!-- Hidden data stores -->
<div id="qc-pairs-container" style="display:none;">
  {''.join(pair_html_blocks)}
</div>
<div id="excel-tables-container" style="display:none;">
  {''.join(excel_html_blocks)}
</div>

<script>
// ── Data ──
const TABLE_LIST = {table_list_json};
const PAIRS_DATA = {pairs_data_json};
const TABLE_IDX_TO_PAIR = {table_idx_to_pair_json};
const PAIR_CONCLUSIONS = {pair_conclusions_json};

let currentView = 'qc';
let activeTableIdx = null;

// ── DOM refs ──
const tableListEl = document.getElementById('tableList');
const contentArea = document.getElementById('contentArea');
const panelTitle = document.getElementById('panelTitle');
const searchInput = document.getElementById('searchInput');

// ── Build left panel ──
function buildTableList(filter) {{
  filter = (filter || '').toLowerCase();
  tableListEl.innerHTML = '';

  const filtered = TABLE_LIST.filter(function(t) {{
    if (!filter) return true;
    return (t.tableNumber || '').toLowerCase().includes(filter) ||
           (t.shortName || '').toLowerCase().includes(filter) ||
           (t.title || '').toLowerCase().includes(filter) ||
           String(t.index + 1).includes(filter);
  }});

  if (filtered.length === 0) {{
    tableListEl.innerHTML = '<div style="padding:20px;text-align:center;color:#94a3b8;">未找到匹配的表格</div>';
    return;
  }}

  filtered.forEach(function(t) {{
    const div = document.createElement('div');
    div.className = 'table-list-item';
    div.dataset.tableIndex = t.index;
    if (activeTableIdx === t.index) div.classList.add('active');

    const badgeClass = 'badge-' + t.conclusionClass;
    const pairInfo = t.pairNum ? ' &rarr; Pair ' + t.pairNum : '';

    div.innerHTML =
      '<span class="table-idx">#' + (t.index + 1) + '</span>' +
      '<span class="table-name">' + escapeHtml(t.tableNumber) + ' ' +
        escapeHtml(t.shortName) + '</span>' +
      '<span class="qc-badge ' + badgeClass + '">' + t.conclusionLabel + pairInfo + '</span>';

    div.addEventListener('mouseenter', function() {{ onTableHover(t.index); }});
    div.addEventListener('mouseleave', onTableLeave);
    div.addEventListener('click', function() {{ onTableClick(t.index); }});

    tableListEl.appendChild(div);
  }});
}}

// ── Interaction ──
function onTableHover(tableIdx) {{
  document.querySelectorAll('.table-list-item').forEach(function(el) {{
    if (parseInt(el.dataset.tableIndex) === tableIdx) el.classList.add('qc-highlight');
    else el.classList.remove('qc-highlight');
  }});
  if (activeTableIdx === null) updateRightPanel(tableIdx);
}}

function onTableLeave() {{
  document.querySelectorAll('.table-list-item').forEach(function(el) {{ el.classList.remove('qc-highlight'); }});
  if (activeTableIdx !== null) updateRightPanel(activeTableIdx);
  else resetRightPanel();
}}

function onTableClick(tableIdx) {{
  activeTableIdx = tableIdx;
  buildTableList(searchInput.value);
  updateRightPanel(tableIdx);
  setTimeout(function() {{
    if (currentView === 'qc') {{
      const pairNum = TABLE_IDX_TO_PAIR[tableIdx];
      if (pairNum) {{
        const qcBlock = document.getElementById('qc-pair-' + pairNum);
        if (qcBlock) {{
          qcBlock.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
          qcBlock.classList.add('highlight');
          setTimeout(function() {{ qcBlock.classList.remove('highlight'); }}, 2500);
        }}
      }}
    }} else {{
      const excelBlock = document.getElementById('excel-block-' + tableIdx);
      if (excelBlock) {{
        excelBlock.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
        excelBlock.classList.add('highlight-table');
        setTimeout(function() {{ excelBlock.classList.remove('highlight-table'); }}, 2500);
      }}
    }}
  }}, 50);
}}

function updateRightPanel(tableIdx) {{
  const item = TABLE_LIST.find(function(t) {{ return t.index === tableIdx; }});
  if (!item) return;

  panelTitle.textContent = item.tableNumber + ' ' + item.shortName;

  if (currentView === 'qc') {{
    const pairNum = TABLE_IDX_TO_PAIR[tableIdx];
    if (pairNum && PAIRS_DATA[String(pairNum)]) {{
      const pairData = PAIRS_DATA[String(pairNum)];
      const conclusion = PAIR_CONCLUSIONS[String(pairNum)] || ['ok', '✅ 无问题'];
      contentArea.innerHTML =
        '<div class="qc-pair-block" id="qc-pair-' + pairNum + '" data-pair="' + pairNum + '">' +
        '  <h2 class="qc-pair-title">Pair ' + pairNum + ': ' +
             escapeHtml((pairData.title || '').replace('Pair ' + pairNum + ': ', '')) + '</h2>' +
        '  <div class="qc-conclusion-badge badge-' + conclusion[0] + '">' + conclusion[1] + '</div>' +
        '  ' + pairData.html +
        '</div>';
    }} else {{
      contentArea.innerHTML =
        '<div class="no-qc-notice">' +
        '  <div class="icon">📭</div>' +
        '  <h3>该表格未被质控</h3>' +
        '  <p>此表格（' + escapeHtml(item.tableNumber) + ' ' + escapeHtml(item.shortName) + '）未包含在质控报告中</p>' +
        '  <p style="font-size:12px;margin-top:12px;color:#cbd5e1;">映射表中未匹配为关键字匹配或人工指定，或是否QC≠是</p>' +
        '</div>';
    }}
  }} else {{
    // "原表格" view — use Excel HTML
    const excelBlock = document.getElementById('excel-block-' + tableIdx);
    if (excelBlock) {{
      contentArea.innerHTML = excelBlock.outerHTML;
      const block = contentArea.querySelector('.excel-block');
      if (block) block.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
    }} else {{
      contentArea.innerHTML = '<div class="empty-state"><p>表格数据不可用</p></div>';
    }}
  }}
}}

function resetRightPanel() {{
  panelTitle.textContent = '质控报告概览';
  const overview = document.getElementById('overviewContainer');
  if (overview) {{
    contentArea.innerHTML = overview.innerHTML;
  }} else {{
    contentArea.innerHTML =
      '<div class="empty-state">' +
      '  <div class="icon">📊</div>' +
      '  <p>质控报告已加载</p>' +
      '  <p style="font-size:13px;margin-top:8px;">悬停左侧表格查看质控结果</p>' +
      '</div>';
  }}
}}

function switchView(view) {{
  currentView = view;
  document.getElementById('btnQcView').classList.toggle('active', view === 'qc');
  document.getElementById('btnTableView').classList.toggle('active', view === 'table');
  if (activeTableIdx !== null) {{
    updateRightPanel(activeTableIdx);
    setTimeout(function() {{
      if (view === 'qc') {{
        const pairNum = TABLE_IDX_TO_PAIR[activeTableIdx];
        if (pairNum && document.getElementById('qc-pair-' + pairNum))
          document.getElementById('qc-pair-' + pairNum).scrollIntoView({{ behavior: 'smooth', block: 'start' }});
      }} else {{
        const block = contentArea.querySelector('.excel-block');
        if (block) block.scrollIntoView({{ behavior: 'smooth', block: 'start' }});
      }}
    }}, 50);
  }} else {{
    resetRightPanel();
  }}
}}

searchInput.addEventListener('input', function() {{ buildTableList(searchInput.value); }});

document.addEventListener('keydown', function(e) {{
  if (e.key === 'Escape') {{
    activeTableIdx = null;
    buildTableList(searchInput.value);
    resetRightPanel();
  }}
  if (e.key === 'ArrowDown' || e.key === 'ArrowUp') {{
    e.preventDefault();
    const items = Array.from(document.querySelectorAll('.table-list-item'));
    if (items.length === 0) return;
    const currentIdx = items.findIndex(function(el) {{ return el.classList.contains('active'); }});
    let nextIdx;
    if (e.key === 'ArrowDown') nextIdx = currentIdx < 0 ? 0 : Math.min(currentIdx + 1, items.length - 1);
    else nextIdx = currentIdx < 0 ? items.length - 1 : Math.max(currentIdx - 1, 0);
    const nextItem = items[nextIdx];
    const tableIdx = parseInt(nextItem.dataset.tableIndex);
    onTableClick(tableIdx);
    nextItem.scrollIntoView({{ block: 'nearest' }});
  }}
}});

function escapeHtml(str) {{
  const div = document.createElement('div');
  div.textContent = String(str);
  return div.innerHTML;
}}

buildTableList('');
</script>
</body>
</html>"""

    return html_page


# ═══════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="构建表格 QC 对照浏览器 HTML 页面 (v2)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""示例:
  %(prog)s --docx 表格附件.docx --md QC报告-汇总.md --mapping 表格-清单-映射表.json --excel-dir 表格/
        """,
    )
    parser.add_argument(
        "--docx", default=None, help="表格附件路径 (.docx)"
    )
    parser.add_argument(
        "--md", default=None, help="QC报告-汇总.md 路径"
    )
    parser.add_argument(
        "--mapping", default=None,
        help="表格-清单-映射表.json 路径（必需）",
    )
    parser.add_argument(
        "--excel-dir", default=None,
        help="Excel 表格文件目录（默认 DOCX 同目录下的「表格」子目录）",
    )
    parser.add_argument(
        "--output", default=None,
        help="输出 HTML 路径（默认 DOCX 同目录下的 qc-viewer.html）",
    )
    args = parser.parse_args()

    # ── Resolve input paths ──
    docx_path = Path(args.docx) if args.docx else None
    md_path = Path(args.md) if args.md else None
    mapping_path = Path(args.mapping) if args.mapping else None
    excel_dir = Path(args.excel_dir) if args.excel_dir else None

    if not docx_path:
        candidates = sorted(
            p for p in BASE_DIR.glob("*.docx")
            if "清单" not in p.name and "方案" not in p.name
        )
        docx_path = candidates[0] if candidates else BASE_DIR / "表格附件.docx"
    if not md_path:
        md_path = BASE_DIR / "QC报告-汇总.md"
    if not mapping_path:
        # 尝试常见位置
        for candidate in [
            BASE_DIR / "表格-清单-映射表.json",
            BASE_DIR / "human_review_output" / "表格-清单-映射表.json",
        ]:
            if candidate.exists():
                mapping_path = candidate
                break
        if not mapping_path:
            print("❌ 未找到 --mapping，请指定映射表 JSON 路径")
            return
    if not excel_dir:
        excel_dir = BASE_DIR / "表格"
        if not excel_dir.exists():
            excel_dir = docx_path.parent / "表格"

    output_path = (
        Path(args.output) if args.output
        else docx_path.parent / "qc-viewer.html"
    )

    # ── Step 1: Extract tables from DOCX ──
    print("📊 Extracting tables from DOCX...")
    print(f"   📄 {docx_path}")
    tables = extract_tables_from_docx(docx_path)
    print(f"   → Found {len(tables)} tables")

    # ── Step 2: Parse QC markdown ──
    print("📝 Parsing QC markdown...")
    print(f"   📄 {md_path}")
    qc_data = parse_qc_markdown(md_path)
    print(f"   → Found {len(qc_data['pairs'])} QC pairs")

    # ── Step 3: Build Pair ↔ Table mapping from JSON ──
    global TABLE_IDX_TO_PAIR, PAIR_CONCLUSIONS, EXCEL_HTML_CACHE

    print(f"🔗 从映射表 JSON 推导 Pair↔表格关系...")
    print(f"   📄 {mapping_path}")
    mapping_entries = parse_mapping_from_json(str(mapping_path))
    print(f"   → {len(mapping_entries)} 条映射记录")

    pair_to_idx, auto_conclusions, table_idx_to_pair = build_pair_table_mapping(
        mapping_entries, tables
    )
    TABLE_IDX_TO_PAIR = table_idx_to_pair
    PAIR_CONCLUSIONS = auto_conclusions

    # Overlay TOC conclusions from QC report
    md_conclusions = parse_conclusions_from_md(md_path.read_text(encoding="utf-8"))
    if md_conclusions:
        for pid, (cls, label) in md_conclusions.items():
            PAIR_CONCLUSIONS[pid] = (cls, label)

    print(f"   → {len(pair_to_idx)} 个表格匹配到 mapping")
    qc_with_conclusion = len(
        [c for c, _ in PAIR_CONCLUSIONS.values() if c != "none"]
    )
    print(f"   → {qc_with_conclusion} 个有 QC 结论（来自 TOC）")

    # ── Step 4: Statistics ──
    qc_tables = [t for t in tables if t["index"] in TABLE_IDX_TO_PAIR]
    no_qc_tables = [t for t in tables if t["index"] not in TABLE_IDX_TO_PAIR]
    print(f"   → {len(qc_tables)} tables have QC results")
    print(f"   → {len(no_qc_tables)} tables have NO QC results")

    # ── Step 5: Load Excel tables ──
    print("📋 Loading Excel original tables...")
    print(f"   📂 {excel_dir}")
    EXCEL_HTML_CACHE = match_excel_to_tables(excel_dir, tables)
    print(f"   → {len(EXCEL_HTML_CACHE)} Excel tables converted to HTML")

    # ── Step 6: Generate HTML ──
    print("🔨 Generating HTML viewer...")
    html = build_html(tables, qc_data, EXCEL_HTML_CACHE)

    output_path.write_text(html, encoding="utf-8")
    size_kb = output_path.stat().st_size / 1024
    print(f"   → Written: {output_path} ({size_kb:.0f} KB)")
    print("✅ Done! Open the HTML file in a browser.")


if __name__ == "__main__":
    main()
