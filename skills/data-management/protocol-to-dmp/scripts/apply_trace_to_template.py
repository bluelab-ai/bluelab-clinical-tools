#!/usr/bin/env python3
"""Apply confirmed DMP trace values to a copy of the Word template."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Iterable

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


LABEL_FIELDS = {
    "版本号：": "DMP版本号",
    "版本日期：": "DMP版本日期",
    "方案编号：": "方案编号",
    "方案版本号：": "方案版本号",
    "方案版本日期：": "方案版本日期",
}

TRIAL_OVERVIEW_FIELDS = {
    "研究名称",
    "研究设计",
    "研究目的",
    "样本量",
    "主要有效性终点",
    "其他终点",
    "统计分析人群",
}

TARGETED_SENTENCE_REPLACEMENTS = [
    ("本项目方案对于阶段性分析的目的和阶段要求为：XXXX", "阶段性分析目的和阶段要求"),
    ("针对需要进行阶段性分析的数据，在递交时需要完成的工作有：", "阶段性分析数据准备要求"),
    ("本项目涉及的外部数据为xxx", "设计的外部数据类型"),
    ("本研究设计类型为xxx", "研究设计"),
]

TEXT_PLACEHOLDERS = {
    "请输入版本号": "DMP版本号",
    "请输入版本日期": "DMP版本日期",
    "请输入临床试验方案名称": "临床试验方案名称",
    "请输入临床试验方案编号": "方案编号",
    "请输入临床试验方案版本号": "方案版本号",
    "请输入方案版本日期": "方案版本日期",
    "请输入申办者名称": "申办者名称",
    "请输入临床监查方名称": "临床监查方名称",
    "请输入临床监察方名称": "临床监查方名称",
    "请输入统计分析单位名称": "统计分析单位名称",
    "请输入统计分析单位": "统计分析单位名称",
    "请输入数据管理单位名称": "数据管理单位名称",
    "请输入申办者": "申办者名称",
    "请输入临床监查方": "临床监查方名称",
    "请输入临床监察方": "临床监查方名称",
}

ALL_TEMPLATE_MARKER_RE = re.compile(r"^\s*/?\*?/?\s*[模|模板版]+[一二三四五六七0-9１２３４５６７]+")

FIELD_ALIASES = {
    "DMP版本号": ["版本号"],
    "DMP版本日期": ["版本日期"],
    "临床试验方案名称": ["方案名称", "研究名称"],
    "申办者名称": ["申办方名称"],
    "临床监查方名称": ["临床监察方名称", "临床监查方", "临床监察方"],
    "临床监察方名称": ["临床监查方名称", "临床监查方", "临床监察方"],
    "数据管理单位名称": ["数据管理单位"],
    "统计分析单位名称": ["统计分析单位", "统计分析方名称"],
    "版本修订记录": ["版本修订内容"],
    "撰写者修订者": ["撰写者/修订者"],
}


def norm(value: str) -> str:
    return re.sub(r"[\s　：:，,。；;（）()、/\\_\-]+", "", str(value)).lower()


def value_for(values: dict[str, str], item: str) -> str | None:
    if values.get(item):
        return values[item]
    aliases = FIELD_ALIASES.get(item, [])
    for alias in aliases:
        if values.get(alias):
            return values[alias]
    target_norms = {norm(item)}
    target_norms.update(norm(alias) for alias in aliases)
    for key, value in values.items():
        if value and norm(key) in target_norms:
            return value
    return None


def _build_hint_reverse_map() -> dict[str, str]:
    """Build a reverse mapping from normalized hint text (after 请输入) to item key."""
    hint_map: dict[str, str] = {}
    for item_key, aliases in FIELD_ALIASES.items():
        for hint in [item_key] + aliases:
            hint_map[norm(hint)] = item_key
    return hint_map


_HINT_REVERSE_MAP = _build_hint_reverse_map()
_RE_PLEASE_INPUT = re.compile(r"请输入[^\s，。,\.；;：:）\)】\]]+")


def cell_text(cell) -> str:
    return cell.text.strip().replace("\n", " | ")


def block_text(block) -> str:
    return "".join(t.text or "" for t in block.iter(qn("w:t"))).strip()


def body_children(doc) -> list:
    return list(doc.element.body.iterchildren())


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_range(doc, start: int, end: int) -> None:
    children = body_children(doc)
    for element in children[start:end]:
        if element.tag != qn("w:sectPr"):
            remove_element(element)


def find_block_index(doc, text: str, start: int = 0) -> int | None:
    for index, child in enumerate(body_children(doc)[start:], start=start):
        if block_text(child) == text:
            return index
    return None


def find_next_block_index(doc, texts: set[str], start: int = 0) -> int | None:
    for index, child in enumerate(body_children(doc)[start:], start=start):
        if block_text(child) in texts:
            return index
    return None


def value_is_yes(value: str | None) -> bool:
    if not value:
        return False
    text = str(value)
    if any(word in text for word in ["否", "无", "不使用", "不涉及", "没有", "不适用"]):
        return False
    return any(word in text for word in ["是", "有", "使用", "涉及", "高标准"])


def data_mode(values: dict[str, str]) -> str | None:
    value = (value_for(values, "项目数据采集模式：EDC / PDC") or "").upper()
    if "EDC" in value:
        return "EDC"
    if "PDC" in value:
        return "PDC"
    return None


def choose_project_type_marker(values: dict[str, str]) -> str | None:
    value = value_for(values, "项目类型：药物 / 器械") or ""
    if "器械" in value:
        return "模板2：适用于器械项目，"
    if "药物" in value:
        return "模板1：适用于药物项目，"
    return None


def choose_edc_system_marker(values: dict[str, str]) -> str | None:
    value = value_for(values, "EDC系统供应商/系统类型") or ""
    if "太美" in value and re.search(r"V\s*6|Version\s*6|6\.0", value, re.I):
        return "/模板5/太美系统V6"
    if "太美" in value and re.search(r"V\s*5|Version\s*5|5\.0", value, re.I):
        return "/模板4/太美系统V5"
    if "赛美斯" in value or "CIMS" in value:
        return "/模板1/赛美斯系统"
    if "青蜂" in value or "医墨" in value:
        return "/模板2/青蜂系统"
    if "里恩" in value:
        return "/模板3/里恩系统"
    if "易迪希" in value or "Clinflash" in value:
        return "/模板6/易迪希系统"
    if value:
        return "/模板7/其他系统"
    return None


def choose_random_system_marker(values: dict[str, str]) -> str | None:
    value = value_for(values, "随机系统供应商/系统类型") or ""
    if "医墨" in value or "医兰德" in value:
        return "/模板1/医墨随机系统"
    if "易迪希" in value:
        return "/模板2/易迪希随机系统"
    if "赛美斯" in value:
        return "/模板3/赛美斯随机系统"
    if value:
        return "/模板4/其他随机系统"
    return None


def select_pair_after_heading(
    doc,
    heading: str,
    first_marker: str,
    second_marker: str,
    keep_marker: str,
    stop_texts: set[str],
    applied: list[str],
) -> None:
    heading_index = find_block_index(doc, heading)
    if heading_index is None:
        return
    first_index = find_block_index(doc, first_marker, heading_index + 1)
    if first_index is None:
        return
    second_index = find_block_index(doc, second_marker, first_index + 1)
    if second_index is None:
        return
    stop_index = find_next_block_index(doc, stop_texts, second_index + 1)
    if stop_index is None:
        stop_index = len(body_children(doc)) - 1

    if keep_marker == first_marker:
        remove_range(doc, second_index, stop_index)
        selected_index = find_block_index(doc, first_marker, heading_index + 1)
        if selected_index is not None:
            remove_range(doc, selected_index, selected_index + 1)
    elif keep_marker == second_marker:
        remove_range(doc, first_index, second_index)
        selected_index = find_block_index(doc, second_marker, heading_index + 1)
        if selected_index is not None:
            remove_range(doc, selected_index, selected_index + 1)
    applied.append(f"模板选择 `{heading}` -> `{keep_marker}`")


def select_alternatives_after_heading(
    doc,
    heading: str,
    markers: list[str],
    keep_marker: str,
    stop_texts: set[str],
    applied: list[str],
) -> None:
    heading_index = find_block_index(doc, heading)
    if heading_index is None:
        return
    positions: list[tuple[int, str]] = []
    search_start = heading_index + 1
    for marker in markers:
        index = find_block_index(doc, marker, search_start)
        if index is not None:
            positions.append((index, marker))
            search_start = index + 1
    if not positions:
        return

    spans: list[tuple[int, int, str]] = []
    for pos_index, (start, marker) in enumerate(positions):
        if pos_index + 1 < len(positions):
            end = positions[pos_index + 1][0]
        else:
            end = find_next_block_index(doc, stop_texts, start + 1) or len(body_children(doc)) - 1
        spans.append((start, end, marker))

    for start, end, marker in reversed(spans):
        if marker != keep_marker:
            remove_range(doc, start, end)

    selected_index = find_block_index(doc, keep_marker, heading_index + 1)
    if selected_index is not None:
        remove_range(doc, selected_index, selected_index + 1)
    applied.append(f"模板选择 `{heading}` -> `{keep_marker}`")


def select_simple_marker_block(
    doc,
    marker: str,
    keep: bool,
    stop_texts: set[str],
    applied: list[str],
    start_after: str | None = None,
) -> None:
    start = 0
    if start_after:
        index = find_block_index(doc, start_after)
        if index is not None:
            start = index + 1
    marker_index = find_block_index(doc, marker, start)
    if marker_index is None:
        return
    if keep:
        remove_range(doc, marker_index, marker_index + 1)
        applied.append(f"模板标记清理 `{marker}`")
        return
    stop_index = find_next_block_index(doc, stop_texts, marker_index + 1) or len(body_children(doc)) - 1
    remove_range(doc, marker_index, stop_index)
    applied.append(f"模板删除 `{marker}`")


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        set_paragraph_text(cell.paragraphs[0], text)
        for paragraph in cell.paragraphs[1:]:
            set_paragraph_text(paragraph, "")
    else:
        cell.text = text


def align_cell_left(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(10.5)


def align_table_left(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            align_cell_left(cell)


def iter_all_paragraphs(doc) -> Iterable:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for part in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            for paragraph in part.paragraphs:
                yield paragraph
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def iter_standalone_paragraphs(doc) -> Iterable:
    for paragraph in doc.paragraphs:
        yield paragraph
    for section in doc.sections:
        for part in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            for paragraph in part.paragraphs:
                yield paragraph


def iter_all_text_elements(doc) -> Iterable:
    yield from doc.element.iter(qn("w:t"))
    for section in doc.sections:
        for part in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            yield from part._element.iter(qn("w:t"))


def iter_all_tables(doc) -> Iterable:
    for table in doc.tables:
        yield table
    for section in doc.sections:
        for part in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            for table in part.tables:
                yield table


def confirmed_values(trace: dict) -> dict[str, str]:
    values: dict[str, str] = {}
    for item in trace.get("items", []):
        if item.get("item") == "页眉":
            continue
        if item.get("status") == "filled" and item.get("value"):
            values[item["item"]] = str(item["value"]).strip()

    metadata = trace.get("metadata", {})
    selection = metadata.get("template_selection", {})
    for item in ["是否使用随机系统", "是否使用登记系统"]:
        if selection.get(item):
            values[item] = str(selection[item]).strip()

    version_records = metadata.get("version_records") or []
    if version_records:
        latest = version_records[-1]
        if latest.get("版本号"):
            values.setdefault("DMP版本号", str(latest["版本号"]).strip())
        if latest.get("版本日期"):
            values.setdefault("DMP版本日期", str(latest["版本日期"]).strip())
        if latest.get("撰写者/修订者"):
            values.setdefault("撰写者修订者", str(latest["撰写者/修订者"]).strip())
        if latest.get("版本修订内容"):
            values.setdefault("版本修订记录", str(latest["版本修订内容"]).strip())
    return values


def fill_label_paragraphs(doc, values: dict[str, str], applied: list[str]) -> None:
    for paragraph in iter_standalone_paragraphs(doc):
        text = paragraph.text.strip()
        for label, item in LABEL_FIELDS.items():
            value = value_for(values, item)
            if value and text == label:
                set_paragraph_text(paragraph, f"{label}{value}")
                applied.append(f"段落标签 `{label}` <- {item}")


def fill_label_cells(doc, values: dict[str, str], applied: list[str]) -> None:
    for table in iter_all_tables(doc):
        for row in table.rows:
            for cell_index, cell in enumerate(row.cells):
                if cell_index + 1 < len(row.cells) and cell_text(row.cells[cell_index + 1]):
                    continue
                text = cell_text(cell)
                for label, item in LABEL_FIELDS.items():
                    value = value_for(values, item)
                    if value and text == label:
                        set_cell_text(cell, f"{label}{value}")
                        applied.append(f"表格标签 `{label}` <- {item}")


def fill_text_placeholders(doc, values: dict[str, str], applied: list[str]) -> None:
    for text_element in iter_all_text_elements(doc):
        original = text_element.text or ""
        new_text = original
        applied_items: list[str] = []
        dmp_version = value_for(values, "DMP版本号")
        dmp_date = value_for(values, "DMP版本日期")
        if dmp_version and dmp_date and "请输入版本，请输入版本日期" in new_text:
            new_text = new_text.replace("请输入版本，请输入版本日期", f"{dmp_version}，{dmp_date}")
            applied_items.extend(["DMP版本号", "DMP版本日期"])
        for placeholder, item in TEXT_PLACEHOLDERS.items():
            value = value_for(values, item)
            if value and placeholder in new_text:
                new_text = new_text.replace(placeholder, value)
                applied_items.append(item)
        if new_text != original:
            remaining = new_text
            for match in _RE_PLEASE_INPUT.finditer(new_text):
                hint = match.group()[4:]
                item_key = _HINT_REVERSE_MAP.get(norm(hint))
                if item_key:
                    value = value_for(values, item_key)
                    if value:
                        remaining = remaining.replace(match.group(), value)
                        applied_items.append(item_key)
            new_text = remaining
        if new_text != original:
            text_element.text = new_text
            applied.append("文本占位替换 <- " + ",".join(dict.fromkeys(applied_items)))


def fill_signature_writer(doc, values: dict[str, str], applied: list[str]) -> None:
    writer = value_for(values, "撰写者修订者")
    if not writer:
        return
    replacements = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        stripped = text.strip()
        if "请输入姓名" not in stripped:
            continue
        if re.search(r"(审核人|复核人|批准人|审批人)[：:]\s*请输入姓名", stripped):
            continue
        if re.search(r"(撰写人|撰写者|修订者|撰写者/修订者)[：:]\s*请输入姓名", stripped):
            new_text = re.sub(
                r"((?:撰写人|撰写者|修订者|撰写者/修订者)[：:])\s*请输入姓名",
                lambda match: match.group(1) + writer,
                text,
            )
        elif stripped in {"请输入姓名", "：请输入姓名", ":请输入姓名"}:
            new_text = "撰写人：" + writer
        else:
            new_text = text.replace("请输入姓名", writer)
        new_text = re.sub(r"(撰写人[：:])\s*撰写人[：:]\s*", r"\1", new_text)
        set_paragraph_text(paragraph, new_text)
        replacements += 1
    if replacements:
        applied.append(f"签署页撰写人 <- 撰写者修订者 ({replacements}处)")


def fill_trial_overview(doc, values: dict[str, str], applied: list[str]) -> None:
    for table_index, table in enumerate(doc.tables):
        is_overview_table = any(
            len(row.cells) >= 2 and cell_text(row.cells[0]) in TRIAL_OVERVIEW_FIELDS
            for row in table.rows
        )
        if is_overview_table:
            align_table_left(table)
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            key = cell_text(row.cells[0])
            value = value_for(values, key)
            if key in TRIAL_OVERVIEW_FIELDS and value:
                existing = cell_text(row.cells[1])
                if existing != value:
                    set_cell_text(row.cells[1], value)
                    align_table_left(table)
                    applied.append(f"试验概述表 `{key}` <- {key} (table {table_index + 1})")
        if is_overview_table:
            applied.append(f"试验概述表左对齐 (table {table_index + 1})")


def remove_table_row(table, row) -> None:
    table._tbl.remove(row._tr)


def revision_records_from_trace(trace: dict, values: dict[str, str]) -> list[dict[str, str]]:
    records = trace.get("metadata", {}).get("version_records") or []
    normalized: list[dict[str, str]] = []
    for record in records:
        normalized.append(
            {
                "版本号": str(record.get("版本号", "")).strip(),
                "版本日期": str(record.get("版本日期", "")).strip(),
                "撰写者/修订者": str(record.get("撰写者/修订者", "")).strip(),
                "版本修订内容": str(record.get("版本修订内容", "")).strip(),
            }
        )
    if normalized:
        normalized.sort(key=lambda r: (r.get("版本日期", ""), r.get("版本号", "")))
        return normalized

    fallback_version = value_for(values, "DMP版本号") or ""
    fallback_date = value_for(values, "DMP版本日期") or ""
    fallback_author = value_for(values, "撰写者修订者") or ""
    fallback_content = value_for(values, "版本修订记录") or ""
    fallback = {
        "版本号": fallback_version,
        "版本日期": fallback_date,
        "撰写者/修订者": fallback_author,
        "版本修订内容": fallback_content,
    }
    return [fallback] if any(fallback.values()) else []


def fill_revision_table(doc, trace: dict, values: dict[str, str], applied: list[str]) -> None:
    records = revision_records_from_trace(trace, values)
    if not records:
        return

    for table_index, table in enumerate(doc.tables):
        if not table.rows:
            continue
        headers = [cell_text(cell) for cell in table.rows[0].cells]
        if headers[:4] != ["版本号", "版本日期", "撰写者/修订者", "版本修订内容"]:
            continue
        target_row_count = len(records) + 1
        while len(table.rows) < target_row_count:
            table.add_row()
        while len(table.rows) > target_row_count:
            remove_table_row(table, table.rows[-1])

        for index, record in enumerate(records, start=1):
            target_row = table.rows[index]
            set_cell_text(target_row.cells[0], record.get("版本号", ""))
            set_cell_text(target_row.cells[1], record.get("版本日期", ""))
            set_cell_text(target_row.cells[2], record.get("撰写者/修订者", ""))
            set_cell_text(target_row.cells[3], record.get("版本修订内容", ""))
        # 版本修订记录表格式：小四(12pt)、左对齐
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
        applied.append(f"版本修订记录表 <- DM日志版本记录 {len(records)} 条 (table {table_index + 1})")


def extract_design_type(design_full: str) -> str:
    """从完整研究设计句式中提取纯类型描述。

    "本研究是一项前瞻性、多中心、单组、目标值法的临床试验。" -> "前瞻性、多中心、单组、目标值法"
    "本试验为随机、平行对照、多中心的开放性临床试验。" -> "随机、平行对照、多中心、开放性"
    """
    result = design_full.strip().rstrip("。；;，,.")
    for prefix in ["本研究是一项", "本试验是一项", "本研究为", "本试验为", "本研究是", "本试验是"]:
        if result.startswith(prefix):
            result = result[len(prefix):]
            break
    for suffix in ["的临床试验", "的临床研究", "临床试验", "临床研究"]:
        if result.endswith(suffix):
            result = result[:-len(suffix)]
            break
    return result.strip("。；;，,. ")


def fill_targeted_sentences(doc, values: dict[str, str], applied: list[str]) -> None:
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        new_text = text
        design_value = value_for(values, "研究设计")
        if design_value and re.search(r"研究设计类型为\s*xxx", new_text):
            design_type = extract_design_type(design_value)
            new_text = re.sub(r"研究设计类型为\s*xxx", f"研究设计类型为{design_type}", new_text)
            applied.append("定向句子 `研究设计类型为xxx` <- 研究设计")
        for marker, item in TARGETED_SENTENCE_REPLACEMENTS:
            value = value_for(values, item)
            if value and marker in new_text:
                if marker.endswith("XXXX"):
                    new_text = new_text.replace(marker, marker.replace("XXXX", value))
                elif marker.endswith("xxx"):
                    new_text = new_text.replace(marker, marker.replace("xxx", value))
                elif marker.endswith("有："):
                    new_text = new_text.replace(marker, marker + value)
                else:
                    new_text = new_text.replace(marker, marker.replace("xxx", value))
                applied.append(f"定向句子 `{marker}` <- {item}")
        if new_text != text:
            set_paragraph_text(paragraph, new_text)


def clean_toc_template_labels(doc, values: dict[str, str], applied: list[str]) -> None:
    use_registry = value_is_yes(value_for(values, "是否使用登记系统"))
    use_random = value_is_yes(value_for(values, "是否使用随机系统"))
    if use_registry and use_random:
        section_8_title = "登记系统/随机系统"
    elif use_registry:
        section_8_title = "登记系统"
    elif use_random:
        section_8_title = "随机系统"
    else:
        section_8_title = "登记系统/随机系统"

    replacements = {
        "8\t/模板1/登记系统  /模板2/随机系统\t18": f"8\t{section_8_title}\t18",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])
            applied.append("目录模板标记清理 <- 是否使用登记系统,是否使用随机系统")


def fill_inline_table_template_options(doc, values: dict[str, str], applied: list[str]) -> None:
    use_pv = value_is_yes(value_for(values, "是否涉及针对有药物警戒系统的项目"))
    replacements = {
        "/模板1/按照《数据核查计划》执行\n/模板2/PV数据库及eCRF定稿后3个工作日": (
            "PV数据库及eCRF定稿后3个工作日" if use_pv else "按照《数据核查计划》执行"
        ),
        "/模板1/按照《数据核查计划》执行\n/模板2/按照《SAE一致性核查计划》执行": (
            "按照《SAE一致性核查计划》执行" if use_pv else "按照《数据核查计划》执行"
        ),
    }
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text in replacements:
                    set_cell_text(cell, replacements[text])
                    applied.append("表格内模板选项 <- 是否涉及针对有药物警戒系统的项目")
                    continue
                normalized = [line.strip() for line in text.splitlines() if line.strip()]
                if len(normalized) == 4 and normalized[0] == "/模板1/" and normalized[2] == "/模板2/":
                    first_option = normalized[1]
                    second_option = normalized[3]
                    if first_option == "按照《数据核查计划》执行" and second_option in {
                        "PV数据库及eCRF定稿后3个工作日",
                        "按照《SAE一致性核查计划》执行",
                    }:
                        set_cell_text(cell, second_option if use_pv else first_option)
                        applied.append("表格内模板选项 <- 是否涉及针对有药物警戒系统的项目")


def apply_template_selection(doc, values: dict[str, str], applied: list[str]) -> None:
    project_type_marker = choose_project_type_marker(values)
    if project_type_marker:
        applied.append("模板判断已处理 <- 项目类型：药物 / 器械")
        select_pair_after_heading(
            doc,
            "遵循的法规和支持性文件",
            "模板1：适用于药物项目，",
            "模板2：适用于器械项目，",
            project_type_marker,
            {"试验概述"},
            applied,
        )

    mode = data_mode(values)

    if mode == "EDC":
        applied.append("模板判断已处理 <- 项目数据采集模式：EDC / PDC")
        select_pair_after_heading(
            doc,
            "项目人员职责",
            "/*模版1*/EDC项目适用",
            "/*模版2*/PDC项目适用",
            "/*模版1*/EDC项目适用",
            {"数据采集/管理系统"},
            applied,
        )
        select_pair_after_heading(
            doc,
            "数据采集方式",
            "/模版1/适用于PDC项目",
            "/模版2/适用于EDC项目",
            "/模版2/适用于EDC项目",
            {"数据库项目创建"},
            applied,
        )
        for heading in ["CRF设计/审核", "eCRF/ 注释CRF设计", "eCRF填写指南"]:
            select_pair_after_heading(
                doc,
                heading,
                "/模版1/适用于PDC项目",
                "/模版2/适用于EDC项目",
                "/模版2/适用于EDC项目",
                {"eCRF/ 注释CRF设计", "eCRF填写指南", "数据核查计划"},
                applied,
            )
        for heading, stop in [
            ("EDC系统基本信息设置", {"EDC动态规则设置"}),
            ("EDC动态规则设置", {"EDC核查规则配置"}),
            ("系统培训和账号管理", {"UAT"}),
            ("UAT", {"数据库上线"}),
            ("数据库上线", {"数据库更新"}),
            ("数据库更新", {"数据录入"}),
            ("数据录入", {"数据质疑"}),
            ("数据质疑", {"外部数据管理"}),
            ("沟通频率及进度报告", {"提供数据管理单位文件"}),
            ("数据库锁定", {"数据库解锁"}),
            ("数据库解锁", {"向统计部门数据递交"}),
        ]:
            select_pair_after_heading(
                doc,
                heading,
                "/*模版1：PDC项目适用)*/",
                "/*模版2：EDC项目适用)*/",
                "/*模版2：EDC项目适用)*/",
                stop,
                applied,
            )
            select_pair_after_heading(
                doc,
                heading,
                "/*模版1(PDC项目适用)*/",
                "/*模版2(EDC项目适用)*/",
                "/*模版2(EDC项目适用)*/",
                stop,
                applied,
            )
            select_pair_after_heading(
                doc,
                heading,
                "/*模版1*/(PDC项目适用)",
                "/*模版2*/(EDC项目适用)",
                "/*模版2*/(EDC项目适用)",
                stop,
                applied,
            )
            select_pair_after_heading(
                doc,
                heading,
                "/*模版1*/（适用于PDC项目）",
                "/*模版2*/（适用于EDC项目）",
                "/*模版2*/（适用于EDC项目）",
                stop,
                applied,
            )
    elif mode == "PDC":
        applied.append("模板判断已处理 <- 项目数据采集模式：EDC / PDC")
        select_pair_after_heading(
            doc,
            "项目人员职责",
            "/*模版1*/EDC项目适用",
            "/*模版2*/PDC项目适用",
            "/*模版2*/PDC项目适用",
            {"数据采集/管理系统"},
            applied,
        )
        select_pair_after_heading(
            doc,
            "数据采集方式",
            "/模版1/适用于PDC项目",
            "/模版2/适用于EDC项目",
            "/模版1/适用于PDC项目",
            {"数据库项目创建"},
            applied,
        )

    edc_marker = choose_edc_system_marker(values)
    if edc_marker:
        applied.append("模板判断已处理 <- EDC系统供应商/系统类型")
        select_alternatives_after_heading(
            doc,
            "数据管理系统",
            [
                "/模板1/赛美斯系统",
                "/模板2/青蜂系统",
                "/模板3/里恩系统",
                "/模板4/太美系统V5",
                "/模板5/太美系统V6",
                "/模板6/易迪希系统",
                "/模板7/其他系统",
            ],
            edc_marker,
            {"用户权限定义"},
            applied,
        )

    use_registry = value_is_yes(value_for(values, "是否使用登记系统"))
    use_random = value_is_yes(value_for(values, "是否使用随机系统"))
    if value_for(values, "是否使用登记系统"):
        applied.append("模板判断已处理 <- 是否使用登记系统")
    if value_for(values, "是否使用随机系统"):
        applied.append("模板判断已处理 <- 是否使用随机系统")
    select_simple_marker_block(doc, "/模板1/登记系统  /模板2/随机系统", True, {"登记系统"}, applied)
    select_simple_marker_block(doc, "/模板1/单组登记系统", use_registry, {"随机系统", "/模板2/随机系统"}, applied)
    select_simple_marker_block(doc, "/模板2/随机系统", use_random, {"/模板3/无随机、单组登记"}, applied)
    if use_random:
        random_marker = choose_random_system_marker(values)
        if random_marker:
            applied.append("模板判断已处理 <- 随机系统供应商/系统类型")
            select_alternatives_after_heading(
                doc,
                "随机系统",
                [
                    "/模板1/医墨随机系统",
                    "/模板2/易迪希随机系统",
                    "/模板3/赛美斯随机系统",
                    "/模板4/其他随机系统",
                ],
                random_marker,
                {"用户权限定义"},
                applied,
            )
        select_simple_marker_block(doc, "/模板3/无随机、单组登记", False, {"数据管理里程碑"}, applied)
    elif use_registry:
        select_simple_marker_block(doc, "/模板3/无随机、单组登记", False, {"数据管理里程碑"}, applied)

    select_pair_after_heading(
        doc,
        "外部数据管理",
        "/*模版1*/",
        "/*模版2*/",
        "/*模版2*/" if value_is_yes(value_for(values, "是否涉及外部数据")) else "/*模版1*/",
        {"医学编码"},
        applied,
    )
    if value_for(values, "是否涉及外部数据"):
        applied.append("模板判断已处理 <- 是否涉及外部数据")
    select_pair_after_heading(
        doc,
        "医学编码",
        "/*模版1*/（含医学编码）",
        "/*模版2*/（不适用）",
        "/*模版1*/（含医学编码）" if value_is_yes(value_for(values, "是否涉及医学编码")) else "/*模版2*/（不适用）",
        {"SAE一致性核查"},
        applied,
    )
    if value_for(values, "是否涉及医学编码"):
        applied.append("模板判断已处理 <- 是否涉及医学编码")
    select_pair_after_heading(
        doc,
        "SAE一致性核查",
        "/*模版1*/针对有药物警戒系统的项目，注意：爱美客项目由PV部门进行一致性核查，及得需改负责方名称。",
        "/*模版2*/针对无药物警戒系统的项目",
        "/*模版1*/针对有药物警戒系统的项目，注意：爱美客项目由PV部门进行一致性核查，及得需改负责方名称。"
        if value_is_yes(value_for(values, "是否涉及针对有药物警戒系统的项目"))
        else "/*模版2*/针对无药物警戒系统的项目",
        {"沟通频率及进度报告"},
        applied,
    )
    if value_for(values, "是否涉及针对有药物警戒系统的项目"):
        applied.append("模板判断已处理 <- 是否涉及针对有药物警戒系统的项目")
    select_pair_after_heading(
        doc,
        "阶段性分析",
        "/*模版1*/",
        "/*模版2*/",
        "/*模版2*/" if value_is_yes(value_for(values, "是否有阶段性分析/中期分析")) else "/*模版1*/",
        {"数据库锁定与解锁"},
        applied,
    )
    if value_for(values, "是否有阶段性分析/中期分析"):
        applied.append("模板判断已处理 <- 是否有阶段性分析/中期分析")
    select_pair_after_heading(
        doc,
        "数据管理报告",
        "/模板1/无此服务范围",
        "/模板2/有此服务范围",
        "/模板2/有此服务范围" if value_is_yes(value_for(values, "是否需要数据管理报告")) else "/模板1/无此服务范围",
        {"向申办者数据递交"},
        applied,
    )
    if value_for(values, "是否需要数据管理报告"):
        applied.append("模板判断已处理 <- 是否需要数据管理报告")
    if value_for(values, "是否需要预递交"):
        applied.append("模板判断已处理 <- 是否需要预递交")
    if value_for(values, "是否包含向申办者数据递交服务范围"):
        applied.append("模板判断已处理 <- 是否包含向申办者数据递交服务范围")
    qc_value = value_for(values, "项目质量控制等级/模板") or ""
    if qc_value:
        applied.append("模板判断已处理 <- 项目质量控制等级/模板")
    select_pair_after_heading(
        doc,
        "数据质量控制",
        "/*模版1*/（适用于高标准项目）",
        "/*模版2*/（适用于低标准项目）",
        "/*模版1*/（适用于高标准项目）" if "高" in qc_value else "/*模版2*/（适用于低标准项目）",
        {"变更管理"},
        applied,
    )


def unresolved_items(trace: dict) -> list[dict]:
    return [
        item
        for item in trace.get("items", [])
        if item.get("status") in {"missing", "uncertain", "conflict", "manual_confirm", "not_processed"}
        and item.get("missing_handling") != "NA"
    ]


def write_report(path: Path, trace: dict, values: dict[str, str], applied: list[str], doc) -> None:
    applied_items = set()
    for line in applied:
        for item in values:
            if item in line:
                applied_items.add(item)
        if line.startswith("版本修订记录表 <-"):
            applied_items.update({"DMP版本号", "DMP版本日期", "版本号", "版本日期", "版本修订记录", "撰写者修订者"})
        if line == "模板判断已处理 <- 项目数据采集模式：EDC / PDC":
            applied_items.add("数据录入和质疑模板")

    filled_not_applied = sorted(set(values) - applied_items)
    unresolved = unresolved_items(trace)

    lines = [
        "# DMP生成报告",
        "",
        f"- 模板：{trace.get('metadata', {}).get('template', '')}",
        f"- 清单：{trace.get('metadata', {}).get('checklist', '')}",
        f"- Word表格数量：{len(doc.tables)}",
        f"- 自动应用项数量：{len(applied)}",
        f"- 已确认但未自动定位项数量：{len(filled_not_applied)}",
        f"- 待确认/未处理项数量：{len(unresolved)}",
        "",
        "## 自动应用项",
    ]
    lines.extend([f"- {line}" for line in applied] or ["- 无"])
    lines.extend(["", "## 已确认但需人工按模板规则处理"])
    lines.extend([f"- {item}: {values[item]}" for item in filled_not_applied] or ["- 无"])
    lines.extend(["", "## 待确认/未处理项"])
    for item in unresolved:
        key = item.get("key", item.get("item", ""))
        lines.append(f"- [{item.get('seq')}] {item.get('section')} - {item.get('item')} ({item.get('status')})  key=`{key}`")
        if item.get("question"):
            lines.append(f"  {item['question']}")
    if not unresolved:
        lines.append("- 无")

    path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Apply confirmed DMP trace values to a copied Word template.")
    parser.add_argument("--template", type=Path)
    parser.add_argument("--trace", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    parser.add_argument("--report", type=Path)
    args = parser.parse_args()

    from docx import Document

    trace = json.loads(args.trace.read_text(encoding="utf-8"))
    values = confirmed_values(trace)
    template_value = args.template or trace.get("metadata", {}).get("template")
    if not template_value:
        raise SystemExit("未提供 --template，且 trace metadata 中没有 template。")
    template_path = Path(template_value)
    doc = Document(template_path)
    applied: list[str] = []

    apply_template_selection(doc, values, applied)
    clean_toc_template_labels(doc, values, applied)
    fill_inline_table_template_options(doc, values, applied)
    fill_text_placeholders(doc, values, applied)
    fill_signature_writer(doc, values, applied)
    fill_label_paragraphs(doc, values, applied)
    fill_label_cells(doc, values, applied)
    fill_trial_overview(doc, values, applied)
    fill_revision_table(doc, trace, values, applied)
    fill_targeted_sentences(doc, values, applied)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)

    if args.report:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        write_report(args.report, trace, values, applied, doc)

    print(
        json.dumps(
            {"applied": len(applied), "filled_values": len(values), "template": str(template_path), "out": str(args.out)},
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
