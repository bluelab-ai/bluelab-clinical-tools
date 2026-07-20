#!/usr/bin/env python3
"""从 Word 文档中提取所有表格为 Excel 文件，按文档标题命名。

用法:
    # 同时处理表格和清单
    python3 extract_tables.py <表格.docx> <清单.docx> [输出目录]

    # 单独处理一个文件
    python3 extract_tables.py <文件.docx> [--out 输出目录]

两步走:
  1. 先遍历 body XML，建立每张表格的标题索引
  2. 再用 doc.tables 提取数据，按索引命名文件
"""

import json
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook


SECTION_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+(.+)$")
TABLE_TITLE_RE = re.compile(r"^表\s*[\d.]+\s+(.+)")


def build_table_index(docx_path, doc_type):
    """第一遍：遍历 body XML，按表格出现顺序返回每张表的 (title, parents)。

    parents: 该表所在的章节标题路径（仅 doc_type='表格' 时填充），如
        ["7. 统计分析", "7.2 人口学信息和基线资料（FAS）", "7.2.1 人口学信息"]
    用途：classify_and_rename.py 在表题未带 (FAS/PPS/SS) 时，回溯父节括号确定人群类型。
    """
    doc = Document(docx_path)
    body = doc.element.body

    entries = []
    last_title = ""
    in_toc = (doc_type == "清单")
    # 标题层级（数字段数 = 层级，如 "7"=1, "7.2"=2, "7.2.1"=3）→ 段落文本
    heading_stack = {}

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            texts = []
            for t in child.findall(".//" + qn("w:t")):
                if t.text:
                    texts.append(t.text)
            text = "".join(texts).strip()
            if not text:
                continue

            if doc_type == "表格":
                # ① 表题：表 X.X.X 标题
                m = TABLE_TITLE_RE.match(text)
                if m:
                    last_title = text
                    continue

                # ② 章节标题：x[.x.x...] 标题（不以"表"开头）
                m = SECTION_RE.match(text)
                if m:
                    num = m.group(1)
                    level = num.count(".") + 1
                    heading_stack[level] = text
                    # 进入新层级时，清掉更深层的旧记录
                    for k in [k for k in heading_stack if k > level]:
                        del heading_stack[k]
            else:
                m = re.match(r"^清单\s+(\d+)\s+(.+)", text)
                if m:
                    name = m.group(2).strip()
                    if in_toc and re.search(r"\d+$", name):
                        continue
                    in_toc = False
                    last_title = text

        elif tag == "tbl":
            parents = ([heading_stack[k] for k in sorted(heading_stack.keys())]
                       if doc_type == "表格" else [])
            entries.append({
                "title": last_title if last_title else "未命名",
                "parents": parents,
            })
            last_title = ""

    return entries


def extract_tables_from_docx(docx_path):
    """返回 docx 中所有表格的列表，每个表格是一个二维列表（行 × 列）。"""
    doc = Document(docx_path)
    tables = []
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)
        tables.append(rows_data)
    return tables


def sanitize_filename(name):
    """去掉文件名中不合法的字符，并限制长度。"""
    name = re.sub(r'[/\\:*?"<>|]', "-", name)
    if len(name) > 120:
        name = name[:120]
    return name


def save_table_to_xlsx(table_data, xlsx_path):
    """将一个二维列表保存为 Excel 文件。"""
    wb = Workbook()
    ws = wb.active
    for row_data in table_data:
        ws.append(row_data)
    wb.save(xlsx_path)


def detect_doc_type(filename):
    """从文件名推断文档类型：表格 / 清单 / 未知"""
    if "清单" in filename:
        return "清单"
    if "表格" in filename or "table" in filename.lower():
        return "表格"
    return "表格"  # 默认按表格处理


def process_one(docx_path, output_dir, doc_type=None):
    """处理单个 docx 文件，导出 Excel 到 output_dir。"""
    if not os.path.exists(docx_path):
        print(f"⚠ 文件不存在，跳过: {docx_path}")
        return

    if doc_type is None:
        doc_type = detect_doc_type(os.path.basename(docx_path))

    print(f"处理: {docx_path}  (类型: {doc_type})")

    entries = build_table_index(docx_path, doc_type)
    print(f"  索引条目: {len(entries)}")

    tables = extract_tables_from_docx(docx_path)
    print(f"  表格数: {len(tables)}")

    if len(entries) != len(tables):
        print(f"  ⚠ 索引数({len(entries)})与表格数({len(tables)})不一致，以表格数为准")

    os.makedirs(output_dir, exist_ok=True)

    meta = {}
    for i, table_data in enumerate(tables):
        idx = f"{i+1:02d}"
        if i < len(entries):
            title = entries[i]["title"]
            parents = entries[i]["parents"]
            safe_title = sanitize_filename(title)
            filename = f"{idx}-{safe_title}.xlsx"
        else:
            title = "未命名"
            parents = []
            filename = f"Table_{idx}.xlsx"

        xlsx_path = os.path.join(output_dir, filename)
        save_table_to_xlsx(table_data, xlsx_path)
        rows = len(table_data)
        cols = max((len(r) for r in table_data), default=0)
        print(f"  → {filename}  ({rows} 行 × {cols} 列)")

        meta[idx] = {"title": title, "parents": parents, "file": filename}

    # 仅 doc_type='表格' 时落盘父节路径（清单模式不需要人群类型推断）
    if doc_type == "表格":
        meta_path = os.path.join(output_dir, "tables_meta.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
        print(f"  → tables_meta.json  ({len(meta)} 条)")

    print(f"  已保存到: {output_dir}\n")


def main():
    args = sys.argv[1:]
    if not args:
        print("用法:")
        print("  python3 extract_tables.py <表格.docx> <清单.docx> [输出目录]")
        print("  python3 extract_tables.py <文件.docx> [--out 输出目录] [--type 表格|清单]")
        sys.exit(1)

    # 解析参数
    output_base = os.getcwd()
    doc_type_override = None
    files = []

    i = 0
    while i < len(args):
        a = args[i]
        if a == "--out" and i + 1 < len(args):
            i += 1
            output_base = args[i]
        elif a == "--type" and i + 1 < len(args):
            i += 1
            doc_type_override = args[i]
        else:
            files.append(a)
        i += 1

    if not files:
        print("错误: 未指定输入文件")
        sys.exit(1)

    # 双文件模式：自动分配 表格/ 和 清单/ 子目录
    if len(files) == 2 and doc_type_override is None:
        for docx_path in files:
            basename = os.path.basename(docx_path)
            doc_type = detect_doc_type(basename)
            out_dir = os.path.join(output_base, doc_type)
            process_one(docx_path, out_dir, doc_type)
    else:
        # 单文件模式
        for docx_path in files:
            doc_type = doc_type_override or detect_doc_type(os.path.basename(docx_path))
            out_dir = os.path.join(output_base, doc_type) if len(files) > 1 else output_base
            process_one(docx_path, out_dir, doc_type)

    print("完成。")


if __name__ == "__main__":
    main()
